"""Generate MCP Discovery & Strategy Document for ClassBridge."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from datetime import date


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2B579A')

    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r % 2 == 1:
                set_cell_shading(cell, 'F2F2F2')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def build_document():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # ─── TITLE PAGE ───
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('MCP Protocol Integration\nDiscovery & Strategy Document')
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(43, 87, 154)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('ClassBridge — AI-Powered Education Platform')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f'Date: {date.today().strftime("%B %d, %Y")}\nAuthor: Solutions Architecture & Design\nVersion: 1.0\nClassification: Internal')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_page_break()

    # ─── TABLE OF CONTENTS ───
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. What is MCP (Model Context Protocol)?',
        '3. Strategic Fit for ClassBridge',
        '4. Phase Timing Analysis',
        '5. Feature Discovery — What MCP Unlocks',
        '   5.1 Multi-Client Distribution',
        '   5.2 MCP Prompts — Guided Workflows',
        '   5.3 Consuming External MCP Servers',
        '   5.4 Features Only Possible with MCP',
        '6. MCP Ecosystem Landscape',
        '   6.1 MCP Client Ecosystem',
        '   6.2 Available External MCP Servers',
        '   6.3 Education-Specific MCP Servers',
        '7. Architecture Design',
        '8. Implementation Plan — Track F Breakdown',
        '   8.1 Phase A: Foundation (#904, #905)',
        '   8.2 Phase B: Core Tools (#906, #907, #908)',
        '   8.3 Phase C: Advanced (#909, #910, #911)',
        '   8.4 Phase D: Quality (#912)',
        '9. Cost Analysis',
        '   9.1 Additional Costs',
        '   9.2 Cost Savings',
        '   9.3 ROI Summary',
        '10. GitHub Issues Created',
        '11. Requirements Files Updated',
        '12. Risks & Mitigations',
        '13. Recommendations & Next Steps',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)

    doc.add_page_break()

    # ─── 1. EXECUTIVE SUMMARY ───
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This document captures the complete discovery and strategic analysis of integrating '
        'the Model Context Protocol (MCP) into ClassBridge. MCP is an open protocol by Anthropic '
        'that standardizes how AI applications connect to external data sources and tools.'
    )
    doc.add_paragraph(
        'The key finding is that MCP transforms ClassBridge from a standalone AI education tool '
        'into a contextual AI learning platform accessible from any MCP-compatible client — '
        'including Claude Desktop (19M+ MAU), GitHub Copilot (20M+ users), ChatGPT (400M+ weekly), '
        'Cursor (1M+), and Google Gemini. This is achieved by exposing ClassBridge\'s existing '
        'FastAPI endpoints as MCP tools, adding custom student context resources, and optionally '
        'consuming external MCP servers for content enrichment.'
    )

    doc.add_heading('Key Conclusions', level=2)
    bullets = [
        'Phase 1 Value: Not required for MVP, but a 1-day "plant the flag" move in Phase 1.5 is recommended (3 lines of code to mount fastapi-mcp).',
        'Competitive Window: No education platform has shipped a production MCP server yet. ClassBridge can be first.',
        'Cost: Near-zero incremental hosting cost (~$0-8/month). Development investment of ~4-6 weeks for full Track F.',
        'ROI: $61K-$135K+ in avoided integration costs (native plugins for Claude, Copilot, Cursor, ChatGPT would cost months each).',
        '26 MCP tools designed across 5 categories: student context, Google Classroom, study materials, AI tutor agent, and teacher communication.',
        '10 GitHub issues created (#903-#912) as Phase 2 Track F, fully integrated into requirements and roadmap.',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_page_break()

    # ─── 2. WHAT IS MCP ───
    doc.add_heading('2. What is MCP (Model Context Protocol)?', level=1)
    doc.add_paragraph(
        'The Model Context Protocol (MCP) is an open protocol created by Anthropic that provides '
        'a standardized way for AI applications to connect to external data sources and tools. '
        'It follows a client-server architecture where:'
    )
    items = [
        'MCP Servers expose tools (actions), resources (read-only data), and prompts (reusable templates).',
        'MCP Clients (Claude Desktop, Cursor, GitHub Copilot, etc.) connect to servers and make tools available to AI models.',
        'Transport: Uses Streamable HTTP (successor to SSE), making it compatible with serverless platforms like Cloud Run.',
        'Auth: Supports standard HTTP authentication (Bearer tokens, API keys).',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('MCP Primitives', level=2)
    add_styled_table(doc,
        ['Primitive', 'Description', 'ClassBridge Use Case'],
        [
            ['Tools', 'Actions the AI model can execute (function calls)', 'Generate study guide, list assignments, send message'],
            ['Resources', 'Read-only data the model can access', 'Student profile, assignment feed, study history, weak areas'],
            ['Prompts', 'Reusable workflow templates surfaced as slash commands', '/study_plan, /practice_quiz, /parent_briefing, /exam_prep'],
        ],
        col_widths=[3, 5.5, 7.5]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'MCP launched in November 2024. By late 2025, Anthropic reported "millions of MCP end-users" '
        'and nearly 2,000 servers in the official MCP Registry. OpenAI adopted MCP in March 2025, '
        'and Google confirmed Gemini MCP support in April 2025.'
    )

    doc.add_page_break()

    # ─── 3. STRATEGIC FIT ───
    doc.add_heading('3. Strategic Fit for ClassBridge', level=1)
    doc.add_paragraph(
        'ClassBridge currently uses the Anthropic Claude API (claude-sonnet-4-5) for AI study tool '
        'generation. The AI interaction model is "user pastes content → AI generates material." '
        'MCP transforms this into "AI has full student context → generates personalized material."'
    )

    doc.add_heading('Current vs. MCP-Enabled Workflow', level=2)
    add_styled_table(doc,
        ['Aspect', 'Current (Phase 1)', 'MCP-Enabled (Phase 2+)'],
        [
            ['AI Input', 'User-pasted text content only', 'Full student context: grades, assignments, study history, weak areas'],
            ['Interaction', 'Web UI / Mobile app only', 'Any MCP client: Claude Desktop, Copilot, Cursor, ChatGPT, voice'],
            ['Context', 'Stateless — each generation is independent', 'Contextual — AI knows what student studied before, quiz scores, trends'],
            ['Study Plans', 'Manual — user decides what to study', 'AI-generated personalized multi-day plans with auto-created tasks'],
            ['Content Enrichment', 'None — only user-provided content', 'External resources via MCP servers (web search, YouTube, Google Drive)'],
            ['Distribution', 'ClassBridge web + mobile only', 'Every MCP client worldwide (hundreds of millions of users)'],
        ],
        col_widths=[3, 5.5, 7.5]
    )

    doc.add_paragraph()
    doc.add_heading('Why MCP (Not a Custom API or Plugin System)', level=2)
    reasons = [
        'Universal compatibility: One MCP server works with Claude Desktop, Copilot, Cursor, ChatGPT, Gemini — no separate plugins needed.',
        'Zero frontend code for AI interactions: MCP clients render their own UI for tools and prompts.',
        'Community ecosystem: 1,200+ existing MCP servers that ClassBridge can consume as a client.',
        'Industry momentum: All major AI providers have adopted MCP. It is becoming the standard.',
        'Low effort: fastapi-mcp library exposes existing FastAPI endpoints with 3 lines of code.',
    ]
    for r in reasons:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_page_break()

    # ─── 4. PHASE TIMING ───
    doc.add_heading('4. Phase Timing Analysis', level=1)
    doc.add_paragraph(
        'The question of whether MCP belongs in Phase 1 or Phase 2 depends on what value it adds '
        'relative to what users need at each phase.'
    )

    doc.add_heading('Phase 1 (MVP) — Not Required', level=2)
    doc.add_paragraph(
        'Phase 1 is 99% complete. Users are still learning the web/mobile UI. MCP adds a new '
        'interaction channel, not a missing core feature. Adding it to Phase 1 would delay the '
        'pilot without clear user demand.'
    )

    doc.add_heading('Phase 1.5 (Post-Pilot) — Recommended Quick Win', level=2)
    doc.add_paragraph(
        'After pilot launch, mounting fastapi-mcp on the existing app takes ~1 day and makes '
        'ClassBridge immediately discoverable to Claude Desktop and Copilot users. This is a '
        '"plant the flag" move with near-zero risk:'
    )
    items = [
        '3 lines of code added to main.py',
        'Zero infrastructure changes (same Cloud Run instance)',
        'Read-only endpoints only (safe by default)',
        'Marketing differentiator: "First education platform with native AI assistant support"',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Phase 2 Track F — Full Implementation', level=2)
    doc.add_paragraph(
        'The complete MCP integration (auth, custom tools, AI tutor agent, external resources) '
        'is a 4-6 week investment that belongs in Phase 2. By this time, pilot users will have '
        'established usage patterns, making contextual AI features immediately valuable.'
    )

    add_styled_table(doc,
        ['Timing', 'Action', 'Effort', 'Rationale'],
        [
            ['Phase 1 (now)', 'Nothing — ship pilot as-is', '0', 'Users still learning the UI'],
            ['Phase 1.5 (post-pilot)', 'Mount fastapi-mcp (#904 only)', '~1 day', '3 lines of code, instant discoverability'],
            ['Phase 2 Track F', 'Full MCP integration (#905-#912)', '~4-6 weeks', 'Auth, custom tools, AI tutor, external resources'],
        ],
        col_widths=[3.5, 5, 2.5, 5]
    )

    doc.add_page_break()

    # ─── 5. FEATURE DISCOVERY ───
    doc.add_heading('5. Feature Discovery — What MCP Unlocks', level=1)

    doc.add_heading('5.1 Multi-Client Distribution', level=2)
    doc.add_paragraph(
        'Without MCP, ClassBridge would need to build native plugins for each AI client separately. '
        'Building integrations for even 3 clients would take months and require ongoing maintenance. '
        'With MCP, one server works everywhere.'
    )
    add_styled_table(doc,
        ['MCP Client', 'User Base', 'ClassBridge Use Case'],
        [
            ['Claude Desktop', '~19M monthly active users', '"Show my child\'s grades and upcoming assignments"'],
            ['GitHub Copilot', '20M+ users (90% of Fortune 100)', 'Teachers query classroom data while creating lesson plans'],
            ['ChatGPT (OpenAI)', '400M+ weekly users', 'Adopted MCP March 2025 — massive reach'],
            ['Google Gemini', 'Hundreds of millions', 'Confirmed MCP support April 2025'],
            ['Cursor', '1M+ users, 360K paying', 'Students query assignments while coding'],
            ['Windsurf / JetBrains', 'Millions combined', 'Full MCP support with plugin stores'],
        ],
        col_widths=[3.5, 4.5, 8]
    )

    doc.add_paragraph()
    doc.add_heading('5.2 MCP Prompts — Guided Workflows Without Frontend Code', level=2)
    doc.add_paragraph(
        'MCP\'s "Prompts" primitive lets ClassBridge define reusable slash commands that appear '
        'natively in every MCP client. Users select them from a command palette, fill in parameters '
        'via the client\'s native UI, and get structured results — without ClassBridge writing '
        'any frontend code for these interactions.'
    )
    add_styled_table(doc,
        ['ClassBridge Prompt', 'User Sees in Claude Desktop', 'What Happens'],
        [
            ['/study_plan', 'Form: Subject, Grade, Days', 'Pulls student context → generates personalized multi-day plan'],
            ['/practice_quiz', 'Form: Topic, Difficulty, # Questions', 'Generates quiz with real-time grade context'],
            ['/parent_briefing', 'Form: Child name', 'Fetches grades + assignments + teacher comms → 1-page summary'],
            ['/exam_prep', 'Form: Exam date, Subject', 'Builds countdown study schedule with daily tasks'],
            ['/review_grades', 'Form: Child, Time period', 'Fetches grades and generates trend analysis with recommendations'],
        ],
        col_widths=[3.5, 5, 7.5]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Why this is hard without MCP: Replicating this requires building a custom UI for '
        'parameterized prompt templates in the web app, then mobile, then a chatbot interface. '
        'MCP externalizes all of that — define the prompt once, every client renders its own native UI.'
    )

    doc.add_paragraph()
    doc.add_heading('5.3 Consuming External MCP Servers (Free Integrations)', level=2)
    doc.add_paragraph(
        'ClassBridge can act as an MCP client, connecting to community MCP servers to enrich '
        'AI-generated study materials without building integrations from scratch.'
    )
    add_styled_table(doc,
        ['External MCP Server', 'What ClassBridge Gets', 'Build Cost Without MCP'],
        [
            ['YouTube MCP', 'Extract full video transcripts → generate study guides from Khan Academy videos', 'Weeks (YouTube API, transcript parsing, quota management)'],
            ['Google Drive MCP', 'Read teacher\'s uploaded rubrics/notes from Drive', 'Days (OAuth, Drive API, file parsing)'],
            ['Brave Search MCP', 'Enrich study guides with supplementary web articles (2K queries/month free)', 'Days (search API integration, result ranking)'],
            ['Memory MCP', 'Persistent student learning profiles across AI sessions (knowledge graph)', 'Weeks (vector DB, knowledge graph implementation)'],
            ['Google Suite MCP', 'Drive + Docs + Sheets + Calendar in one server', 'Weeks (multi-API OAuth, file format handling)'],
        ],
        col_widths=[3.5, 6.5, 6]
    )

    doc.add_paragraph()
    doc.add_heading('5.4 Features Only Possible with MCP', level=2)
    add_styled_table(doc,
        ['Feature', 'Why MCP Is Required'],
        [
            ['"Add ClassBridge to Claude"', 'Users configure ClassBridge as a data source in Claude Desktop in 30 seconds. No app install needed.'],
            ['Voice-to-ClassBridge', 'Claude Desktop has voice input. Parents speak "what\'s my kid\'s homework tonight?" and get answers from ClassBridge API.'],
            ['Multi-tool AI composition', 'A Claude agent can combine ClassBridge data + Google Calendar + email MCP servers in one conversation — something no single app can do.'],
            ['Developer/school-IT ecosystem', 'Schools can build custom automations that query ClassBridge data via MCP without ClassBridge building a custom API.'],
            ['Consistent AI pedagogy', 'ClassBridge ships authoritative tutoring prompts that enforce "help learn, don\'t give answers" rules across all AI clients.'],
        ],
        col_widths=[4, 12]
    )

    doc.add_page_break()

    # ─── 6. MCP ECOSYSTEM ───
    doc.add_heading('6. MCP Ecosystem Landscape', level=1)

    doc.add_heading('6.1 MCP Client Ecosystem', level=2)
    doc.add_paragraph(
        'As of early 2026, MCP has been adopted by every major AI platform. The protocol launched '
        'in November 2024 and reached "millions of end-users" within its first year.'
    )
    add_styled_table(doc,
        ['Client', 'Users / Scale', 'MCP Support Status'],
        [
            ['GitHub Copilot (VS Code)', '20M+ users, 90% of Fortune 100', 'Full — agent mode with MCP rolled out to all VS Code users'],
            ['Claude Desktop / Claude Code', '~18.9M MAU; ~300M with API', 'Full — original MCP client (Anthropic created MCP)'],
            ['ChatGPT (OpenAI)', '400M+ weekly users', 'Adopted MCP March 2025'],
            ['Google Gemini', 'Hundreds of millions', 'Confirmed MCP support April 2025'],
            ['Cursor', '1M+ users, $1B ARR, 360K paying', 'Full'],
            ['Windsurf (Codeium)', 'Millions', 'Full — built-in MCP Plugin Store'],
            ['JetBrains AI Assistant', 'Tens of millions of IDE users', 'Full (since 2025.2)'],
            ['LibreChat', 'Open-source, self-hosted', 'Full'],
        ],
        col_widths=[4, 4.5, 7.5]
    )

    doc.add_paragraph()
    doc.add_heading('6.2 Available External MCP Servers', level=2)
    doc.add_paragraph(
        'The ecosystem has 1,200+ catalogued servers across directories like mcp-awesome.com, '
        'mcpservers.org, and mcp.so. Key servers relevant to ClassBridge:'
    )
    add_styled_table(doc,
        ['Category', 'Server', 'Free Tier', 'Relevance to ClassBridge'],
        [
            ['Web Search', 'Brave Search MCP', '2,000 queries/month', 'Enrich study guides with supplementary articles'],
            ['Web Search', 'Tavily Search MCP', '1,000 credits/month', 'Multi-source research aggregation for AI tutoring'],
            ['Google', 'Google Drive MCP (Anthropic official)', 'Unlimited (user\'s Drive)', 'Read teacher\'s rubrics, notes, syllabi from Drive'],
            ['Google', 'Google Suite MCP', 'Unlimited', 'Drive + Docs + Sheets + Calendar combined'],
            ['Video', 'YouTube MCP', '10K queries/day (Data API)', 'Extract transcripts, generate study guides from Khan Academy'],
            ['Memory', 'Memory MCP (Anthropic official)', 'Unlimited (local)', 'Persistent student learning profiles across sessions'],
            ['Files', 'Filesystem MCP (Anthropic official)', 'Unlimited (local)', 'Secure local file operations for school lab scenarios'],
        ],
        col_widths=[2.5, 4.5, 3, 6]
    )

    doc.add_paragraph()
    doc.add_heading('6.3 Education-Specific MCP Servers', level=2)
    doc.add_paragraph(
        'The education-specific MCP server ecosystem is early but growing. No production-grade '
        'education MCP server exists yet — this is ClassBridge\'s competitive opportunity.'
    )
    add_styled_table(doc,
        ['Server', 'What It Does', 'Maturity'],
        [
            ['EduChain MCP', 'Generates MCQs, lesson plans, flashcards via Claude Desktop', 'Community / experimental'],
            ['EduBase MCP', 'Quiz creation, exam scheduling, result analysis', 'Community'],
            ['Educational Tutor MCP', 'Grading tools, cognitive diagnosis, knowledge tracing, learning paths', 'Community'],
            ['MCP Server Learning', 'Flashcard generation with Anki integration, Zotero, Obsidian', 'Community'],
            ['Skolverket MCP', 'Swedish curriculum data — curricula, courses, school lookup', 'Community'],
            ['ClassBridge MCP (planned)', 'Production-grade: real OAuth, real student data, real institutional backing', 'Planned — Phase 2 Track F'],
        ],
        col_widths=[3.5, 8, 4.5]
    )
    doc.add_paragraph(
        'Key insight: Khan Academy\'s Khanmigo has not published an MCP server. There is no '
        'dominant player in the "education MCP server" space. ClassBridge can establish first-mover advantage.'
    )

    doc.add_page_break()

    # ─── 7. ARCHITECTURE ───
    doc.add_heading('7. Architecture Design', level=1)

    doc.add_heading('System Architecture', level=2)
    arch_text = """
┌─────────────────────────────────────────────────────────┐
│                    MCP Clients                           │
│  Claude Desktop │ Copilot │ Cursor │ ChatGPT │ Mobile   │
└────────┬────────┴────┬────┴────┬───┴────┬────┴─────────┘
         │             │         │        │
         ▼             ▼         ▼        ▼
┌─────────────────────────────────────────────────────────┐
│            ClassBridge MCP Server (/mcp)                 │
│  Transport: Streamable HTTP  │  Auth: JWT + API Key      │
│                                                          │
│  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐   │
│  │  Resources   │  │    Tools     │  │   Prompts    │   │
│  │  (4 types)   │  │  (26 tools)  │  │  (4+ types)  │   │
│  │              │  │              │  │              │   │
│  │ Profile      │  │ Study Gen    │  │ Study Plan   │   │
│  │ Assignments  │  │ Classroom    │  │ Exam Prep    │   │
│  │ History      │  │ Messages     │  │ Briefing     │   │
│  │ Weak Areas   │  │ Tutor Agent  │  │ Review       │   │
│  └──────────────┘  └──────────────┘  └──────────────┘   │
│                                                          │
│  RBAC: Role-filtered tools per user role                 │
└────────────────────────┬────────────────────────────────┘
                         │
                         ▼
┌─────────────────────────────────────────────────────────┐
│       ClassBridge as MCP Client (opt-in)                 │
│  Brave Search │ YouTube │ Google Drive │ Curriculum      │
└─────────────────────────────────────────────────────────┘
"""
    p = doc.add_paragraph()
    run = p.add_run(arch_text.strip())
    run.font.name = 'Consolas'
    run.font.size = Pt(7.5)

    doc.add_heading('Role-Based Access Control (RBAC)', level=2)
    add_styled_table(doc,
        ['Tool Category', 'Parent', 'Student', 'Teacher', 'Admin'],
        [
            ['List children + child data', 'Yes', 'No', 'No', 'Yes'],
            ['View own courses/assignments', 'Yes', 'Yes', 'Yes', 'Yes'],
            ['Generate study materials', 'Yes', 'Yes', 'Yes', 'No'],
            ['View grades/analytics', 'Yes', 'Yes', 'Yes', 'Yes'],
            ['View/send messages', 'Yes', 'No', 'Yes', 'Yes'],
            ['List students (roster)', 'No', 'No', 'Yes', 'Yes'],
            ['Admin tools', 'No', 'No', 'No', 'Yes'],
        ],
        col_widths=[5, 2, 2, 2, 2]
    )

    doc.add_paragraph()
    doc.add_heading('Tool Summary (26 Total)', level=2)
    add_styled_table(doc,
        ['Category', 'Tools', 'Count', 'Issue'],
        [
            ['Student Context', 'get_student_summary, identify_knowledge_gaps', '2', '#906'],
            ['Google Classroom', 'list_google_courses, list_course_assignments, get_course_materials, get_classroom_grades, sync_classroom_data, get_sync_status', '6', '#907'],
            ['Study Materials', 'generate_study_guide, generate_quiz, generate_flashcards, list_study_materials, get_study_material, search_study_materials, convert_study_material', '7', '#908'],
            ['AI Tutor Agent', 'create_study_plan, get_study_recommendations, analyze_study_effectiveness', '3', '#909'],
            ['Communication', 'list_conversations, get_conversation_messages, send_message, get_teacher_communication_summary, get_unread_count', '5', '#910'],
            ['Auto-exposed', '~15 read endpoints via fastapi-mcp (courses, assignments, tasks, notifications, analytics)', '~15', '#904'],
        ],
        col_widths=[3, 7.5, 1.5, 1.5]
    )

    doc.add_page_break()

    # ─── 8. IMPLEMENTATION PLAN ───
    doc.add_heading('8. Implementation Plan — Track F Breakdown', level=1)

    doc.add_heading('8.1 Phase A: Foundation (#904, #905)', level=2)
    doc.add_paragraph('Priority: P0 — All other MCP work depends on this.')

    doc.add_heading('#904 — FastAPI-MCP Server Foundation', level=3)
    items = [
        'Install fastapi-mcp (pip install fastapi-mcp)',
        'Mount MCP server at /mcp endpoint in main.py (3 lines of code)',
        'Configure MCP_BASE_URL in app/core/config.py',
        'Filter endpoints: expose safe read + AI generation routes, exclude auth/admin/delete/OAuth',
        'Add MCP status to /health endpoint',
        'Estimated effort: 1 day',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('#905 — MCP Authentication & Authorization', level=3)
    items = [
        'JWT authentication: reuse existing get_current_user() from app/api/deps.py',
        'API key authentication: new api_keys table with hashed keys, CRUD endpoints, rate limiting',
        'Role-based tool filtering: parents, students, teachers, admins see different tool sets',
        'Audit logging: all MCP tool invocations logged to audit_logs table',
        'Frontend: API key management UI on Profile page',
        'Estimated effort: 3-4 days',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('8.2 Phase B: Core Tools (#906, #907, #908)', level=2)
    doc.add_paragraph('Priority: P1 — The core value of MCP integration.')

    doc.add_heading('#906 — Student Academic Context', level=3)
    items = [
        '4 MCP resources: student profile, assignment feed, study history, weak areas',
        '2 MCP tools: get_student_summary, identify_knowledge_gaps',
        'Resources cached with 5-minute TTL',
        'Role-based access: parent→own children, student→self, teacher→enrolled',
        'Estimated effort: 3-4 days',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('#907 — Google Classroom MCP Tools', level=3)
    items = [
        '6 MCP tools wrapping existing Google Classroom integration',
        'Tools query local DB (synced data), not Google API directly',
        'sync_classroom_data triggers actual Google API call (rate limited: 1/5min)',
        'Clear error when Google is not connected',
        'Estimated effort: 2-3 days',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('#908 — Study Material Generation MCP Tools', level=3)
    items = [
        '7 MCP tools for study material CRUD and generation',
        'Generation tools are synchronous (wait for AI response)',
        'Duplicate detection (content_hash) prevents redundant AI calls',
        'Critical dates extraction still creates Task records',
        'Parent notifications fire when student generates via MCP',
        'Estimated effort: 3-4 days',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('8.3 Phase C: Advanced (#909, #910, #911)', level=2)
    doc.add_paragraph('Priority: P2 — Advanced features building on foundation.')

    doc.add_heading('#909 — AI Tutor Agent (Capstone Feature)', level=3)
    items = [
        'create_study_plan: Multi-day personalized study plans with auto-created Task records',
        'get_study_recommendations: Quick "what to study next" with rationale',
        'analyze_study_effectiveness: Trend-based analysis of study impact',
        'Multi-step AI prompt: (1) Analyze context → (2) Generate plan',
        'References existing study materials by ID when available',
        'Estimated effort: 5-7 days',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('#910 — Teacher Communication MCP Tools', level=3)
    items = [
        '5 MCP tools for messaging and teacher communication',
        'send_message requires user confirmation in MCP client',
        'Teacher communication summaries reuse existing AI summarization',
        'Estimated effort: 2-3 days',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('#911 — MCP Client (External Resources)', level=3)
    items = [
        'ClassBridge as MCP client consuming external servers',
        'Web search enrichment adds resource links to study guides',
        'Feature flag controlled (disabled by default)',
        'Graceful degradation when external servers unavailable',
        'No student PII sent to external servers',
        'External calls don\'t add >2s to generation latency (parallel + 5s timeout)',
        'Estimated effort: 3-4 days',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('8.4 Phase D: Quality (#912)', level=2)
    doc.add_paragraph('Priority: P1 — Required before shipping.')
    items = [
        '~40 unit tests across all MCP tools, resources, auth, and filtering',
        '3+ integration test flows (parent, student, teacher journeys)',
        'Performance benchmarks: <500ms reads, <30s AI generation, 10 concurrent clients',
        'Claude Desktop configuration guide (step-by-step with screenshots)',
        'Cursor IDE configuration guide',
        'Custom MCP client setup guide (Python SDK example)',
        'Estimated effort: 3-5 days',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ─── 9. COST ANALYSIS ───
    doc.add_heading('9. Cost Analysis', level=1)

    doc.add_heading('9.1 Additional Costs', level=2)
    add_styled_table(doc,
        ['Cost Item', 'Amount', 'Notes'],
        [
            ['Hosting (Cloud Run)', '~$0/month', 'MCP server is another HTTP endpoint on existing Cloud Run instance. Free tier covers 2M requests/month. At scale: ~$8/month at 1M requests.'],
            ['fastapi-mcp library', '$0', 'Open source (11.6K GitHub stars)'],
            ['MCP Python SDK', '$0', 'Official Anthropic SDK, open source'],
            ['External MCP servers', '$0-$5/month', 'Brave Search: 2,000 queries/month free. YouTube Data API: 10K queries/day free.'],
            ['Development time', '~4-6 weeks', 'Full Track F implementation (Phase A through D)'],
            ['API key storage', '~$0', 'Uses existing database (one new table)'],
        ],
        col_widths=[3.5, 3, 9.5]
    )

    doc.add_paragraph()
    doc.add_heading('9.2 Cost Savings', level=2)
    add_styled_table(doc,
        ['Saving', 'Estimated Impact', 'How'],
        [
            ['Avoid building native integrations', '$50K-$100K+ saved', 'A Claude Desktop plugin + Copilot extension + Cursor plugin would each take weeks of dev. MCP replaces all with one server.'],
            ['Reduce AI generation costs', '30-50% on enrichment', 'External MCP servers provide context that makes AI prompts more targeted → shorter, cheaper completions.'],
            ['Study Guide Repository leverage', '67% AI cost savings', 'MCP search tools let AI find existing materials before generating new ones (builds on #573).'],
            ['Support ticket reduction', '~20% fewer', 'Parents get instant answers from Claude Desktop instead of logging support requests.'],
            ['No custom chatbot needed', '$10K-$30K saved', 'No in-app chatbot/assistant needed — every MCP client IS the chatbot.'],
        ],
        col_widths=[3.5, 3, 9.5]
    )

    doc.add_paragraph()
    doc.add_heading('9.3 ROI Summary', level=2)
    add_styled_table(doc,
        ['Item', 'Year 1 Cost', 'Year 1 Value'],
        [
            ['Development', '~$15K-$25K (dev time)', '—'],
            ['Hosting', '~$0-$100', '—'],
            ['Avoided native integrations', '—', '$50K-$100K'],
            ['Avoided chatbot build', '—', '$10K-$30K'],
            ['AI cost savings (enrichment + dedup)', '—', '$1K-$5K'],
            ['Differentiation (no competitor has this)', '—', 'Hard to quantify — significant for fundraising/partnerships'],
            ['Net', '~$15K-$25K', '$61K-$135K+ saved'],
        ],
        col_widths=[5, 3.5, 5]
    )

    doc.add_page_break()

    # ─── 10. GITHUB ISSUES ───
    doc.add_heading('10. GitHub Issues Created', level=1)
    doc.add_paragraph(
        'All issues have been created in the ClassBridge GitHub repository (theepangnani/emai-dev-03) '
        'under the "mcp" label (purple #7057ff).'
    )
    add_styled_table(doc,
        ['Issue', 'Title', 'Priority', 'Labels'],
        [
            ['#903', 'EPIC: MCP Protocol Integration — AI-Powered Contextual Learning', '—', 'enhancement, mcp, ai, infrastructure'],
            ['#904', 'FastAPI-MCP server foundation with endpoint filtering', 'P0', 'enhancement, mcp, backend, infrastructure'],
            ['#905', 'MCP authentication and role-based tool authorization', 'P0', 'enhancement, mcp, backend, security'],
            ['#906', 'Student Academic Context MCP resources and tools', 'P1', 'enhancement, mcp, backend, ai'],
            ['#907', 'Google Classroom MCP tools', 'P1', 'enhancement, mcp, backend, integration'],
            ['#908', 'Study Material Generation MCP tools', 'P1', 'enhancement, mcp, backend, ai'],
            ['#909', 'AI Tutor Agent — contextual study plan generation via MCP', 'P2', 'enhancement, mcp, ai, backend'],
            ['#910', 'Teacher Communication MCP tools', 'P2', 'enhancement, mcp, backend'],
            ['#911', 'MCP Client — external educational resource discovery', 'P2', 'enhancement, mcp, backend, ai'],
            ['#912', 'MCP integration tests and Claude Desktop configuration', 'P1', 'enhancement, mcp, testing, documentation'],
        ],
        col_widths=[1.5, 7, 1.5, 6]
    )

    doc.add_paragraph()
    doc.add_heading('Implementation Roadmap', level=2)
    roadmap = """
Phase A (Foundation)       Phase B (Core Tools)        Phase C (Advanced)         Phase D (Quality)
┌───────────────────┐     ┌────────────────────┐      ┌────────────────────┐     ┌───────────────────┐
│ #904 MCP Server   │──→  │ #906 Student Ctx   │──→   │ #909 AI Tutor      │──→  │ #912 Tests +      │
│ #905 Auth/RBAC    │──→  │ #907 Classroom     │──→   │ #910 Messages      │     │      Docs         │
│                   │──→  │ #908 Study Tools   │      │ #911 Ext. Res.     │     │                   │
└───────────────────┘     └────────────────────┘      └────────────────────┘     └───────────────────┘
     ~4-5 days                  ~8-11 days                  ~10-14 days                ~3-5 days
"""
    p = doc.add_paragraph()
    run = p.add_run(roadmap.strip())
    run.font.name = 'Consolas'
    run.font.size = Pt(7.5)

    doc.add_page_break()

    # ─── 11. REQUIREMENTS FILES ───
    doc.add_heading('11. Requirements Files Updated', level=1)
    add_styled_table(doc,
        ['File', 'Change'],
        [
            ['requirements/features-part4.md', 'NEW — Full feature specification for §6.52 (MCP Protocol Integration), 9 sub-features, 26 MCP tools, architecture diagram'],
            ['requirements/roadmap.md', 'Added Phase 2+ Track F table with all 9 issues, priorities, dependencies, implementation phases'],
            ['REQUIREMENTS.md', 'Added features-part4.md to document structure table'],
            ['requirements/tracking.md', 'Added Phase 2+ Track F section with all 10 issues (#903-#912)'],
        ],
        col_widths=[5, 11]
    )

    doc.add_page_break()

    # ─── 12. RISKS & MITIGATIONS ───
    doc.add_heading('12. Risks & Mitigations', level=1)
    add_styled_table(doc,
        ['Risk', 'Impact', 'Likelihood', 'Mitigation'],
        [
            ['MCP protocol changes', 'Medium — could require code updates', 'Low — protocol is stabilizing', 'Use official MCP Python SDK which tracks spec changes; fastapi-mcp abstracts transport layer'],
            ['Low user adoption of MCP clients', 'Low — core product unaffected', 'Medium — MCP is still early', 'MCP server is additive, not a replacement. Web/mobile UIs remain primary. MCP is zero-cost to maintain if unused.'],
            ['Security: MCP exposes sensitive data', 'High — student data exposure', 'Low — controlled via RBAC', 'JWT/API key auth required; role-based tool filtering; excluded auth/admin/delete endpoints; audit logging'],
            ['AI cost increase from MCP usage', 'Medium — more AI API calls', 'Medium', 'Rate limiting matches web API limits (10/min); BYOK (#578) shifts costs to users; duplicate detection prevents redundant calls'],
            ['External MCP server unavailability', 'Low — enrichment is opt-in', 'Medium', 'Feature flag (disabled by default); graceful degradation; 5-second timeout; cached results (1-hour TTL)'],
            ['Privacy: student PII sent externally', 'High — PIPEDA/FERPA violation', 'Low — designed to prevent', 'Only topic name + grade level sent to external servers; never student PII; configurable per-server'],
        ],
        col_widths=[3, 2.5, 2, 8.5]
    )

    doc.add_page_break()

    # ─── 13. RECOMMENDATIONS ───
    doc.add_heading('13. Recommendations & Next Steps', level=1)

    doc.add_heading('Immediate (This Sprint)', level=2)
    items = [
        'Decision: Approve pulling #904 (FastAPI-MCP foundation) into Phase 1.5 as a 1-day quick win.',
        'No code changes needed for Phase 1 pilot — MCP work starts post-pilot.',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Phase 2 Sprint Planning', level=2)
    items = [
        'Sprint 1: Phase A — #904 + #905 (MCP server + auth). 4-5 days.',
        'Sprint 2: Phase B — #906 + #907 + #908 (student context, classroom, study tools). 8-11 days.',
        'Sprint 3: Phase C — #909 + #910 + #911 (AI tutor agent, messages, external resources). 10-14 days.',
        'Sprint 4: Phase D — #912 (tests, documentation, Claude Desktop setup guide). 3-5 days.',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Marketing & Positioning', level=2)
    items = [
        'Announce "ClassBridge + Claude Desktop" integration on launch of Phase A.',
        'Create demo video: parent asks Claude Desktop "what\'s my daughter\'s homework?" and gets real-time ClassBridge data.',
        'Add "Works with Claude Desktop, Copilot, Cursor, and ChatGPT" badge to landing page.',
        'Target tech-savvy early adopters who already use AI coding assistants.',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Long-Term (Phase 3+)', level=2)
    items = [
        'Ontario Curriculum MCP server (#571) — curriculum-aligned content enrichment.',
        'MCP-powered exam preparation engine (#576) — combines all data sources.',
        'Premium tier: advanced AI tutor agent features behind paid subscription.',
        'Open ClassBridge MCP server to third-party developers (API marketplace play).',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— End of Document —')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.italic = True

    return doc


if __name__ == '__main__':
    output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs')
    os.makedirs(output_dir, exist_ok=True)

    doc = build_document()
    output_path = os.path.join(output_dir, 'ClassBridge_MCP_Discovery_Strategy.docx')
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
