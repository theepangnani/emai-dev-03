"""Generate ClassBridge Project Status Report as Word document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('ClassBridge (EMAI)')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Comprehensive Project Status Report')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}')
run.font.size = Pt(12)

ver_p = doc.add_paragraph()
ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver_p.add_run('Prepared by: Claude Code Analysis')
run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Overall Project Health',
    '3. GitHub Issues Analysis',
    '4. Phase-by-Phase Status',
    '   4.1 Phase 1 (MVP)',
    '   4.2 Phase 1.5 (Calendar, Content, Mobile)',
    '   4.3 Phase 2 (Advanced Features)',
    '   4.4 Phase 3 (Course Planning)',
    '   4.5 Phase 4 (Tutor Marketplace)',
    '   4.6 Phase 5 (AI Email Agent)',
    '   4.7 Mobile App',
    '5. Feature Completion by Section',
    '6. All Incomplete Items',
    '7. Open Issues Breakdown',
    '8. Risk Assessment',
    '9. Recommended Priority Plan',
    '   9.1 Tier 1 \u2014 Fix Now (Before April 14)',
    '   9.2 Tier 2 \u2014 Build for Growth (Q2 2026)',
    '   9.3 Tier 3 \u2014 School Board Readiness',
    '   9.4 Tier 4 \u2014 Future Phases',
    '10. Conclusion',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'ClassBridge is an AI-powered education platform connecting parents, students, teachers, '
    'and administrators. The platform integrates with Google Classroom, provides AI study tools '
    '(study guides, quizzes, flashcards, mind maps), secure messaging, teacher email monitoring, '
    'and a comprehensive task management system.'
)
doc.add_paragraph(
    'The platform soft-launched on March 6, 2026 at classbridge.ca with a waitlist flow. '
    'The next planned launch milestone is April 14, 2026. The Phase 1 MVP is 99% complete '
    'with 945 GitHub issues closed out of 1,137 total (83% close rate). The overall requirements '
    'completion stands at 74% (491 of 661 checkboxes completed across all requirements documents).'
)

doc.add_heading('Key Highlights', level=2)
highlights = [
    'Phase 1 MVP: 99% complete \u2014 core platform fully functional',
    'Soft launch: Live at classbridge.ca since March 6, 2026',
    '945 issues closed, 192 open across all phases',
    'Mobile MVP built with Expo SDK 54 (React Native)',
    'AI study tools: Study guides, quizzes, flashcards, mind maps all operational',
    '4-role RBAC system: Parent, Student, Teacher, Admin',
    'Google Classroom integration: On-demand sync working',
]
for h in highlights:
    doc.add_paragraph(h, style='List Bullet')

doc.add_page_break()

# ============================================================
# 2. OVERALL PROJECT HEALTH
# ============================================================
doc.add_heading('2. Overall Project Health', level=1)

add_table(
    ['Metric', 'Value'],
    [
        ['GitHub Issues \u2014 Closed', '945'],
        ['GitHub Issues \u2014 Open', '192'],
        ['GitHub Issue Close Rate', '83%'],
        ['Requirements Checkboxes \u2014 Completed', '491'],
        ['Requirements Checkboxes \u2014 Remaining', '170'],
        ['Requirements Completion Rate', '74%'],
        ['Requirements Files', '7 (features-part1/2/3, roadmap, dashboards, mobile, technical)'],
        ['Soft Launch Date', 'March 6, 2026'],
        ['Next Launch Target', 'April 14, 2026'],
        ['Deployment Platform', 'GCP Cloud Run'],
        ['Production URL', 'https://www.classbridge.ca'],
    ],
    col_widths=[6, 10]
)

# ============================================================
# 3. GITHUB ISSUES ANALYSIS
# ============================================================
doc.add_heading('3. GitHub Issues Analysis', level=1)

doc.add_heading('Issue Statistics', level=2)
add_table(
    ['Category', 'Count', 'Percentage'],
    [
        ['Closed Issues', '945', '83%'],
        ['Open Issues', '192', '17%'],
        ['Total Issues', '1,137', '100%'],
    ]
)

doc.add_heading('Recent Closed Issues (Last 7 Days)', level=2)
add_table(
    ['#', 'Title', 'Date'],
    [
        ['1656', 'fix: add session idle timeout and reduce refresh token lifetime', '2026-03-13'],
        ['1654', 'fix: duplicate teacher entries when same email exists via class and manual link', '2026-03-13'],
        ['1648', 'fix: Help Chatbot items not clickable and chat state not persisted', '2026-03-12'],
        ['1646', 'fix: SpeedDial FAB overlaps bottom of content on course material pages', '2026-03-12'],
        ['1645', 'fix: Study guide content truncated for math-heavy assignments', '2026-03-12'],
        ['1640', 'fix: Help Chatbot not fullscreen on mobile devices', '2026-03-12'],
        ['1639', 'UX: Add View Source Files button next to Upload Document', '2026-03-12'],
        ['1634', 'fix: Edit Class modal missing Teacher and Students fields', '2026-03-12'],
        ['1633', 'Remove gradient style toggle', '2026-03-12'],
        ['1628', 'fix: broken images in document tab \u2014 relative markdown refs', '2026-03-12'],
        ['1626', 'MyKidsPage panel headers inconsistent with ParentDashboard', '2026-03-10'],
        ['1625', 'Upload Material modal doesn\'t show which child is selected', '2026-03-10'],
        ['1622', 'fix: MyKidsPage polish \u2014 school name, View navigation', '2026-03-10'],
        ['1621', 'fix: Document download 404 for pre-migration content', '2026-03-10'],
        ['1620', 'fix: Dashboard quick actions reorganization and additions', '2026-03-10'],
        ['1619', 'feat: Create Class wizard polish and enrollment follow-ups', '2026-03-10'],
        ['1618', 'fix: Filter parent Recent Activity to relevant items only', '2026-03-10'],
        ['1617', 'fix: Collapsible dashboard panels, simplified view', '2026-03-10'],
        ['1616', 'fix: Upload wizard loses class selection and resets', '2026-03-10'],
        ['1615', 'fix: Chat FAB icon and Study Guide UI polish iterations', '2026-03-10'],
    ]
)

doc.add_page_break()

# ============================================================
# 4. PHASE-BY-PHASE STATUS
# ============================================================
doc.add_heading('4. Phase-by-Phase Status', level=1)

doc.add_heading('Phase Overview', level=2)
add_table(
    ['Phase', 'Description', 'Status', 'Completion'],
    [
        ['Phase 1', 'MVP \u2014 Core Platform', 'Nearly Complete', '99%'],
        ['Phase 1.5', 'Calendar, Content, Mobile & School Integration', 'In Progress', '~90%'],
        ['Phase 2', 'Advanced Features, Compliance, Monetization', 'Partially Started', '~45%'],
        ['Phase 3', 'Course Planning & Guidance', 'Not Started', '0%'],
        ['Phase 4', 'Tutor Marketplace', 'Not Started', '0%'],
        ['Phase 5', 'AI Email Agent', 'Not Started', '0%'],
        ['Mobile', 'Expo SDK 54 React Native App', 'MVP Built', '77%'],
    ]
)

# Phase 1
doc.add_heading('4.1 Phase 1 \u2014 MVP (99% Complete)', level=2)
doc.add_paragraph(
    'Phase 1 is essentially complete. The core platform is fully functional with Google Classroom '
    'integration, AI study tools, RBAC, messaging, calendar, tasks, file upload, and comprehensive '
    'UI polish. Only 1 item remains:'
)
for item in [
    'Multi-Google account support for teachers (#945)',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Key Phase 1 achievements (257 items completed):')
for h in [
    'Google Classroom integration (on-demand sync)',
    'AI study tools: study guides, quizzes, flashcards, mind maps with difficulty levels',
    'Parent, Student, Teacher, Admin dashboards with role-based views',
    'Calendar with drag-and-drop, day/week/month views, child filtering',
    'Task system with CRUD, entity linking, reminders, archival',
    'Secure parent-teacher messaging with email notifications',
    'File upload with OCR, multi-file support, security hardening (20MB limit, magic bytes)',
    'Theme system (Light/Dark/Focus modes)',
    'Audit logging with admin UI',
    'JWT auth with refresh tokens, password reset, case-insensitive email login',
    'Loading skeletons, styled confirmation modals, lazy chunk retry',
    'Non-blocking AI generation with background processing',
    'Content moderation (Claude Haiku K-12 safety check)',
    'Print & PDF export for all study material types',
    'Waitlist system for soft launch',
]:
    doc.add_paragraph(h, style='List Bullet')

# Phase 1.5
doc.add_heading('4.2 Phase 1.5 \u2014 Calendar, Content, Mobile & School Integration (~90%)', level=2)
doc.add_paragraph('Most Phase 1.5 items are complete. 3 items remain:')
add_table(
    ['Item', 'Issue #', 'Status', 'Notes'],
    [
        ['Google Calendar push integration', '#943', 'Not Started', 'Sync tasks/reminders to Google Calendar'],
        ['School board email integration', '#942', 'Blocked', 'Waiting on DTAP approval'],
        ['Central document repository', '\u2014', 'Not Started', 'Shared document storage'],
    ]
)

# Phase 2
doc.add_heading('4.3 Phase 2 \u2014 Advanced Features (~45% Complete)', level=2)
doc.add_paragraph(
    'Phase 2 is the largest phase with the most remaining work. It covers monetization, compliance, '
    'advanced AI features, dashboard redesigns, and platform hardening. Many features are epics with '
    'multiple sub-tasks.'
)

add_table(
    ['Feature Area', 'Section', 'Status', '% Done', 'Open Issues'],
    [
        ['Upload Modal Redesign', '\u00a76.28', 'Near Complete', '91%', '#1266-#1273'],
        ['Mobile Responsiveness', '\u00a76.18', 'In Progress', '24%', '#1641'],
        ['Multi-Role Support', '\u00a76.24', 'Partial', '69%', '#255'],
        ['Show/Hide Password Toggle', '\u00a76.45', 'Not Started', '0%', '#420'],
        ['Lottie Animation Loader', '\u00a76.46', 'Not Started', '0%', '#424'],
        ['Admin Email Templates', '\u00a76.49', 'Not Started', '0%', '#513'],
        ['Broadcast History', '\u00a76.50', 'Not Started', '0%', '#514'],
        ['AI Usage Tracking', '\u00a76.54', 'Partial', '70%', '#1650, #1651'],
        ['Contextual Notes', '\u00a76.55', 'Partial', '80%', '#1090'],
        ['Help Chatbot RAG', '\u00a76.59', 'Not Started', '0%', '#1355-#1363'],
        ['Digital Wallet & Subscriptions', '\u00a76.60', 'Not Started', '0%', '#1384-#1392'],
        ['Smart Daily Briefing', '\u00a76.61', 'Not Started', '0%', '#1403'],
        ['Help My Kid', '\u00a76.62', 'Not Started', '0%', '#1407'],
        ['Dashboard Redesign v2', '\u00a76.65', 'Not Started', '0%', '#1418-#1419'],
        ['Learn Your Way', '\u00a76.69', 'Not Started', '0%', '#1436, #1441'],
        ['Smart Data Import', '\u00a76.67', 'Not Started', '0%', '#1431-#1433'],
        ['VASP/DTAP Compliance', '\u2014', 'Not Started', '0%', '#778-#806 (18 issues)'],
        ['MCP Protocol Integration', '\u2014', 'Not Started', '0%', '#903-#912 (7 issues)'],
        ['Teacher Grade Entry', '\u2014', 'Not Started', '0%', '#665'],
        ['Study Material Versioning', '\u2014', 'Not Started', '0%', '#884'],
        ['Premium Accounts', '\u2014', 'Not Started', '0%', '#1007'],
        ['BYOK (User AI API Key)', '\u2014', 'Not Started', '0%', '#578'],
        ['Student Progress Analysis', '\u2014', 'Not Started', '0%', '#575'],
    ]
)

# Phase 3
doc.add_heading('4.4 Phase 3 \u2014 Course Planning & Guidance (0% Complete)', level=2)
doc.add_paragraph('Phase 3 focuses on Ontario curriculum integration and academic planning. Not yet started.')
for item in [
    'Ontario Curriculum Management \u2014 download, parse, serve curriculum data (#571)',
    'School Board Integration \u2014 board-specific catalogs, student-board linking (#511, #500)',
    'Course Catalog Model \u2014 board-scoped high school course database (#500)',
    'Academic Plan Model \u2014 multi-year course plan Grade 9-12 (#501)',
    'Prerequisite & Graduation Requirements Engine \u2014 OSSD validation (#502)',
    'AI Course Recommendations \u2014 personalized guidance (#503)',
    'Semester Planner UI \u2014 per-semester course selection (#504)',
    'Multi-Year Planner UI \u2014 visual Grade 9-12 grid (#505)',
    'University Pathway Alignment \u2014 post-secondary program mapping (#506)',
    'Course Planning Dashboard Integration (#507)',
    'Course Planning Tests (#508)',
    'Multi-language support',
    'Advanced AI personalization',
    'Admin analytics',
]:
    doc.add_paragraph(item, style='List Bullet')

# Phase 4
doc.add_heading('4.5 Phase 4 \u2014 Tutor Marketplace (0% Complete)', level=2)
doc.add_paragraph('Phase 4 introduces a marketplace for private tutors. Not yet started.')
for item in ['Private tutor profiles (availability, rates, subjects)', 'Parent/student tutor search and discovery',
             'AI tutor matching', 'Booking workflow', 'Ratings & reviews', 'Payment integration']:
    doc.add_paragraph(item, style='List Bullet')

# Phase 5
doc.add_heading('4.6 Phase 5 \u2014 AI Email Agent (0% Complete)', level=2)
doc.add_paragraph('Phase 5 adds AI-powered email capabilities. Not yet started.')
for item in ['AI email sending', 'Reply ingestion', 'AI summaries', 'Searchable archive']:
    doc.add_paragraph(item, style='List Bullet')

# Mobile
doc.add_heading('4.7 Mobile App (77% Complete)', level=2)
doc.add_paragraph(
    'The mobile MVP is built with Expo SDK 54 / React Native. Core screens (login, parent dashboard, '
    'child overview, calendar, messages, notifications, profile) are implemented. Remaining work:'
)
for item in [
    'Device testing \u2014 physical iOS and Android via Expo Go (#375)',
    'Pilot launch checklist \u2014 verify production API connection (#376)',
    'Unit & component testing \u2014 8 test suites with Jest + RNTL (#490-#494)',
    'Pilot success criteria \u2014 7 acceptance tests not yet verified',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================
# 5. FEATURE COMPLETION BY SECTION
# ============================================================
doc.add_heading('5. Feature Completion by Section', level=1)

add_table(
    ['Section', 'Feature', 'Done', 'Remaining', '% Complete', 'Status'],
    [
        ['\u00a76.1', 'Google Classroom Integration', '7', '0', '100%', 'Complete'],
        ['\u00a76.2', 'AI Study Assistant', '7', '0', '100%', 'Complete'],
        ['\u00a76.3', 'Registration & Auth', 'All', '0', '100%', 'Complete'],
        ['\u00a76.4', 'Courses & Content', 'All', '0', '100%', 'Complete'],
        ['\u00a76.5', 'Analytics', 'All', '0', '100%', 'Complete'],
        ['\u00a76.6', 'Communication', 'All', '0', '100%', 'Complete'],
        ['\u00a76.7', 'Teacher Features', 'All', '0', '100%', 'Complete'],
        ['\u00a76.8-6.14', 'Tasks, Audit, etc.', 'All', '0', '100%', 'Complete'],
        ['\u00a76.15', 'Theme System', 'All', '0', '100%', 'Complete'],
        ['\u00a76.16', 'Layout Redesign', '3', '2', '60%', 'Partial'],
        ['\u00a76.18', 'Mobile Support', '4', '13', '24%', 'In Progress'],
        ['\u00a76.24', 'Multi-Role Support', '11', '5', '69%', 'Partial'],
        ['\u00a76.28', 'Upload Modal Redesign', '29', '3', '91%', 'Near Complete'],
        ['\u00a76.34', 'Course Enrollment', 'Most', '1', '99%', 'Near Complete'],
        ['\u00a76.35', 'Teacher Invites', 'Most', '4', '80%', 'Partial'],
        ['\u00a76.42', 'Admin Messaging', 'Most', '1', '95%', 'Near Complete'],
        ['\u00a76.45', 'Password Toggle', '0', '5', '0%', 'Not Started'],
        ['\u00a76.46', 'Lottie Loader', '0', '5', '0%', 'Not Started'],
        ['\u00a76.49', 'Email Templates', '0', '7', '0%', 'Not Started'],
        ['\u00a76.50', 'Broadcast History', '0', '7', '0%', 'Not Started'],
        ['\u00a76.53', 'Waitlist System', 'All', '0', '100%', 'Complete'],
        ['\u00a76.54', 'AI Usage Limits', 'Most', '3', '85%', 'Partial'],
        ['\u00a76.55', 'Contextual Notes', 'Most', '3', '80%', 'Partial'],
        ['\u00a76.56', 'Tutorial Pages', 'Most', '2', '90%', 'Near Complete'],
        ['\u00a76.57-6.58', 'Teacher Links, Images', 'All', '0', '100%', 'Complete'],
        ['\u00a76.59', 'Help Chatbot RAG', '0', '17', '0%', 'Not Started'],
        ['\u00a76.60', 'Digital Wallet', '0', '16', '0%', 'Not Started'],
        ['\u00a76.61', 'Daily Briefing', '0', '~8', '0%', 'Not Started'],
        ['\u00a76.62', 'Help My Kid', '0', '~6', '0%', 'Not Started'],
        ['\u00a76.65', 'Dashboard Redesign v2', '0', '~8', '0%', 'Not Started'],
        ['\u00a76.67', 'Smart Data Import', '0', '~6', '0%', 'Not Started'],
        ['\u00a76.69', 'Learn Your Way', '0', '~10', '0%', 'Not Started'],
        ['\u00a76.70-6.91', 'Recent Features', 'Most', '~5', '90%', 'Near Complete'],
    ]
)

doc.add_page_break()

# ============================================================
# 6. ALL INCOMPLETE ITEMS
# ============================================================
doc.add_heading('6. All Incomplete Items', level=1)

doc.add_heading('6.1 Bugs & Fixes (Immediate)', level=2)
add_table(['Issue', 'Description', 'Effort'], [
    ['#1653', 'Mind map branches overlap on desktop (radial layout)', 'Small'],
])

doc.add_heading('6.2 Phase 1 Remaining (1 item)', level=2)
add_table(['Item', 'Issue', 'Priority'], [
    ['Multi-Google account support for teachers', '#945', 'Medium'],
])

doc.add_heading('6.3 Phase 1.5 Remaining (3 items)', level=2)
add_table(['Item', 'Issue', 'Status'], [
    ['Google Calendar push integration', '#943', 'Not Started'],
    ['School board email integration', '#942', 'Blocked (DTAP)'],
    ['Central document repository', '\u2014', 'Not Started'],
])

doc.add_heading('6.4 Phase 2 \u2014 Core Features Not Yet Built', level=2)
add_table(['Feature', 'Section', 'Issue(s)', 'Notes'], [
    ['Mobile Responsiveness Audit', '\u00a76.18', '#1641', '55+ CSS files missing breakpoints'],
    ['Show/Hide Password Toggle', '\u00a76.45', '#420', '5 locations (Login, Register, Reset, Invite)'],
    ['Lottie Animation Loader', '\u00a76.46', '#424', 'Replace placeholder with branded animation'],
    ['Admin Email Template Management', '\u00a76.49', '#513', '7 sub-tasks: DB, CRUD, editor, preview'],
    ['Broadcast History Reuse/Resend', '\u00a76.50', '#514', '7 sub-tasks: detail, reuse, resend'],
    ['Help Chatbot RAG Full Build', '\u00a76.59', '#1355-#1363', '17 sub-tasks: YAML, embeddings, RAG, API, widget'],
    ['Help Chatbot Global Search', '\u00a76.59.9', '#1630', 'Unified SQL search across entities'],
    ['Digital Wallet & Subscriptions', '\u00a76.60', '#1384-#1392', '16 sub-tasks: Stripe, plans, wallet, invoicing'],
    ['Smart Daily Briefing', '\u00a76.61', '#1403', 'Proactive parent intelligence digest'],
    ['Help My Kid \u2014 One-Tap Actions', '\u00a76.62', '#1407', 'Quick study actions for parents'],
    ['Teacher Dashboard v2', '\u00a76.65.3', '#1418', 'Class overview, student alerts, quick actions'],
    ['Admin Dashboard v2', '\u00a76.65.4', '#1419', 'Platform health, user activity, quick actions'],
    ['Smart Data Import', '\u00a76.67', '#1431-#1433', 'Photo capture + email forwarding'],
    ['Learn Your Way', '\u00a76.69', '#1436, #1441', 'Interest-based personalized learning + paywall'],
    ['AI Token/Cost Tracking', '\u00a76.54', '#1650', 'Track token counts and cost per generation'],
    ['AI Regenerate Tracking', '\u00a76.54', '#1651', 'Track original vs regenerated content'],
    ['View All Activities Page', '\u2014', '#1547', 'Unified activity feed'],
    ['Hierarchical Study Guides', '\u2014', '#1594', 'Child guides from parent topics'],
    ['Admin Role Management UI', '\u00a76.24', '#255', 'Add/remove roles for any user'],
    ['Premium Accounts & Limits', '\u2014', '#1007', 'Subscription tiers with configurable limits'],
    ['BYOK \u2014 User AI API Key', '\u2014', '#578', 'Encrypted user-provided OpenAI key'],
    ['GCS File Storage Migration', '\u2014', '#1643, #572', 'Move source files from DB blob to cloud'],
    ['Teacher Grade & Feedback Entry', '\u2014', '#665', 'Per-student per-term grade input'],
    ['Student Progress Analysis', '\u2014', '#575', 'OCR score extraction, AI recommendations'],
    ['Data Privacy & User Rights', '\u2014', '#787', 'Account deletion, data export, consent'],
])

doc.add_heading('6.5 Phase 2 \u2014 VASP/DTAP Compliance (18 Issues)', level=2)
doc.add_paragraph(
    'These compliance items are mandatory before any Ontario school board partnership. '
    'They cover privacy law (MFIPPA/PIPEDA), security standards (SOC 2), and data residency requirements.'
)
add_table(['Issue', 'Priority', 'Description'], [
    ['#780', 'CRITICAL', 'OpenAI API data residency (US data transfer risk)'],
    ['#779', 'CRITICAL', 'Migrate infrastructure to GCP Canada region'],
    ['#781', 'CRITICAL', 'Update Privacy Policy & ToS for MFIPPA/PIPEDA'],
    ['#782', 'CRITICAL', 'Create Privacy Impact Assessment (PIA)'],
    ['#784', 'CRITICAL', 'Implement MFA/2FA support'],
    ['#785', 'HIGH', 'Implement SSO/SAML for school board integration'],
    ['#787', 'HIGH', 'Data export and right to erasure'],
    ['#789', 'HIGH', 'Dependency vulnerability scanning in CI/CD'],
    ['#790', 'HIGH', 'Annual penetration testing program'],
    ['#791', 'HIGH', 'SOC 2 Type II readiness assessment'],
    ['#792', 'HIGH', 'Obtain cyber liability insurance'],
    ['#793', 'HIGH', 'Create Data Processing Agreement (DPA) template'],
    ['#794', 'HIGH', 'Formal breach notification procedures'],
    ['#795', 'MEDIUM', 'Complete K-12CVAT vendor questionnaire'],
    ['#798', 'MEDIUM', 'Formal data retention and automated purging'],
    ['#799', 'MEDIUM', 'Designate privacy officer'],
    ['#800', 'MEDIUM', 'Enhance audit logging for SOC 2 / VASP'],
    ['#802', 'MEDIUM', 'Identify pilot school board partner'],
    ['#803', 'CRITICAL', 'EPIC: DTAP Compliance \u2014 Ontario School Board Approval'],
    ['#804', 'MEDIUM', 'WAF/DDoS protection (Cloud Armor)'],
    ['#806', 'MEDIUM', 'Formal data classification and inventory'],
])

doc.add_heading('6.6 Phase 2 \u2014 MCP Protocol Integration (7 Issues)', level=2)
add_table(['Issue', 'Description'], [
    ['#903', 'EPIC: MCP Protocol Integration \u2014 AI-Powered Contextual Learning'],
    ['#906', 'Student Academic Context MCP resources and tools'],
    ['#907', 'Google Classroom MCP tools'],
    ['#908', 'Study Material Generation MCP tools'],
    ['#909', 'AI Tutor Agent \u2014 contextual study plan generation'],
    ['#910', 'Teacher Communication MCP tools'],
    ['#911', 'MCP Client \u2014 external educational resource discovery'],
    ['#912', 'MCP integration tests and Claude Desktop configuration'],
])

doc.add_page_break()

# ============================================================
# 7. OPEN ISSUES BREAKDOWN
# ============================================================
doc.add_heading('7. Open Issues Breakdown (192 Total)', level=1)

add_table(
    ['Category', 'Count', 'Key Issues'],
    [
        ['Active Bugs', '~1', '#1653'],
        ['Phase 1 Remaining', '~1', '#945'],
        ['Phase 1.5', '~3', '#942, #943'],
        ['Phase 2 Features', '~60', 'Subscriptions, dashboards, AI features, data import'],
        ['Phase 2 Compliance (VASP/DTAP)', '~20', '#778-#806 (privacy, security, SOC 2)'],
        ['Phase 2 MCP Integration', '~7', '#903-#912'],
        ['Phase 2 Monetization', '~8', '#1384-#1392 (Stripe, wallet, invoicing)'],
        ['Phase 2 LMS Integrations', '~3', '#777, #778 (Brightspace/D2L)'],
        ['Phase 3 (Course Planning)', '~12', '#500-#508, #511, #571, #576'],
        ['Phase 4 (Marketplace)', '~5', '#968'],
        ['Phase 5 (AI Email)', '~4', '\u2014'],
        ['Mobile App', '~5', '#963, device testing'],
        ['Documentation/Strategy', '~8', '#761, #762, #1429, #1430, #1435'],
        ['Infrastructure', '~5', '#971, #1643'],
        ['Upload Wizard', '~7', '#1266-#1273'],
        ['Other Enhancements', '~45', 'Various phase-2 labeled items'],
    ]
)

doc.add_paragraph(
    'Note: Many open issues are Phase 2+ epics that were intentionally created as future roadmap items. '
    'The 192 open issues do not indicate 192 bugs or blockers \u2014 the majority are planned features for '
    'future phases.'
)

doc.add_page_break()

# ============================================================
# 8. RISK ASSESSMENT
# ============================================================
doc.add_heading('8. Risk Assessment', level=1)

add_table(
    ['Risk', 'Severity', 'Impact', 'Mitigation'],
    [
        ['Source files stored as DB blobs (#1643)',
         'HIGH', 'Database performance degradation as user base grows; backup size explosion',
         'Migrate to GCS with signed URLs (#572). Plan for Q2 2026.'],
        ['No monetization infrastructure',
         'HIGH', 'Cannot generate revenue; 16 sub-tasks for Stripe integration',
         'Prioritize \u00a76.60 Digital Wallet in Q2 2026.'],
        ['55+ CSS files missing mobile breakpoints (#1641)',
         'HIGH', 'Poor mobile UX for launch users; classbridge.ca on phones',
         'Systematic CSS audit before April 14 launch.'],
        ['VASP/DTAP compliance not started',
         'HIGH', 'Cannot partner with Ontario school boards without compliance',
         'Begin with data residency (#779, #780) and privacy docs (#781, #782).'],
        ['OpenAI API data residency (#780)',
         'CRITICAL', 'Student data processed in US; violates MFIPPA for Ontario boards',
         'Evaluate Azure OpenAI (Canada region) or data processing agreements.'],
        ['No MFA/2FA (#784)',
         'MEDIUM', 'Security gap for school board requirements',
         'Implement TOTP-based MFA in Phase 2.'],
        ['Session security (#1656)',
         'RESOLVED', 'Idle timeout and reduced refresh token lifetime implemented',
         'Fixed on 2026-03-13.'],
        ['Help Chatbot not fully built (#1355)',
         'LOW', 'Basic chatbot exists but RAG/knowledge base not implemented',
         'Phase 2 feature \u2014 current basic chatbot is functional.'],
    ]
)

doc.add_page_break()

# ============================================================
# 9. RECOMMENDED PRIORITY PLAN
# ============================================================
doc.add_heading('9. Recommended Priority Plan', level=1)

doc.add_heading('9.1 Tier 1 \u2014 Fix Now (Before April 14, 2026 Launch)', level=2)
doc.add_paragraph(
    'These items directly impact the user experience for the upcoming launch. '
    'They should be completed in the next 4 weeks.'
)
add_table(
    ['Priority', 'Task', 'Issue(s)', 'Effort', 'Rationale'],
    [
        ['1', 'Fix mind map overlap on desktop', '#1653', 'Small (1 day)', 'Visual bug, tagged as bug'],
        ['2', 'Mobile responsiveness audit', '#1641', 'Large (1-2 weeks)', '55+ CSS files need breakpoints'],
        ['3', 'AI usage token tracking', '#1650, #1651', 'Medium (3-4 days)', 'Cost control before scaling'],
        ['4', 'Show/hide password toggle', '#420', 'Small (1 day)', 'Basic UX, 5 form fields'],
        ['5', 'Replace tutorial SVG placeholders', '#1209', 'Small (1 day)', 'Polish \u2014 real screenshots'],
    ]
)

doc.add_heading('9.2 Tier 2 \u2014 Build for Growth (Q2 2026, Post-Launch)', level=2)
doc.add_paragraph(
    'These features drive user engagement and revenue. Target completion: Q2 2026.'
)
add_table(
    ['Priority', 'Task', 'Issue(s)', 'Effort', 'Rationale'],
    [
        ['8', 'Migrate files from DB to GCS', '#1643, #572', 'Large (1-2 weeks)', 'Scaling blocker'],
        ['9', 'Digital Wallet & Subscriptions', '#1384-#1392', 'Epic (4-6 weeks)', 'Monetization \u2014 Stripe'],
        ['10', 'Help Chatbot RAG (full)', '#1355-#1363', 'Epic (3-4 weeks)', 'Differentiated support'],
        ['11', 'Teacher & Admin Dashboard v2', '#1418, #1419', 'Large (2-3 weeks)', 'Teacher/admin UX basic'],
        ['12', 'Help My Kid \u2014 One-Tap Actions', '#1407', 'Medium (1-2 weeks)', 'Parent engagement'],
        ['13', 'Smart Daily Briefing', '#1403', 'Medium (1-2 weeks)', 'Proactive parent value'],
        ['14', 'Teacher grade & feedback entry', '#665', 'Medium (1-2 weeks)', 'Core teacher workflow'],
        ['15', 'Premium accounts & limits', '#1007', 'Medium (1 week)', 'Enable freemium model'],
    ]
)

doc.add_heading('9.3 Tier 3 \u2014 School Board Readiness (Before Board Partnerships)', level=2)
doc.add_paragraph(
    'Compliance is a prerequisite for any Ontario school board partnership. '
    'Estimate 2-3 months of dedicated work. These items are non-negotiable for board approval.'
)
add_table(
    ['Priority', 'Task', 'Issue(s)', 'Effort', 'Rationale'],
    [
        ['16', 'GCP Canada region migration', '#779', 'Large (1 week)', 'Data residency requirement'],
        ['17', 'OpenAI data residency solution', '#780', 'Large (2 weeks)', 'Student data in US is legal risk'],
        ['18', 'Privacy Policy + PIA', '#781, #782', 'Medium (1-2 weeks)', 'MFIPPA/PIPEDA compliance'],
        ['19', 'MFA/2FA support', '#784', 'Large (2 weeks)', 'Security requirement for boards'],
        ['20', 'SSO/SAML for school boards', '#785', 'Large (2-3 weeks)', 'Board IT integration'],
        ['21', 'Data export + right to erasure', '#787', 'Medium (1 week)', 'Privacy regulation'],
        ['22', 'Vulnerability scanning in CI/CD', '#789', 'Medium (3-5 days)', 'Security baseline'],
        ['23', 'Penetration testing program', '#790', 'Medium (ongoing)', 'Annual assessment'],
        ['24', 'SOC 2 readiness assessment', '#791', 'Large (2-3 months)', 'Enterprise trust certification'],
        ['25', 'DPA template + breach procedures', '#793, #794', 'Medium (1-2 weeks)', 'Legal documentation'],
    ]
)

doc.add_heading('9.4 Tier 4 \u2014 Future Phases (Q3+ 2026)', level=2)
doc.add_paragraph(
    'These are longer-term roadmap items that should be planned after Tiers 1-3 are substantially complete.'
)
for item in [
    'Phase 3: Ontario curriculum integration, academic planning, university pathway alignment (12+ items)',
    'Phase 4: Tutor marketplace \u2014 profiles, search, AI matching, booking, ratings (6 items)',
    'Phase 5: AI email agent \u2014 sending, ingestion, summaries, archive (4 items)',
    'LMS integrations: Brightspace/D2L, Canvas, Moodle (#967, #777, #778)',
    'MCP Protocol Integration: 7 sub-tasks for AI-powered contextual learning (#903-#912)',
    'Mobile app store distribution: iOS App Store + Google Play (#963)',
    'Multi-language support',
    'Advanced AI personalization',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================
# 10. CONCLUSION
# ============================================================
doc.add_heading('10. Conclusion', level=1)

doc.add_paragraph(
    'ClassBridge has made exceptional progress with a 74% requirements completion rate and 83% GitHub '
    'issue close rate. The Phase 1 MVP is effectively complete at 99%, delivering a fully functional '
    'AI-powered education platform with Google Classroom integration, comprehensive study tools, '
    'role-based dashboards, and a polished UI.'
)

doc.add_paragraph(
    'The platform is live at classbridge.ca since the March 6 soft launch with a waitlist system. '
    'The immediate priorities for the April 14 launch are: fixing 1 remaining bug (#1653 mind map overlap), '
    'completing the mobile responsiveness audit (55+ CSS files), and implementing AI usage tracking for cost control.'
)

doc.add_paragraph(
    'Looking ahead, the two most impactful workstreams are: (1) monetization infrastructure via the '
    'Digital Wallet & Subscription system (\u00a76.60), which is entirely unbuilt and required for revenue, '
    'and (2) VASP/DTAP compliance (18 issues), which is the gatekeeper for any Ontario school board '
    'partnership. Both should be active priorities in Q2 2026.'
)

doc.add_paragraph(
    'The project has strong momentum \u2014 945 issues closed, a working product in production, and a clear '
    'multi-phase roadmap. The foundation is solid for scaling from soft launch to full market entry.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\u2014 End of Report \u2014')
run.italic = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ============================================================
# SAVE
# ============================================================
output_path = r'c:\dev\emai\emai-dev-03\docs\ClassBridge_Project_Status_Report.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
