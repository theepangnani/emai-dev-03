"""
Export March 6 Phase 1 Pilot Launch Plan to Excel and Word.
Run: python scripts/export_launch_plan.py
Outputs: docs/pilot/March6_Launch_Plan.xlsx and docs/pilot/March6_Launch_Plan.docx
"""

import os
from datetime import datetime

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as WdAlignParagraph
from docx.oxml.ns import qn

# ──────────────────────────── DATA ────────────────────────────

TIMELINE = [
    ("Today (Feb 27)", "4 dev days remaining"),
    ("Mar 3 (Mon)", "Deploy freeze — no more merges to master"),
    ("Mar 4 (Tue)", "Dress rehearsal — full end-to-end walkthrough"),
    ("Mar 5 (Wed)", "Final smoke tests, pre-launch backup"),
    ("Mar 6 (Thu)", "LAUNCH DAY"),
]

BLOCKERS = [
    ("#589", "Submit Google OAuth app for production verification", "Process", "External — Google review takes days", ""),
    ("#587", "Record Google OAuth demo video for scope verification", "Process", "1-2 hrs", ""),
    ("#457", "Build and distribute ClassBridge mobile app via EAS Build", "Build", "2-4 hrs", ""),
    ("#460", "Link EAS project and update app.json placeholders", "Config", "30 min", ""),
    ("#461", "Register iOS pilot devices for ad hoc distribution", "Config", "1 hr", ""),
    ("#462", "Update pilot docs with real download links after EAS Build", "Docs", "30 min", ""),
    ("#375", "Device testing on physical iOS + Android devices", "Testing", "2-3 hrs", ""),
]

LAUNCH_PROCEDURES = {
    "Mar 3 — Deploy Freeze": [
        "Send deploy freeze announcement to all contributors",
        "Verify all feature branches stay unmerged",
    ],
    "Mar 4 — Dress Rehearsal": [
        "Web: Login as Parent, Student, Teacher, Admin — verify dashboards",
        "Web: Test messaging, notifications, calendar, AI study tools",
        "Mobile: Login as Parent on iOS and Android",
        "Mobile: Navigate all tabs (Dashboard, Child Overview, Calendar, Messages, Notifications, Profile)",
        "Cross-platform: Verify data syncs between web and mobile",
        "Verify all pilot accounts can log in",
        "Verify correct parent-child linkages",
        "Verify demo data is in place (assignments, messages, notifications)",
    ],
    "Mar 5 — Final Smoke Tests": [
        "Health check: curl -s https://www.classbridge.ca/health",
        "Check logs for errors",
        "Verify Cloud SQL is RUNNABLE",
        "All pilot accounts still log in",
        "Pre-launch backup: ./scripts/backup/manual-backup.sh pre-launch",
        "Alert policies active (5xx rate, latency, uptime)",
        "Welcome emails drafted and ready",
        "Mobile download links ready to share",
    ],
    "Mar 6 — Launch Morning": [
        "Verify services (health check, web loads, Cloud SQL, no errors)",
        "Verify all 4 roles on web + parent on mobile",
        "Send welcome emails (parent, teacher, student templates)",
        "Share mobile download links with pilot parents",
        "Open monitoring dashboards (Cloud Run + Cloud SQL metrics)",
        "Designate support contact available",
    ],
    "Mar 6 — Post-Launch Afternoon": [
        "Confirm all pilot participants received welcome emails",
        "Confirm at least one parent logged in via mobile",
        "Review Cloud Run metrics — no elevated error rates",
        "Review Cloud SQL metrics — no connection issues",
        "Check GitHub Issues for reported bugs",
        "Send 'launch successful' status update to stakeholders",
    ],
}

SHOULD_DO = [
    ("#509", "Welcome email on registration", "Polish — branded onboarding", "2-3 hrs", ""),
    ("#510", "Verification acknowledgement email", "Polish — post-verify marketing", "2-3 hrs", ""),
    ("#260", "Inspirational messages in emails", "Polish — adds personality", "3-4 hrs", ""),
    ("#833", "Teacher dashboard enhancement (SVG icons, dynamic counts)", "UX — teachers are pilot users", "4-6 hrs", ""),
    ("#834", "Student engagement (streaks, spaced repetition)", "UX — students are pilot users", "4-6 hrs", ""),
    ("#811", "Replace header back button with unified PageNav", "UX consistency", "2-3 hrs", ""),
    ("#142", "Input validation & field length limits", "Security hygiene", "3-4 hrs", ""),
    ("#73", "Missing database indexes", "Performance under load", "1-2 hrs", ""),
    ("#188", "Replace deprecated dependencies (python-jose, PyPDF2)", "Tech debt / security", "2-3 hrs", ""),
]

DEFERRED = [
    ("Architecture refactoring", "#127, #128, #129, #131, #132, #133, #134", "High risk, zero user impact"),
    ("Phase 1 New Workflows", "#546-#552", "Large features, destabilizing"),
    ("DTAP/VASP Compliance", "#779-#807 (26 issues)", "Phase 2, months of work"),
    ("Flat/non-gradient style", "#486, #487, #488, #489", "Cosmetic, low risk to defer"),
    ("Teacher multi-Google", "#41, #62", "Not needed for pilot"),
    ("Parent Dashboard Redesign", "#710-#719 (epic)", "Too large, current dashboard works"),
    ("Railway migration", "#769-#774", "Infrastructure change = high risk"),
    ("LMS abstraction/Brightspace", "#775-#778", "Phase 2 feature"),
    ("PostgreSQL test CI", "#156", "Doesn't affect users"),
    ("Alembic migrations", "#185", "Infrastructure, not user-facing"),
    ("HCD Tier 3", "#838, #839", "Grade integration + submissions = Phase 2"),
    ("Mobile unit tests", "#490-#494", "Nice-to-have, not blocking launch"),
]

ACTION_PLAN = {
    "Day 1 (Feb 27 — Today)": [
        "Start #587 (record OAuth demo video) -> submit #589 immediately (Google review needs lead time)",
        "Start mobile build chain: #460 -> #457",
        "Pick up #73 (DB indexes) — quick win, low risk",
    ],
    "Day 2 (Feb 28)": [
        "Complete #457 (EAS build) -> #461 (register devices) -> #462 (update docs)",
        "#375 — Device testing on physical iOS + Android",
        "Pick up #509 + #510 (welcome/verification emails)",
    ],
    "Day 3 (Mar 1)": [
        "#833 or #834 (teacher or student dashboard polish)",
        "#260 (inspirational messages in emails)",
        "#142 (input validation) or #188 (deprecated deps)",
    ],
    "Day 4 (Mar 2)": [
        "Buffer day — fix anything that broke",
        "Final PR reviews and merges",
        "Prep demo data for dress rehearsal",
    ],
}

SUMMARY = [
    ("Launch blockers (march-6-pilot)", "7 open", "Must close by Mar 2"),
    ("Launch procedures (checklist items)", "~25 items", "Mar 3-6"),
    ("High-value pre-freeze work", "~9 issues", "Best effort Feb 27-Mar 2"),
    ("Deferred post-launch", "~80+ issues", "After pilot"),
]

# ──────────────────────────── COLORS ────────────────────────────

BRAND_TEAL = "0D9488"
BRAND_DARK = "1E293B"
RED = "DC2626"
AMBER = "D97706"
GREEN = "16A34A"
LIGHT_GRAY = "F1F5F9"
WHITE = "FFFFFF"

# ──────────────────────────── EXCEL ────────────────────────────


def create_excel(output_path: str):
    wb = Workbook()

    header_font = Font(name="Calibri", bold=True, color=WHITE, size=11)
    header_fill = PatternFill(start_color=BRAND_TEAL, end_color=BRAND_TEAL, fill_type="solid")
    dark_fill = PatternFill(start_color=BRAND_DARK, end_color=BRAND_DARK, fill_type="solid")
    stripe_fill = PatternFill(start_color=LIGHT_GRAY, end_color=LIGHT_GRAY, fill_type="solid")
    title_font = Font(name="Calibri", bold=True, size=14, color=BRAND_DARK)
    subtitle_font = Font(name="Calibri", bold=True, size=11, color=BRAND_TEAL)
    bold_font = Font(name="Calibri", bold=True, size=11)
    normal_font = Font(name="Calibri", size=11)
    thin_border = Border(
        left=Side(style="thin", color="D1D5DB"),
        right=Side(style="thin", color="D1D5DB"),
        top=Side(style="thin", color="D1D5DB"),
        bottom=Side(style="thin", color="D1D5DB"),
    )
    wrap = Alignment(wrap_text=True, vertical="top")

    def style_header(ws, row, cols, fill=None):
        for c in range(1, cols + 1):
            cell = ws.cell(row=row, column=c)
            cell.font = header_font
            cell.fill = fill or header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = thin_border

    def style_row(ws, row, cols, striped=False):
        for c in range(1, cols + 1):
            cell = ws.cell(row=row, column=c)
            cell.font = normal_font
            cell.alignment = wrap
            cell.border = thin_border
            if striped:
                cell.fill = stripe_fill

    # ── Sheet 1: Overview ──
    ws = wb.active
    ws.title = "Overview"
    ws.sheet_properties.tabColor = BRAND_TEAL

    ws.merge_cells("A1:E1")
    ws["A1"].value = "ClassBridge - March 6 Phase 1 Pilot Launch Plan"
    ws["A1"].font = Font(name="Calibri", bold=True, size=18, color=BRAND_DARK)
    ws["A1"].alignment = Alignment(horizontal="center")

    ws.merge_cells("A2:E2")
    ws["A2"].value = f"Generated {datetime.now().strftime('%B %d, %Y')}"
    ws["A2"].font = Font(name="Calibri", size=10, color="64748B")
    ws["A2"].alignment = Alignment(horizontal="center")

    # Timeline
    r = 4
    ws.cell(row=r, column=1, value="TIMELINE").font = subtitle_font
    r += 1
    for h in ["Date", "Milestone"]:
        ws.cell(row=r, column=["Date", "Milestone"].index(h) + 1, value=h)
    style_header(ws, r, 2, dark_fill)
    r += 1
    for date, milestone in TIMELINE:
        ws.cell(row=r, column=1, value=date).font = bold_font
        ws.cell(row=r, column=2, value=milestone).font = normal_font
        style_row(ws, r, 2, (r - 6) % 2 == 0)
        if "LAUNCH" in milestone:
            ws.cell(row=r, column=2).font = Font(name="Calibri", bold=True, size=11, color=RED)
        r += 1

    # Summary
    r += 1
    ws.cell(row=r, column=1, value="SUMMARY").font = subtitle_font
    r += 1
    for h_idx, h in enumerate(["Category", "Count", "Deadline"]):
        ws.cell(row=r, column=h_idx + 1, value=h)
    style_header(ws, r, 3)
    r += 1
    for cat, cnt, deadline in SUMMARY:
        ws.cell(row=r, column=1, value=cat)
        ws.cell(row=r, column=2, value=cnt)
        ws.cell(row=r, column=3, value=deadline)
        style_row(ws, r, 3, (r % 2 == 0))
        r += 1

    # Risk callout
    r += 1
    ws.merge_cells(f"A{r}:E{r}")
    cell = ws.cell(row=r, column=1)
    cell.value = "WARNING: #1 RISK — Google OAuth verification (#589) requires Google review (3-7+ business days). Submit TODAY."
    cell.font = Font(name="Calibri", bold=True, size=12, color=RED)
    cell.fill = PatternFill(start_color="FEF2F2", end_color="FEF2F2", fill_type="solid")
    cell.alignment = Alignment(horizontal="center")

    ws.column_dimensions["A"].width = 35
    ws.column_dimensions["B"].width = 55
    ws.column_dimensions["C"].width = 20
    ws.column_dimensions["D"].width = 15
    ws.column_dimensions["E"].width = 15

    # ── Sheet 2: Launch Blockers ──
    ws2 = wb.create_sheet("1. Blockers")
    ws2.sheet_properties.tabColor = RED

    ws2.merge_cells("A1:F1")
    ws2["A1"].value = "LAUNCH BLOCKERS - Must close before March 3 freeze"
    ws2["A1"].font = Font(name="Calibri", bold=True, size=14, color=RED)

    r = 3
    headers = ["Done?", "Issue", "Task", "Type", "Estimate", "Notes"]
    for i, h in enumerate(headers):
        ws2.cell(row=r, column=i + 1, value=h)
    style_header(ws2, r, len(headers), PatternFill(start_color=RED, end_color=RED, fill_type="solid"))
    r += 1

    for issue, task, typ, est, notes in BLOCKERS:
        ws2.cell(row=r, column=1, value="").alignment = Alignment(horizontal="center", vertical="top")
        ws2.cell(row=r, column=2, value=issue)
        ws2.cell(row=r, column=3, value=task)
        ws2.cell(row=r, column=4, value=typ)
        ws2.cell(row=r, column=5, value=est)
        ws2.cell(row=r, column=6, value=notes)
        style_row(ws2, r, len(headers), (r % 2 == 0))
        r += 1

    r += 1
    ws2.cell(row=r, column=1, value="Critical Path:").font = bold_font
    ws2.cell(row=r, column=2, value="#460 -> #457 -> #461 -> #462 -> #375 (mobile build chain is sequential)").font = normal_font
    ws2.merge_cells(f"B{r}:F{r}")
    r += 1
    ws2.cell(row=r, column=2, value="#587 -> #589 (OAuth video then submit — Google review has external lead time)").font = normal_font
    ws2.merge_cells(f"B{r}:F{r}")

    for col, w in [(1, 8), (2, 10), (3, 55), (4, 12), (5, 15), (6, 30)]:
        ws2.column_dimensions[get_column_letter(col)].width = w

    # ── Sheet 3: Launch Procedures ──
    ws3 = wb.create_sheet("2. Launch Procedures")
    ws3.sheet_properties.tabColor = BRAND_TEAL

    ws3.merge_cells("A1:C1")
    ws3["A1"].value = "LAUNCH WEEK PROCEDURES (Mar 3-6)"
    ws3["A1"].font = Font(name="Calibri", bold=True, size=14, color=BRAND_TEAL)

    r = 3
    for phase, items in LAUNCH_PROCEDURES.items():
        ws3.merge_cells(f"A{r}:C{r}")
        ws3.cell(row=r, column=1, value=phase).font = Font(name="Calibri", bold=True, size=12, color=WHITE)
        for c in range(1, 4):
            ws3.cell(row=r, column=c).fill = dark_fill
            ws3.cell(row=r, column=c).border = thin_border
        r += 1

        headers = ["Done?", "#", "Task"]
        for i, h in enumerate(headers):
            ws3.cell(row=r, column=i + 1, value=h)
        style_header(ws3, r, 3)
        r += 1

        for idx, item in enumerate(items, 1):
            ws3.cell(row=r, column=1, value="").alignment = Alignment(horizontal="center", vertical="top")
            ws3.cell(row=r, column=2, value=idx)
            ws3.cell(row=r, column=3, value=item)
            style_row(ws3, r, 3, idx % 2 == 0)
            r += 1
        r += 1

    ws3.column_dimensions["A"].width = 8
    ws3.column_dimensions["B"].width = 5
    ws3.column_dimensions["C"].width = 75

    # ── Sheet 4: Should-Do ──
    ws4 = wb.create_sheet("3. Should-Do Pre-Freeze")
    ws4.sheet_properties.tabColor = AMBER

    ws4.merge_cells("A1:F1")
    ws4["A1"].value = "HIGH-VALUE WORK - Best effort before Mar 3 freeze"
    ws4["A1"].font = Font(name="Calibri", bold=True, size=14, color=AMBER)

    r = 3
    headers = ["Done?", "Issue", "Task", "Impact", "Estimate", "Notes"]
    for i, h in enumerate(headers):
        ws4.cell(row=r, column=i + 1, value=h)
    style_header(ws4, r, len(headers), PatternFill(start_color=AMBER, end_color=AMBER, fill_type="solid"))
    r += 1

    for issue, task, impact, est, notes in SHOULD_DO:
        ws4.cell(row=r, column=1, value="").alignment = Alignment(horizontal="center", vertical="top")
        ws4.cell(row=r, column=2, value=issue)
        ws4.cell(row=r, column=3, value=task)
        ws4.cell(row=r, column=4, value=impact)
        ws4.cell(row=r, column=5, value=est)
        ws4.cell(row=r, column=6, value=notes)
        style_row(ws4, r, len(headers), (r % 2 == 0))
        r += 1

    for col, w in [(1, 8), (2, 10), (3, 55), (4, 30), (5, 12), (6, 30)]:
        ws4.column_dimensions[get_column_letter(col)].width = w

    # ── Sheet 5: Action Plan ──
    ws5 = wb.create_sheet("4. Action Plan")
    ws5.sheet_properties.tabColor = GREEN

    ws5.merge_cells("A1:C1")
    ws5["A1"].value = "RECOMMENDED ACTION PLAN (Feb 27 - Mar 2)"
    ws5["A1"].font = Font(name="Calibri", bold=True, size=14, color=GREEN)

    r = 3
    for day, tasks in ACTION_PLAN.items():
        ws5.merge_cells(f"A{r}:C{r}")
        ws5.cell(row=r, column=1, value=day).font = Font(name="Calibri", bold=True, size=12, color=WHITE)
        for c in range(1, 4):
            ws5.cell(row=r, column=c).fill = PatternFill(start_color=GREEN, end_color=GREEN, fill_type="solid")
            ws5.cell(row=r, column=c).border = thin_border
        r += 1

        for idx, task in enumerate(tasks, 1):
            ws5.cell(row=r, column=1, value="").alignment = Alignment(horizontal="center", vertical="top")
            ws5.cell(row=r, column=2, value=idx)
            ws5.cell(row=r, column=3, value=task)
            style_row(ws5, r, 3, idx % 2 == 0)
            r += 1
        r += 1

    ws5.column_dimensions["A"].width = 8
    ws5.column_dimensions["B"].width = 5
    ws5.column_dimensions["C"].width = 80

    # ── Sheet 6: Deferred ──
    ws6 = wb.create_sheet("5. Deferred Post-Launch")
    ws6.sheet_properties.tabColor = "94A3B8"

    ws6.merge_cells("A1:C1")
    ws6["A1"].value = "DEFERRED - Do NOT touch before March 6"
    ws6["A1"].font = Font(name="Calibri", bold=True, size=14, color="64748B")

    r = 3
    headers = ["Category", "Issues", "Reason to Defer"]
    for i, h in enumerate(headers):
        ws6.cell(row=r, column=i + 1, value=h)
    style_header(ws6, r, 3, PatternFill(start_color="94A3B8", end_color="94A3B8", fill_type="solid"))
    r += 1

    for cat, issues, reason in DEFERRED:
        ws6.cell(row=r, column=1, value=cat)
        ws6.cell(row=r, column=2, value=issues)
        ws6.cell(row=r, column=3, value=reason)
        style_row(ws6, r, 3, (r % 2 == 0))
        r += 1

    ws6.column_dimensions["A"].width = 30
    ws6.column_dimensions["B"].width = 35
    ws6.column_dimensions["C"].width = 45

    for ws_item in wb.worksheets:
        ws_item.sheet_view.showGridLines = False

    wb.save(output_path)
    print(f"  Excel saved: {output_path}")


# ──────────────────────────── WORD ────────────────────────────


def _shade_cell(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color,
    })
    shading.append(shading_elem)


def add_table(doc, headers, rows, color=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WdAlignParagraph.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
        _shade_cell(cell, color or BRAND_TEAL)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                _shade_cell(table.rows[r_idx + 1].cells[c_idx], LIGHT_GRAY)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(6.5 / len(headers))

    return table


def create_word(output_path: str):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading("ClassBridge - March 6 Phase 1 Pilot Launch Plan", level=0)
    title.runs[0].font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    p = doc.add_paragraph(f"Generated {datetime.now().strftime('%B %d, %Y')}  |  classbridge.ca")
    p.alignment = WdAlignParagraph.CENTER
    p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    p.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # Risk callout
    risk_p = doc.add_paragraph()
    risk_run = risk_p.add_run(
        "WARNING  #1 RISK: Google OAuth verification (#589) — requires Google review (3-7+ business days). Submit TODAY."
    )
    risk_run.bold = True
    risk_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    risk_run.font.size = Pt(12)

    doc.add_paragraph()

    # ── Section 1: Timeline ──
    doc.add_heading("Timeline", level=1)
    add_table(doc, ["Date", "Milestone"], TIMELINE, BRAND_DARK)
    doc.add_paragraph()

    # ── Section 2: Summary ──
    doc.add_heading("Summary", level=1)
    add_table(doc, ["Category", "Count", "Deadline"], SUMMARY)
    doc.add_paragraph()

    # ── Section 3: Launch Blockers ──
    h = doc.add_heading("1. Launch Blockers", level=1)
    h.runs[0].font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    p = doc.add_paragraph("Must close before March 3 deploy freeze. 7 open issues tagged ")
    p.add_run("march-6-pilot").bold = True
    p.add_run(".")

    blocker_rows = [(b[0], b[1], b[2], b[3]) for b in BLOCKERS]
    add_table(doc, ["Issue", "Task", "Type", "Estimate"], blocker_rows, RED)

    doc.add_paragraph()
    cp = doc.add_paragraph()
    cp.add_run("Critical path: ").bold = True
    cp.add_run("#460 -> #457 -> #461 -> #462 -> #375 (mobile build chain)")
    cp2 = doc.add_paragraph()
    cp2.add_run("OAuth path: ").bold = True
    cp2.add_run("#587 -> #589 (demo video then submit for Google review)")
    doc.add_paragraph()

    # ── Section 4: Launch Procedures ──
    h = doc.add_heading("2. Launch Week Procedures", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x0D, 0x94, 0x88)

    for phase, items in LAUNCH_PROCEDURES.items():
        doc.add_heading(phase, level=2)
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("[ ]  ")
            p.add_run(item)

    doc.add_paragraph()

    # ── Section 5: Should-Do ──
    h = doc.add_heading("3. High-Value Pre-Freeze Work", level=1)
    h.runs[0].font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
    doc.add_paragraph("Best effort before March 3. Not launch blockers, but improves pilot experience.")

    should_rows = [(s[0], s[1], s[2], s[3]) for s in SHOULD_DO]
    add_table(doc, ["Issue", "Task", "Impact", "Estimate"], should_rows, AMBER)
    doc.add_paragraph()

    # ── Section 6: Action Plan ──
    h = doc.add_heading("4. Recommended Action Plan", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)

    for day, tasks in ACTION_PLAN.items():
        doc.add_heading(day, level=2)
        for task in tasks:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("[ ]  ")
            p.add_run(task)

    doc.add_paragraph()

    # ── Section 7: Deferred ──
    doc.add_heading("5. Deferred Post-Launch", level=1)
    doc.add_paragraph("Do NOT touch before March 6. ~80+ issues deferred to post-pilot.")

    add_table(doc, ["Category", "Issues", "Reason to Defer"], DEFERRED, "94A3B8")
    doc.add_paragraph()

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph("ClassBridge 2026  |  classbridge.ca  |  Confidential")
    footer.alignment = WdAlignParagraph.CENTER
    footer.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    footer.runs[0].font.size = Pt(9)

    doc.save(output_path)
    print(f"  Word saved: {output_path}")


# ──────────────────────────── MAIN ────────────────────────────

if __name__ == "__main__":
    out_dir = os.path.join(os.path.dirname(__file__), "..", "docs", "pilot")
    os.makedirs(out_dir, exist_ok=True)

    xlsx_path = os.path.join(out_dir, "March6_Launch_Plan.xlsx")
    docx_path = os.path.join(out_dir, "March6_Launch_Plan.docx")

    print("Exporting March 6 Launch Plan...")
    create_excel(xlsx_path)
    create_word(docx_path)
    print("\nDone! Files saved to docs/pilot/")
