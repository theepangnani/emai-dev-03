"""Generate the March 10 requirements/design/plan Word document."""
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def bold_para(bold_text, normal_text=''):
    p = doc.add_paragraph()
    r = p.add_run(bold_text)
    r.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p


def make_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = str(text)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table


def bullets(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


# ── Title Page ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ClassBridge')
r.font.size = Pt(36)
r.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Requirements, Design & Development Plan')
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x4A, 0x4A, 0x8A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Updated: March 10, 2026')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Author: Theepan Gnanasabapathy')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('classbridge.ca')
r.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

doc.add_page_break()

# ── TOC ──
heading('Table of Contents')
for item in [
    '1. Executive Summary',
    '2. Project Status Overview',
    '3. Launch Timeline',
    '4. Phase 1 (MVP) - Completed Features',
    '5. Phase 1.5 (Polish & UX) - Completed Features',
    '6. WOW Features - April 14 Target',
    '7. Phase 2 - Planned Features',
    '8. GitHub Issue Tracking',
    '9. Technical Architecture',
    '10. April 14 Launch Plan',
    '11. Risk Register',
]:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# ── 1. Executive Summary ──
heading('1. Executive Summary')
doc.add_paragraph(
    'ClassBridge is a unified, AI-powered education platform that connects parents, students, '
    'teachers, and administrators. It integrates with school systems (Google Classroom), provides '
    'AI-driven study tools (study guides, quizzes, flashcards, mind maps), simplifies communication, '
    "and enables parents to actively support their children's education."
)
doc.add_paragraph(
    'The platform soft-launched on March 6, 2026 with a "Join Waitlist" flow at classbridge.ca. '
    'The next major launch is planned for April 14, 2026, targeting early adopters from the waitlist '
    'with WOW features that differentiate ClassBridge from existing tools.'
)

# ── 2. Project Status ──
heading('2. Project Status Overview')
make_table(
    ['Metric', 'Count', 'Percentage', 'Notes'],
    [
        ['Roadmap Items Completed', '257', '81%', 'Out of 317 total items'],
        ['Roadmap Items Remaining', '60', '19%', 'Mostly Phase 2-5'],
        ['GitHub Issues (Total)', '1,627+', '-', 'Highest issue #1627'],
        ['GitHub Issues (Closed)', '1,440+', '88%+', 'Resolved/implemented'],
        ['GitHub Issues (Open)', '186', '12%', 'Phase 2+ features'],
        ['Requirements Sections', '90', '-', 'Sections 6.1 through 6.90'],
    ],
)
doc.add_paragraph()
bold_para('Phase Completion:')
bullets([
    'Phase 1 (MVP): ~95% - Core platform, AI tools, dashboards, messaging, calendar',
    'Phase 1.5 (Polish): ~90% - UI audit fixes, themes, mobile responsive, skeletons',
    'Phase 2 (Monetization): ~15% - Digital wallet structure created, implementation pending',
    'Phase 3 (Marketplace): 0% - Private tutor marketplace (future)',
    'Phase 4 (Scale): 0% - School board partnerships (future)',
])
doc.add_page_break()

# ── 3. Launch Timeline ──
heading('3. Launch Timeline')
make_table(
    ['Date', 'Milestone', 'Details'],
    [
        ['March 6, 2026', 'Soft Launch', 'Join Waitlist flow live at classbridge.ca'],
        ['March 10, 2026', 'Requirements Reassessment', 'Aligned docs + GitHub issues with implementation'],
        ['March 10 - April 11', 'WOW Feature Sprint', 'Build Smart Briefing, Help My Kid, Dashboard v2'],
        ['April 11, 2026', 'Code Freeze', 'Feature complete, testing only'],
        ['April 14, 2026', 'Launch #2', 'Open to waitlist users with WOW features'],
    ],
)
doc.add_page_break()

# ── 4. Phase 1 Completed ──
heading('4. Phase 1 (MVP) - Completed Features')
for name, desc in [
    ('Google Classroom Integration', 'On-demand sync of courses, assignments, teachers'),
    ('Role-Based Dashboards', 'Parent, Student, Teacher, Admin with tailored views'),
    ('AI Study Tools', 'Study guides, quizzes, flashcards with GPT-4o-mini'),
    ('Parent-Teacher Messaging', 'Secure in-app messaging with email notifications'),
    ('Calendar System', 'Day/3-Day/Week/Month views, drag-and-drop, sticky note tasks'),
    ('Task Management', 'CRUD, archival, entity linking, detail page, calendar integration'),
    ('Course Materials', 'Upload, OCR extraction, tabbed detail view (Original/Guide/Quiz/Flash)'),
    ('Parent-Child Linking', 'Many-to-many, email invites, Google Classroom discovery'),
    ('My Kids Page', 'Child cards, teacher linking, colored avatars, progress bars'),
    ('Multi-Role Support', 'Users can hold multiple roles, role switcher in UI'),
    ('Theme System', 'Light/Dark/Focus modes, 50+ CSS variables, OS preference detection'),
    ('Audit Logging', 'Comprehensive action logging with admin UI and retention policies'),
    ('Security', 'JWT + refresh tokens, rate limiting, security headers, RBAC'),
    ('Notification System', 'In-app + email, task reminders, message notifications'),
    ('Non-Blocking AI Generation', 'Background generation with pulsing placeholders'),
    ('Styled Confirmation Modals', 'Custom ConfirmModal replacing all window.confirm()'),
    ('Loading Skeletons', 'Animated skeleton screens for all major pages'),
    ('Mobile Responsive CSS', 'Breakpoints for 5+ pages'),
    ('Backend Tests', '288+ route tests'),
]:
    bold_para(name + ' - ', desc)
doc.add_page_break()

# ── 5. Phase 1.5 Completed ──
heading('5. Phase 1.5 (Polish & UX) - Completed Features')
for name, desc in [
    ('Waitlist System (6.53)', 'Public registration, admin approval, welcome emails, capacity limits'),
    ('AI Usage Limits (6.54)', 'Per-user daily caps, admin config, usage tracking, Request More flow'),
    ('AI Help Chatbot (6.59)', 'RAG-based help bot, persistent FAB, video embeds, role-aware'),
    ('Contextual Notes (6.55)', 'Inline notes on courses, materials, tasks; per-user private notes'),
    ('Upload Wizard Redesign (6.28)', 'Multi-step upload modal with drag-drop, progress, course assignment'),
    ('Flat Style Default (6.15.2)', 'Replaced 30+ gradients with solid accent colors'),
    ('Welcome/Verification Emails (6.48)', 'SendGrid templates, invite flow, password reset emails'),
    ('Mind Map Generation (6.74)', 'AI-generated mind maps from study materials'),
    ('Notes Revision History (6.75)', 'Version tracking for contextual notes'),
    ('Course Material Grouping (6.76)', 'Organize materials by topic/unit'),
    ('Daily Morning Email Digest (6.77)', 'Automated daily summary email for parents'),
    ('Tutorial Completion Tracking (6.79)', 'Track onboarding tutorial progress per user'),
    ('Command Palette (6.80)', 'Keyboard shortcut search across the app'),
    ('Recent Activity Panel (6.81)', 'Activity feed on dashboard'),
    ('LaTeX Math Rendering (6.82)', 'KaTeX rendering in study guides and notes'),
    ('Chat FAB & Study Guide UI (6.84)', 'Floating action button for help chatbot'),
    ('Collapsible Panels (6.86)', 'Expandable/collapsible dashboard sections'),
    ('Activity Feed Filter (6.87)', 'Filter activity by type, date, child'),
    ('Create Class Wizard Polish (6.88)', 'Improved class creation flow'),
    ('Quick Actions Reorg (6.89)', 'Reorganized quick action buttons'),
    ('MyKidsPage Polish (6.90)', 'SectionPanel component, school name display'),
]:
    bold_para(name + ' - ', desc)
doc.add_page_break()

# ── 6. WOW Features ──
heading('6. WOW Features - April 14 Target')
doc.add_paragraph(
    'These features transform ClassBridge from a passive information viewer into an active parenting '
    'tool. They address the pilot feedback: "I don\'t see a WOW factor." The WOW comes when '
    "ClassBridge does things parents can't do anywhere else."
)

heading('Tier 1 - Must Ship (Weeks 1-3)', 2)

heading('#1403 - Smart Daily Briefing (6.61)', 3)
doc.add_paragraph(
    "Proactive daily summary card on the parent dashboard. Shows what's due today, upcoming deadlines, "
    'recent grades, and suggested actions. AI-generated, personalized per child. The #1 engagement '
    'driver - gives parents a reason to open the app every morning.'
)
bullets([
    'Backend: aggregate child data, generate briefing via AI',
    'Frontend: briefing card on parent dashboard (#1405)',
    'Morning email digest integration',
])

heading('#1407 - Help My Kid (6.62)', 3)
doc.add_paragraph(
    'One-tap study actions from the parent dashboard. Parent sees "Your child has a math test '
    'tomorrow" and taps "Create Study Guide" - ClassBridge generates a targeted guide instantly. '
    'Bridges the gap between knowing and doing.'
)

heading('#1547 - View All Activities Page', 3)
doc.add_paragraph(
    'Dedicated page showing all platform activity (assignments, tasks, study guides, messages) in '
    'a filterable, chronological feed. Completes the activity system.'
)

heading('Tier 2 - High Impact (Weeks 3-4)', 2)

heading('#1418 - Teacher Dashboard v2 (6.65.3)', 3)
doc.add_paragraph(
    'Class overview, student alerts, quick actions. Makes the teacher experience compelling enough '
    'that teachers actively want to use ClassBridge.'
)

heading('#1419 - Admin Dashboard v2 (6.65.4)', 3)
doc.add_paragraph(
    'Platform health metrics, user activity tracking, quick admin actions. Essential for managing '
    'the platform as real users onboard.'
)

heading('#1594 - Hierarchical Study Guides', 3)
doc.add_paragraph(
    'Generate child study guides from parent topics. E.g., "Grade 8 Math" spawns guides for each '
    'unit. Deepens the core AI study tool.'
)

heading('Tier 3 - Defer Post-April', 2)
bullets([
    'Digital Wallet & Subscriptions (6.60) - monetization after user base',
    'Smart Data Import (6.67) - photo capture + email parsing, complex',
    'Learn Your Way (6.69) - interest-based personalization, ambitious',
    'Weekly Progress Pulse (6.63) - nice but not critical for launch',
    'Parent-Child Study Link (6.64) - collaborative features, later phase',
])
doc.add_page_break()

# ── 7. Phase 2 Planned ──
heading('7. Phase 2 - Planned Features')
for name, desc in [
    ('Digital Wallet & Subscriptions (6.60)', 'Stripe integration, subscription tiers, AI credit purchases, invoicing'),
    ('Smart Daily Briefing (6.61)', 'AI-generated daily parent summary - Tier 1 for April 14'),
    ('Help My Kid (6.62)', 'One-tap study actions - Tier 1 for April 14'),
    ('Weekly Progress Pulse (6.63)', 'Automated weekly reports for parents'),
    ('Parent-Child Study Link (6.64)', 'Collaborative study sessions between parent and child'),
    ('Dashboard Redesign v2 (6.65)', 'Teacher and Admin dashboard overhauls'),
    ('Responsible AI Parent Tools (6.66)', 'AI transparency, usage reports, content controls'),
    ('Smart Data Import (6.67)', 'Photo capture -> OCR, email forwarding -> auto-parse'),
    ('Learn Your Way (6.69)', 'Interest-based personalized learning paths'),
    ('Per-Category Notifications (6.70)', 'Granular notification preferences'),
    ('Premium Storage Limits (6.71)', 'Tiered storage by subscription level'),
]:
    bold_para(name + ' - ', desc)
doc.add_page_break()

# ── 8. GitHub Tracking ──
heading('8. GitHub Issue Tracking')
make_table(
    ['Category', 'Count', 'Status'],
    [
        ['Total Issues Created', '1,627+', 'Comprehensive tracking'],
        ['Issues Closed', '1,440+', '88%+ resolved'],
        ['Issues Open', '186', 'Phase 2+ features'],
        ['Phase 2 / Monetization', '~20', 'Digital wallet, subscriptions'],
        ['WOW Features', '~15', 'Smart Briefing, Help My Kid, etc.'],
        ['Documentation', '~10', 'Strategy docs, design decisions'],
    ],
)
doc.add_paragraph()
heading('Key Open Issues for April 14', 2)
bullets([
    '#1403 - Smart Daily Briefing (6.61) - Epic',
    '#1405 - Daily briefing frontend card',
    '#1407 - Help My Kid one-tap actions (6.62)',
    '#1418 - Teacher Dashboard v2 (6.65.3)',
    '#1419 - Admin Dashboard v2 (6.65.4)',
    '#1547 - View All Activities page',
    '#1594 - Hierarchical study guides',
])
doc.add_page_break()

# ── 9. Technical Architecture ──
heading('9. Technical Architecture')
heading('Stack', 2)
for name, desc in [
    ('Backend', 'FastAPI, Python 3.13+, SQLAlchemy 2.0, Pydantic 2.x'),
    ('Frontend', 'React 19, TypeScript, Vite, React Router 7, TanStack Query'),
    ('Database', 'SQLite (dev), PostgreSQL (prod)'),
    ('AI', 'OpenAI GPT-4o-mini (study tools, chatbot, briefings)'),
    ('Auth', 'JWT + OAuth2 (Google), refresh tokens'),
    ('Email', 'SendGrid transactional emails'),
    ('Deploy', 'GCP Cloud Run, auto-deploy on merge to master'),
    ('Mobile', 'Expo SDK 54, React Native (MVP complete)'),
    ('Domain', 'classbridge.ca'),
]:
    bold_para(name + ': ', desc)
doc.add_page_break()

# ── 10. April 14 Launch Plan ──
heading('10. April 14 Launch Plan')

heading('Week 1 (March 10-14)', 2)
bullets([
    'Begin Smart Daily Briefing backend (#1403)',
    'Begin Help My Kid backend (#1407)',
    'View All Activities page (#1547)',
])

heading('Week 2 (March 17-21)', 2)
bullets([
    'Smart Daily Briefing frontend card (#1405)',
    'Help My Kid frontend integration',
    'Begin Teacher Dashboard v2 (#1418)',
])

heading('Week 3 (March 24-28)', 2)
bullets([
    'Complete Teacher Dashboard v2',
    'Admin Dashboard v2 (#1419)',
    'Hierarchical Study Guides (#1594)',
])

heading('Week 4 (March 31 - April 4)', 2)
bullets([
    'Integration testing across all new features',
    'Bug fixes from testing',
    'Performance optimization',
])

heading('Week 5 (April 7-11)', 2)
bullets([
    'Final QA and regression testing',
    'Waitlist user communication prep',
    'April 11: Code freeze',
])

heading('April 14 - Launch Day', 2)
bullets([
    'Open platform to waitlist users',
    'Send launch emails with login instructions',
    'Monitor dashboards for errors',
    'Support channel active',
])
doc.add_page_break()

# ── 11. Risk Register ──
heading('11. Risk Register')
make_table(
    ['Risk', 'Impact', 'Likelihood', 'Mitigation'],
    [
        ['WOW features not ready by April 14', 'High', 'Medium', 'Tier system - launch with Tier 1 minimum'],
        ['Google OAuth verification delays', 'High', 'Low', 'Already submitted; fallback to test mode'],
        ['AI costs spike with real users', 'Medium', 'Medium', 'Usage limits already in place (6.54)'],
        ['Low waitlist conversion', 'Medium', 'Medium', 'WOW features + onboarding email sequence'],
        ['Performance under load', 'Medium', 'Low', 'Cloud Run auto-scaling; load test before launch'],
    ],
)

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ClassBridge 2026 | classbridge.ca | Confidential')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

output = r'C:\Users\tgnan\OneDrive\Theepan\Business\EMAI\Notes\Latest.req.update.analyss.March.10.2026.docx'
doc.save(output)
print(f'Saved: {output}')
