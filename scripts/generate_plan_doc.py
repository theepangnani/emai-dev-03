"""Generate Word document for Image Retention Plan."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

# -- Title Page --
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Image Retention in Study Guides')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Implementation Plan & Cost Analysis')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4B, 0x5B, 0x6B)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in ['Project: EMAI (ClassBridge)', 'Date: March 7, 2026', 'GitHub Issues: #1308 - #1313']:
    meta.add_run(line + '\n').font.size = Pt(12)

doc.add_page_break()

# -- 1. Problem Statement --
doc.add_heading('1. Problem Statement', level=1)
doc.add_paragraph(
    'When users upload documents (PDF, DOCX, PPTX) containing images \u2014 diagrams, charts, '
    'formulas, screenshots \u2014 the study guide generation pipeline extracts text via OCR but '
    'discards the original image binaries. The resulting study guides are text-only, losing '
    'valuable visual context that is critical for learning.'
)
doc.add_heading('Impact', level=2)
impacts = [
    'Students lose diagrams, charts, and visual aids essential for science, math, and geography',
    "Teachers' carefully crafted visual materials are reduced to text descriptions",
    'Study guide quality is significantly lower than the source material',
]
for imp in impacts:
    doc.add_paragraph(imp, style='List Bullet')

# -- 2. Current Architecture --
doc.add_heading('2. Current Architecture', level=1)
doc.add_paragraph(
    'The current pipeline follows three stages: Document Upload (file_processor.py extracts text '
    'and OCRs images but discards binaries), Study Guide Generation (ai_service.py sends text-only '
    'to Claude AI), and Frontend Rendering (MarkdownBody renders text-only markdown).'
)

# -- 3. Proposed Solution --
doc.add_heading('3. Proposed Solution: Image Extraction + Reference Embedding', level=1)
doc.add_paragraph(
    'The solution works in three layers: (1) Extract and store images as ContentImage records during '
    'upload, (2) Include image metadata in AI prompts so the AI places {{IMG-N}} markers, and '
    '(3) Resolve markers to real image URLs in the frontend.'
)

phases = [
    ('Phase 1 \u2014 Extract & Store', 'During document upload, extract embedded images from PDF/DOCX/PPTX. '
     'Capture surrounding text context. Compress to max 800px width. Store as ContentImage records. '
     'Reuse existing Vision OCR descriptions (no new AI cost).'),
    ('Phase 2 \u2014 AI-Aware Placement', 'Include image metadata in the AI prompt: '
     '"[IMG-1] Photosynthesis diagram (near: Light reactions...)". '
     'AI returns markdown with ![description]({{IMG-1}}) markers at appropriate locations.'),
    ('Phase 3 \u2014 Frontend Rendering', 'MarkdownBody component parses {{IMG-N}} markers and replaces '
     'them with <img> tags pointing to the image serving API endpoint. '
     'PDF export embeds images as base64.'),
]
for phase_title, phase_desc in phases:
    doc.add_heading(phase_title, level=2)
    doc.add_paragraph(phase_desc)

# -- 4. Implementation Tasks --
doc.add_heading('4. Implementation Tasks', level=1)

doc.add_heading('Batch 1 \u2014 Foundation', level=2)
table = doc.add_table(rows=2, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Task', 'Issue', 'Description', 'Key Files']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
row = table.rows[1]
row.cells[0].text = 'ContentImage model'
row.cells[1].text = '#1308'
row.cells[2].text = 'New SQLAlchemy model + DB migration'
row.cells[3].text = 'content_image.py, main.py'

doc.add_heading('Batch 2 \u2014 Core Features (parallel)', level=2)
table2 = doc.add_table(rows=4, cols=4)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(headers):
    table2.rows[0].cells[i].text = h
    table2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
batch2 = [
    ('Image extraction', '#1309', 'Extract images from PDF/DOCX/PPTX', 'file_processor.py'),
    ('AI prompt integration', '#1310', 'Image metadata in AI prompts', 'ai_service.py, study.py'),
    ('Image serving endpoint', '#1311', 'API to serve stored images', 'course_contents.py'),
]
for i, (task, issue, desc, files) in enumerate(batch2):
    row = table2.rows[i + 1]
    row.cells[0].text = task
    row.cells[1].text = issue
    row.cells[2].text = desc
    row.cells[3].text = files

doc.add_heading('Batch 3 \u2014 Frontend & Polish (parallel)', level=2)
table3 = doc.add_table(rows=3, cols=4)
table3.style = 'Light Grid Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(headers):
    table3.rows[0].cells[i].text = h
    table3.rows[0].cells[i].paragraphs[0].runs[0].bold = True
batch3 = [
    ('Frontend rendering', '#1312', 'Render images inline in study guides', 'MarkdownBody.tsx'),
    ('Fallback section', '#1313', 'Append unplaced images at bottom', 'study.py'),
]
for i, (task, issue, desc, files) in enumerate(batch3):
    row = table3.rows[i + 1]
    row.cells[0].text = task
    row.cells[1].text = issue
    row.cells[2].text = desc
    row.cells[3].text = files

# -- 5. Cost Analysis --
doc.add_heading('5. Cost Analysis', level=1)

doc.add_heading('Current Costs Per Study Guide Generation', level=2)
table4 = doc.add_table(rows=5, cols=4)
table4.style = 'Light Grid Accent 1'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
cost_headers = ['Step', 'Model', 'Tokens (in/out)', 'Cost']
for i, h in enumerate(cost_headers):
    table4.rows[0].cells[i].text = h
    table4.rows[0].cells[i].paragraphs[0].runs[0].bold = True
cost_rows = [
    ('Content safety check', 'Haiku 4.5', '150 / 20', '~$0.0002'),
    ('Vision OCR (images)', 'Haiku 4.5', '10,000 / 4,096', '~$0.024'),
    ('Study guide generation', 'Configured', '2,000-4,000 / 2,000', '~$0.01-0.03'),
    ('CURRENT TOTAL', '', '', '$0.03-0.05'),
]
for i, (step, model, tokens, cost) in enumerate(cost_rows):
    row = table4.rows[i + 1]
    row.cells[0].text = step
    row.cells[1].text = model
    row.cells[2].text = tokens
    row.cells[3].text = cost
# Bold the total row
for cell in table4.rows[4].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

doc.add_heading('Added Costs With Image Retention', level=2)
table5 = doc.add_table(rows=4, cols=3)
table5.style = 'Light Grid Accent 1'
table5.alignment = WD_TABLE_ALIGNMENT.CENTER
add_headers = ['New Step', 'What Changes', 'Added Cost']
for i, h in enumerate(add_headers):
    table5.rows[0].cells[i].text = h
    table5.rows[0].cells[i].paragraphs[0].runs[0].bold = True
added = [
    ('Image description storage', 'Reuse existing OCR (currently discarded)', '$0.00'),
    ('Extra prompt tokens', '+500-1,000 tokens for image metadata', '~$0.002-0.005'),
    ('Image binary storage', '50-200KB/image compressed', 'Storage only'),
]
for i, (step, change, cost) in enumerate(added):
    row = table5.rows[i + 1]
    row.cells[0].text = step
    row.cells[1].text = change
    row.cells[2].text = cost

doc.add_heading('Monthly Cost Comparison', level=2)
table6 = doc.add_table(rows=4, cols=4)
table6.style = 'Light Grid Accent 1'
table6.alignment = WD_TABLE_ALIGNMENT.CENTER
comp_headers = ['Metric', 'Current', 'With Feature', 'Change']
for i, h in enumerate(comp_headers):
    table6.rows[0].cells[i].text = h
    table6.rows[0].cells[i].paragraphs[0].runs[0].bold = True
comp = [
    ('Per generation', '$0.03-0.05', '$0.035-0.055', '+5-10%'),
    ('Monthly (500 gens)', '$15-25', '$17.50-27.50', '+$2.50'),
    ('Storage (100 docs/mo)', '\u2014', '50-300MB', '~$0.05-0.09/mo'),
]
for i, (metric, curr, new, change) in enumerate(comp):
    row = table6.rows[i + 1]
    row.cells[0].text = metric
    row.cells[1].text = curr
    row.cells[2].text = new
    row.cells[3].text = change

# -- 6. Risks --
doc.add_heading('6. Risks & Mitigations', level=1)
table7 = doc.add_table(rows=6, cols=3)
table7.style = 'Light Grid Accent 1'
table7.alignment = WD_TABLE_ALIGNMENT.CENTER
risk_headers = ['Risk', 'Impact', 'Mitigation']
for i, h in enumerate(risk_headers):
    table7.rows[0].cells[i].text = h
    table7.rows[0].cells[i].paragraphs[0].runs[0].bold = True
risks = [
    ('AI ignores image markers', 'Low', 'Fallback "Additional Figures" section'),
    ('Large docs (50+ images)', 'Medium', 'Cap at 20 images, skip tiny/decorative'),
    ('Image extraction fails', 'Low', 'Graceful degradation to text-only'),
    ('DB bloat', 'Medium', 'Compress/resize, cleanup orphans'),
    ('Page load performance', 'Medium', 'Cache headers, lazy loading'),
]
for i, (risk, impact, mitigation) in enumerate(risks):
    row = table7.rows[i + 1]
    row.cells[0].text = risk
    row.cells[1].text = impact
    row.cells[2].text = mitigation

# -- 7. Rollout --
doc.add_heading('7. Rollout Plan', level=1)
rollout = [
    'Batch 1: Merge & deploy ContentImage model (no user-facing changes)',
    'Batch 2: Merge & deploy extraction + AI + endpoint (images stored and referenced)',
    'Batch 3: Merge & deploy frontend rendering (images visible to users)',
]
for i, step in enumerate(rollout, 1):
    doc.add_paragraph(step, style='List Number')
doc.add_paragraph(
    'Each batch is independently deployable with no breaking changes. '
    'The feature is backward-compatible \u2014 existing study guides without images continue to work normally.'
)

# Save
out_dir = os.path.join(os.path.dirname(__file__), '..', 'docs')
docx_path = os.path.join(out_dir, 'Image-Retention-Plan.docx')
doc.save(docx_path)
print(f"Word document saved: {os.path.abspath(docx_path)}")
