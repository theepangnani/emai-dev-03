"""Generate ClassBridge CEO Platform Review as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Styles --
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Calibri"
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles["Normal"]
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.style = doc.styles["Normal"]
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()


def add_code_block(text):
    """Add a monospace code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading("ClassBridge", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading("CEO Platform Review", level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

meta_items = [
    "Mode: SCOPE EXPANSION -- Pure 10x Vision",
    "Date: 2026-03-22",
    "Reviewer: Claude Opus 4.6 (1M context) -- Mega Plan Review",
    "Branch: master",
    "Project: ClassBridge (EMAI) -- AI-powered education platform",
]
for item in meta_items:
    p = doc.add_paragraph(item)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
doc.add_heading("1. Executive Summary", level=1)

doc.add_paragraph(
    "ClassBridge is a substantial Phase 1 education platform -- soft-launched March 6, 2026 "
    "at classbridge.ca with an April 14, 2026 full launch planned. The platform serves parents, "
    "students, teachers, and administrators with Google Classroom integration, AI-powered study "
    "tools (streaming generation), gamification (XP/streaks/badges), messaging, a RAG-powered "
    "chatbot, and a waitlist/survey system."
)

doc.add_paragraph(
    "This review was conducted in SCOPE EXPANSION mode -- focused on identifying the architectural "
    "moves that transform ClassBridge from a feature-rich MVP into a category-defining "
    "'Learning Operating System.'"
)

doc.add_heading("Key Findings", level=2)
add_table(
    ["Area", "Score", "Critical Items"],
    [
        ["Architecture", "C- (current), B+ target", "No event bus, no domain boundaries, business logic in routes"],
        ["Error Handling", "4 CRITICAL GAPS", "Content safety fail-open, wallet non-atomic, silent email, Google credential silent"],
        ["Security", "7/10", "Prompt injection partial, rate limiting ineffective, content moderation gaps"],
        ["Testing", "Good baseline, gaps", "195 backend + 258 frontend tests, but zero E2E, SQLite-only"],
        ["Performance", "2 critical N+1s", "Student search O(N), data export N+1"],
        ["Observability", "Minimal", "Text logs only, no metrics, no tracing, no alerting"],
        ["Deployment", "Functional but risky", "No staging, no canary, no migration rollback"],
    ],
)

doc.add_heading("Key Decisions Made", level=2)
add_table(
    ["#", "Decision", "Choice"],
    [
        ["1", "Review mode", "SCOPE EXPANSION (Pure 10x vision)"],
        ["2", "Event bus timing", "Start now -- foundation for Sept bundle"],
        ["3", "Content safety", "Fail-closed with bypass flag"],
        ["4", "Prompt injection defense", "Full input scanning on all AI inputs"],
        ["5", "Empty content handling", "Block + explain when < 50 chars extracted"],
        ["6", "E2E test introduction", "Incremental with each batch"],
        ["7", "Stash cleanup", "Audit and clean now"],
    ],
)

doc.add_page_break()

# ============================================================
# 2. SYSTEM AUDIT
# ============================================================
doc.add_heading("2. System Audit", level=1)

doc.add_heading("Current System State", level=2)
items = [
    "Backend: 53 API route files, 47 models, 50+ services",
    "Frontend: 100+ pages across 4 role dashboards (Parent, Student, Teacher, Admin)",
    "Tech Stack: FastAPI + React 19 + TypeScript + Vite + TanStack Query",
    "Database: SQLite (dev), PostgreSQL (prod)",
    "AI: Anthropic Claude (streaming, safety) + OpenAI GPT-4o-mini (quizzes, flashcards, chatbot)",
    "Deploy: GCP Cloud Run, auto-deploys on merge to master",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("What's In Flight", level=2)
items = [
    "49 git stashes -- significant accumulated WIP, some over a year old",
    "Recent commits: streaming study guide generation, bug report enhancements, race condition fixes",
    "5 untracked docs in docs/ directory",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Taste Calibration -- Well-Designed (Emulate)", level=2)
add_table(
    ["File", "Quality", "Why"],
    [
        ["xp_service.py (665 LOC)", "A+", "Fail-safe (never raises), excellent logging, anti-gaming checks, clean constants"],
        ["deps.py (104 LOC)", "A", "Token blacklist cache with TTL + LRU. require_role() factory. Security-first."],
        ["useStudyGuideStream.ts (740 LOC)", "A", "SSE streaming with buffered flushing, abort controller, mounted-state tracking"],
        ["api/client.ts", "A", "Exponential backoff, queue-based token refresh, deploy readiness gates"],
    ],
)

doc.add_heading("Taste Calibration -- Anti-Patterns (Avoid)", level=2)
add_table(
    ["File", "Quality", "Issue"],
    [
        ["main.py (~150 lines migrations)", "C-", "Raw ALTER TABLE in startup. Exception swallowing. No rollback."],
        ["study.py (routes)", "C-", "45+ imports. Business logic in route handler. Untestable."],
        ["ParentDashboard.tsx (893 LOC)", "C", "15+ useState hooks, localStorage scattered, God component"],
        ["StudentDashboard.tsx (925 LOC)", "C", "Same problems. Duplicates utilities from ParentDashboard."],
    ],
)

doc.add_page_break()

# ============================================================
# 3. STEP 0: SCOPE CHALLENGE
# ============================================================
doc.add_heading("3. Step 0: Scope Challenge & Mode Selection", level=1)

doc.add_heading("Premise Challenge", level=2)
doc.add_paragraph(
    "Is this the right problem to solve? ClassBridge has validated the problem space. "
    "The real question: is the foundation ready to scale from soft launch to a platform "
    "school boards adopt?"
)
doc.add_paragraph(
    "What if we did nothing? The feature set is rich. But 49 stashes, SQLite-only tests, "
    "C- architecture grade suggest the codebase is accumulating velocity debt."
)

doc.add_heading("10x Check", level=2)
doc.add_paragraph(
    "The 10x version transforms ClassBridge from a tool parents use into a Learning Operating "
    "System that families and schools can't live without:"
)
items = [
    "Real-time learning graph: Every interaction feeds a per-student knowledge graph.",
    "Predictive interventions: Predict grade outcomes and trigger interventions before tests.",
    "Teacher-as-first-class-citizen: Teachers upload curriculum, ClassBridge generates differentiated materials.",
    "School board dashboard: Board-level analytics, cohort tracking, compliance reporting.",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Platonic Ideal", level=2)
doc.add_paragraph(
    '"I open ClassBridge and it already knows what my kid should be working on today. '
    "It shows me a 30-second summary of yesterday's progress, flags that a math test is "
    "in 3 days and my kid hasn't reviewed Chapter 7, and gives me a single button: "
    "'Help them prepare.' I tap it, and 2 minutes later my kid has a personalized study guide, "
    'practice quiz, and flashcard set for exactly the topics they\'re weak on."'
)

doc.add_heading("Delight Opportunities", level=2)
add_table(
    ["#", "Delight", "Effort", "Impact"],
    [
        ["1", "Study streak shared with parent -- milestone notifications", "30 min", "High emotional connection"],
        ["2", "Quiz of the Day -- auto-generated daily challenge", "2 hours", "Daily habit loop"],
        ["3", "Teacher gratitude -- one-tap thank you with counter", "1 hour", "Unique differentiator"],
        ["4", "Smart study time suggestions -- data-driven insights", "1 hour", "Parents love data"],
        ["5", "Weekly family report card email -- shareable to grandparents", "2 hours", "Viral growth mechanic"],
    ],
)

doc.add_page_break()

# ============================================================
# 4. ARCHITECTURE REVIEW
# ============================================================
doc.add_heading("4. Section 1: Architecture Review", level=1)

doc.add_paragraph(
    "The current architecture is a monolithic CRUD app with AI sprinkled on top. "
    "To become a learning operating system, ClassBridge needs to evolve into an "
    "event-driven platform with a knowledge core."
)

doc.add_heading("Key Architectural Moves", level=2)

doc.add_heading("Move 1: Event Bus (the spine)", level=3)
doc.add_paragraph(
    "Every user action becomes a domain event. Currently xp_service.award_xp() is called "
    "directly from 6 endpoints. In the 10x version, those endpoints emit events (StudentStudied, "
    "QuizCompleted, etc.) and gamification, knowledge graph, notifications, and analytics "
    "all subscribe independently. This is the single most important architectural change."
)

doc.add_heading("Move 2: Knowledge Graph Domain (the brain)", level=3)
doc.add_paragraph(
    "Doesn't exist yet. Every MaterialUploaded feeds topics into the graph. Every QuizCompleted "
    "updates mastery scores. Every GradeReceived calibrates predictions. Becomes the source of "
    "truth for 'what does this student know?'"
)

doc.add_heading("Move 3: Intelligence Domain (the nervous system)", level=3)
doc.add_paragraph(
    "AI evolves from a utility to a decision-maker: 'This student should study Chapter 7 "
    "because their mastery is 40% and the test is in 3 days.'"
)

doc.add_heading("Move 4: Plugin Architecture for Integrations", level=3)
doc.add_paragraph(
    "Google Classroom is hard-wired. Adding TeachAssist or Brightspace should be a new plugin "
    "with a standard interface (sync, import, export), not a codebase change."
)

doc.add_heading("Scaling Characteristics", level=2)
add_table(
    ["At 10x (500 users)", "At 100x (5,000 users)"],
    [
        ["AI costs ~$500/mo", "AI costs ~$5,000/mo without BYOK/caching"],
        ["SQLite/PG divergence causes bugs", "Need read replicas"],
        ["In-memory rate limit resets on deploy", "Need Redis-backed rate limiting"],
        ["Single Cloud Run instance bottleneck", "Need horizontal scaling + WebSocket"],
    ],
)

doc.add_heading("Single Points of Failure", level=2)
items = [
    "OpenAI API -- all study generation, quizzes, flashcards, content safety. No fallback.",
    "main.py migrations -- startup failure = broken schema, no rollback.",
    "SendGrid -- all email. No fallback.",
    "GCS -- all file storage. If down, no uploads.",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_page_break()

# ============================================================
# 5. ERROR & RESCUE MAP
# ============================================================
doc.add_heading("5. Section 2: Error & Rescue Map", level=1)

doc.add_heading("Error & Rescue Registry", level=2)
add_table(
    ["Codepath", "Exception", "Rescued?", "User Sees", "Critical?"],
    [
        ["ai_service: content_safe", "anthropic.*Error", "YES (fail-open)", "NOTHING (silent)", "CRITICAL"],
        ["ai_service: non-stream gen", "anthropic.*Error", "NO RETRY", "500 error", "HIGH"],
        ["ai_service: stream gen", "anthropic.*Error", "YES (retry 2x)", "SSE error event", "OK"],
        ["email: batch send", "Exception per-email", "YES (skip)", "Silent skip", "CRITICAL"],
        ["email: inspiration footer", "Exception", "YES (swallow)", "Missing footer", "MEDIUM"],
        ["google: credential refresh", "RefreshError", "pass (bare)", "Silent stale", "CRITICAL"],
        ["gmail: parse_message", "KeyError/TypeError", "NO", "500 error", "HIGH"],
        ["gcs: upload/download", "GoogleCloudError", "NO", "500 error", "HIGH"],
        ["wallet: debit_wallet", "DB flush error", "NO", "Partial debit", "CRITICAL"],
        ["notification: DB add", "DB error", "NO", "500 error", "MEDIUM"],
        ["file: empty extraction", "N/A", "PARTIAL", "Empty guide", "HIGH"],
        ["scheduler: job duplicate", "N/A", "NO", "Duplicate email", "MEDIUM"],
    ],
)

doc.add_heading("4 CRITICAL GAPS", level=2)

doc.add_paragraph(
    "1. Content safety fail-open (ai_service.py lines 86-88) -- If Anthropic API is down, "
    "unsafe content passes through to K-12 platform. Decision: Fix to fail-closed.",
    style="List Number",
)
doc.add_paragraph(
    "2. Email delivery not tracked (email_service.py lines 67-68) -- Batch emails silently "
    "skip failed recipients. Parent never knows they missed a notification.",
    style="List Number",
)
doc.add_paragraph(
    "3. Google credential refresh silently ignored (google_classroom.py lines 143-147) -- "
    "Bare 'pass' on refresh failure. No logging. Downstream calls fail mysteriously.",
    style="List Number",
)
doc.add_paragraph(
    "4. Wallet debit non-atomic (wallet_service.py lines 84-91) -- DB flush failure after "
    "in-memory debit = inconsistent state. User may be double-charged on retry.",
    style="List Number",
)

doc.add_page_break()

# ============================================================
# 6. SECURITY & THREAT MODEL
# ============================================================
doc.add_heading("6. Section 3: Security & Threat Model", level=1)

doc.add_paragraph("Overall Security Posture: 7/10")

doc.add_heading("HIGH Risk Findings", level=2)
add_table(
    ["#", "Threat", "Likelihood", "Impact", "Status"],
    [
        ["1", "LLM Prompt Injection", "High", "High", "Partial -- custom_prompt and assignment_description bypass safety"],
        ["2", "Rate Limiting (Cloud Run)", "High", "Medium", "NOT mitigated -- in-memory resets per instance"],
        ["3", "Content Safety Gaps (K-12)", "Medium", "High", "Partial -- teacher uploads, notes, messages unscanned"],
        ["4", "Token Blacklist Persistence", "Medium", "High", "Partial -- resets on cold start"],
    ],
)

doc.add_heading("What's Well-Secured", level=2)
items = [
    "IDOR protections: Tasks, messages, notes, grades properly scoped with whitelist patterns",
    "Admin endpoints: require_role(ADMIN) on all admin routes",
    "SQL injection: escape_like() + SQLAlchemy parameterization throughout",
    "Password policy: 8-char min with uppercase/lowercase/digit/special",
    "CI security: Bandit (SAST), GitLeaks (secrets), pip-audit + npm audit",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("10x Security Target for School Board Adoption", level=2)
add_table(
    ["Current (7/10)", "10x Target (9.5/10)"],
    [
        ["RBAC + IDOR guards", "All current controls"],
        ["JWT + blacklist", "+ Redis-backed rate limiting"],
        ["In-memory rate limiting", "+ Content moderation on ALL user-generated content"],
        ["Partial prompt guard", "+ Prompt injection defense on ALL AI inputs"],
        ["No content scanning", "+ SOC2 Type II certification"],
        ["No compliance cert", "+ Annual penetration test"],
        ["No pen test", "+ FERPA compliance documentation"],
    ],
)

doc.add_page_break()

# ============================================================
# 7. DATA FLOW & EDGE CASES
# ============================================================
doc.add_heading("7. Section 4: Data Flow & Interaction Edge Cases", level=1)

doc.add_heading("Shadow Path Analysis", level=2)
add_table(
    ["Node", "Shadow Path", "Handled?", "Gap?"],
    [
        ["UPLOAD: nil file", "No file attached", "YES", "OK"],
        ["UPLOAD: 0 bytes", "Empty file", "PARTIAL", "Unclear error message"],
        ["VALIDATION: too large", ">20MB", "YES", "OK"],
        ["VALIDATION: spoofed magic", ".pdf but .exe", "YES", "OK"],
        ["EXTRACTION: empty text", "Image-only PDF", "PARTIAL", "GAP -- guide from nothing"],
        ["EXTRACTION: corrupt PDF", "Malformed", "PARTIAL", "GAP -- may return empty"],
        ["AI GEN: API down", "Unreachable", "PARTIAL", "Inconsistent stream vs non-stream"],
        ["AI GEN: malformed JSON", "Invalid quiz", "PARTIAL", "GAP -- no schema validation"],
        ["DISPLAY: missing images", "IMG markers no file", "YES", "OK -- fallback section"],
    ],
)

doc.add_heading("Interaction Edge Cases", level=2)
add_table(
    ["Interaction", "Edge Case", "Handled?"],
    [
        ["File upload", "Double-click submit", "YES -- useRef guard + 60s dedup"],
        ["File upload", "Navigate away mid-upload", "PARTIAL -- no resume"],
        ["Study guide stream", "User navigates away", "YES -- abort controller"],
        ["Quiz submission", "Double-click", "UNCLEAR -- needs verification"],
        ["Dashboard", "50+ courses", "PARTIAL -- no virtual scrolling"],
        ["Scheduled jobs", "Run twice (Cloud Run)", "NO -- no dedup"],
        ["Wallet", "Double-debit race", "NO -- no idempotency key"],
        ["XP award", "Rapid-fire gaming", "YES -- anti-gaming cooldowns"],
    ],
)

doc.add_page_break()

# ============================================================
# 8. CODE QUALITY
# ============================================================
doc.add_heading("8. Section 5: Code Quality Review", level=1)

doc.add_heading("DRY Violations", level=2)
add_table(
    ["Violation", "Where", "Impact"],
    [
        ["formatRelativeDate()", "ParentDashboard + StudentDashboard", "Identical function duplicated"],
        ["Urgency tier calculation", "ParentDashboard + StudentDashboard", "Same logic duplicated"],
        ["Email template loading", "auth.py lines 50, 71", "Should be in EmailService"],
        ["Email lowercase", "auth.py lines 128, 148", "Normalize once at entry"],
        ["Local type redefinitions", "Dashboard components", "Should import from api types"],
        ["Error message extraction", "10+ components", "No shared utility"],
        ["localStorage try/catch", "Every dashboard", "No useLocalStorage hook"],
    ],
)

doc.add_heading("Cyclomatic Complexity Flags", level=2)
add_table(
    ["Method/Component", "Branches", "Issue"],
    [
        ["ParentDashboard.tsx render", "20+", "God component -- needs decomposition"],
        ["StudentDashboard.tsx render", "15+", "God component"],
        ["auth.py register()", "10+", "Extract to AuthService"],
        ["study.py generate endpoints", "8+ per endpoint", "Extract to StudyService"],
        ["main.py startup migrations", "30+ if-checks", "Extract to MigrationRunner"],
    ],
)

doc.add_heading("Under-Engineering Concerns", level=2)
add_table(
    ["Area", "Issue"],
    [
        ["No repository layer", "Services query DB directly; knowledge graph will duplicate"],
        ["No error taxonomy", "Bare Exception catches or no catches at all"],
        ["No circuit breaker", "Every request tries failed external APIs"],
        ["No idempotency", "Wallet debits, quiz saves lack idempotency keys"],
        ["No request correlation", "No trace ID flows through the system"],
    ],
)

doc.add_page_break()

# ============================================================
# 9. TEST REVIEW
# ============================================================
doc.add_heading("9. Section 6: Test Review", level=1)

doc.add_paragraph("Backend: 195 tests across 64 files (~21,900 lines)")
doc.add_paragraph("Frontend: 258 tests (vitest) across 18 files")
doc.add_paragraph("E2E: 0 tests | Load: 0 tests")

doc.add_heading("Test Coverage Map", level=2)
add_table(
    ["Area", "Unit", "Integration", "E2E", "Gap?"],
    [
        ["Auth", "YES", "YES", "NO", "E2E gap"],
        ["Study guide generation", "YES", "YES (streaming)", "NO", "E2E gap"],
        ["XP/Streak/Badge", "YES (40+)", "YES", "NO", "OK"],
        ["File upload", "PARTIAL", "YES", "NO", "Missing corrupt file tests"],
        ["Wallet credit/debit", "YES", "YES", "NO", "Missing race condition tests"],
        ["Google Classroom", "YES", "YES", "NO", "Missing OAuth refresh test"],
        ["Email delivery", "PARTIAL", "YES", "NO", "Missing batch failure test"],
        ["Scheduled jobs", "PARTIAL", "PARTIAL", "NO", "Missing job duplicate tests"],
        ["Frontend components", "YES (258)", "NO", "NO", "No Playwright/Cypress"],
    ],
)

doc.add_heading("Critical Missing Tests", level=2)
items = [
    "Upload corrupt PDF -> verify clear error, not blank study guide",
    "Two concurrent wallet debits -> verify only one succeeds",
    "Anthropic 500 -> verify non-streaming retries (streaming has it)",
    "Google OAuth token expires mid-sync -> verify re-auth prompt",
    "Scheduled digest runs twice -> verify no duplicate emails",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_page_break()

# ============================================================
# 10. PERFORMANCE
# ============================================================
doc.add_heading("10. Section 7: Performance Review", level=1)

doc.add_heading("Slowest Endpoints (P99 Latency)", level=2)
add_table(
    ["Endpoint", "P99", "Bottleneck", "Severity"],
    [
        ["GET /courses/students/search", "10-18s", "Full table scan + N+1", "CRITICAL"],
        ["POST /study/generate", "18-25s", "AI generation (expected)", "MEDIUM"],
        ["POST /users/me/export", "1.5-3s", "N+1 on assignments", "HIGH"],
        ["GET /courses/teachers/search", "1-3s", "Lazy loading .user", "MEDIUM"],
    ],
)

doc.add_heading("N+1 Query Patterns Found", level=2)
items = [
    "/courses/students/search -- loads ALL students, then 1 User query per student. 1000 students = 1001 queries.",
    "/courses/teachers/search -- lazy-loads .user per teacher without eager loading.",
    "Data export _collect_assignments() -- queries Assignment per StudentAssignment record.",
    "Data export _collect_children_data() -- queries Student + User per child.",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Missing Database Indexes", level=2)
add_table(
    ["Model", "Column", "Used In"],
    [
        ["Student", "user_id", "Every parent view, dashboard, export"],
        ["CourseContent", "created_by_user_id", "Content creator filtering"],
        ["StudyGuide", "guide_type", "Filtering guides vs quizzes vs flashcards"],
    ],
)

doc.add_page_break()

# ============================================================
# 11. OBSERVABILITY
# ============================================================
doc.add_heading("11. Section 8: Observability & Debuggability", level=1)

add_table(
    ["Aspect", "Status"],
    [
        ["Logging", "Good baseline -- rotating files, request middleware"],
        ["Structured logging", "NO -- text-based, can't query in GCP"],
        ["Metrics", "NO -- no Prometheus/StatsD/CloudMonitoring"],
        ["Tracing", "NO -- no trace IDs, no request correlation"],
        ["Alerting", "NO -- no Sentry/PagerDuty"],
        ["Dashboards", "NO -- no Grafana/Cloud Monitoring"],
        ["Health check", "YES -- /health returns version + environment"],
        ["Frontend errors", "Partial -- POST /api/errors/log"],
    ],
)

doc.add_paragraph(
    "10x Target: The 'Observatory Dashboard' -- a single screen showing platform health, "
    "user activity, AI cost tracking, job execution status, and error rates in real-time."
)

doc.add_page_break()

# ============================================================
# 12. DEPLOYMENT
# ============================================================
doc.add_heading("12. Section 9: Deployment & Rollout", level=1)

add_table(
    ["Aspect", "Status", "Detail"],
    [
        ["Migration safety", "RISKY", "Raw ALTER TABLE in startup. No rollback."],
        ["Feature flags", "MINIMAL", "Only WAITLIST_ENABLED, GOOGLE_CLASSROOM_ENABLED"],
        ["Rollback plan", "MANUAL", "Git revert + redeploy (~5 min)"],
        ["Zero-downtime", "YES", "Cloud Run rolling update + readiness gate"],
        ["Canary deploy", "NO", "All-or-nothing"],
        ["Staging environment", "NO", "Deploys directly to production"],
        ["Post-deploy verification", "MANUAL", "Health check only"],
    ],
)

doc.add_paragraph(
    "10x Target: Push to main -> CI -> Docker -> Staging -> Smoke tests -> Canary (10% traffic) "
    "-> Monitor -> 50% -> 100%. Auto-rollback on error spike > 2x baseline."
)

doc.add_page_break()

# ============================================================
# 13. LONG-TERM TRAJECTORY
# ============================================================
doc.add_heading("13. Section 10: Long-Term Trajectory", level=1)

doc.add_heading("Technical Debt Inventory", level=2)
add_table(
    ["Debt Type", "Items", "Impact"],
    [
        ["Code debt", "God components (893/925 LOC), 45-import routes, duplicate utils", "Slows feature velocity"],
        ["Architecture debt", "No event bus, no repository layer, logic in routes", "Coupling increases per feature"],
        ["Operational debt", "No structured logging, no metrics, no alerting", "Can't debug production"],
        ["Testing debt", "Zero E2E, SQLite-only tests, no load tests", "PG divergences cause bugs"],
        ["Documentation debt", "No TODOS.md (now created), no ADRs", "Knowledge concentration"],
        ["Hygiene debt", "49 git stashes, untracked docs", "Cognitive overhead"],
    ],
)

doc.add_paragraph("Reversibility Rating: 3/5")

doc.add_heading("Phase Trajectory", level=2)
add_table(
    ["Phase", "Timeline", "Key Features", "Architecture Needs"],
    [
        ["Phase 2", "Sept 2026", "XP, Streaks, Badges, Multilingual, Pomodoro, Report Cards", "Event bus, Feature flags, E2E tests"],
        ["Phase 3", "2027", "Ontario Curriculum, Prereq Engine, School Board Integration", "Knowledge graph, Plugin architecture, Multi-tenant"],
        ["Phase 4", "2027-28", "Tutor Marketplace, Booking, Payment, AI Matching", "Payment domain, Marketplace domain, Trust/safety"],
    ],
)

doc.add_page_break()

# ============================================================
# 14. TODOS
# ============================================================
doc.add_heading("14. TODOS Created", level=1)

doc.add_heading("P0 -- Critical (Before April 14)", level=2)
add_table(
    ["ID", "Description", "Effort"],
    [
        ["TODO-001", "Content Safety Fail-Closed Gate -- block AI generation when safety API unavailable", "S (1 day)"],
        ["TODO-002", "Fix 4 Critical Silent Failures -- wallet atomicity, Google credential logging, email tracking, empty extraction gate", "M (2 days)"],
    ],
)

doc.add_heading("P1 -- High (Before September Bundle)", level=2)
add_table(
    ["ID", "Description", "Effort"],
    [
        ["TODO-003", "Lightweight In-Process Event Bus -- decouple 53 route files from cross-cutting concerns", "L (3-5 days)"],
        ["TODO-004", "Structured JSON Logging + Request Correlation IDs", "M (2 days)"],
        ["TODO-005", "Fix N+1 Queries + Add Missing Indexes", "S (1 day)"],
        ["TODO-006", "Redis-Backed Rate Limiting -- replace in-memory slowapi", "M (1 day + infra)"],
    ],
)

doc.add_heading("P2 -- Medium (Backlog)", level=2)
add_table(
    ["ID", "Description", "Effort"],
    [
        ["TODO-007", "Audit and Clean 49 Git Stashes", "S (30 min)"],
        ["TODO-008", "Migration Framework (Alembic or equivalent)", "L (3-5 days)"],
    ],
)

doc.add_heading("Vision -- Delight Opportunities", level=2)
add_table(
    ["ID", "Description", "Effort"],
    [
        ["DELIGHT-001", "Study Streak Shared with Parent -- milestone notifications", "S (30 min)"],
        ["DELIGHT-002", "Quiz of the Day -- auto-generated daily challenge", "M (2 hours)"],
        ["DELIGHT-003", "Teacher Gratitude -- one-tap thank you with counter", "S (1 hour)"],
        ["DELIGHT-004", "Smart Study Time Suggestions -- data-driven insights", "S (1 hour)"],
        ["DELIGHT-005", "Weekly Family Report Card Email -- shareable, viral", "M (2 hours)"],
    ],
)

doc.add_page_break()

# ============================================================
# 15. COMPLETION SUMMARY
# ============================================================
doc.add_heading("15. Completion Summary", level=1)

add_table(
    ["Item", "Result"],
    [
        ["Mode selected", "SCOPE EXPANSION (Pure 10x Vision)"],
        ["System Audit", "53 routes, 47 models, 50+ services, 49 stashes, C- arch"],
        ["Step 0", "EXPANSION + event bus before Sept bundle"],
        ["Section 1 (Architecture)", "1 issue: no event bus (decided: build)"],
        ["Section 2 (Errors)", "12 error paths mapped, 4 CRITICAL GAPS"],
        ["Section 3 (Security)", "16 threats assessed, 4 High severity"],
        ["Section 4 (Data/UX)", "14 shadow paths mapped, 5 unhandled"],
        ["Section 5 (Quality)", "7 DRY violations, 5 complexity flags"],
        ["Section 6 (Tests)", "Diagram produced, 5 critical test gaps"],
        ["Section 7 (Performance)", "4 N+1 patterns, 3 missing indexes"],
        ["Section 8 (Observability)", "5 gaps (no metrics/traces/alerts)"],
        ["Section 9 (Deployment)", "3 risks (no staging/canary/migrations)"],
        ["Section 10 (Future)", "Reversibility: 3/5, 6 debt categories"],
        ["NOT in scope", "7 items"],
        ["What already exists", "10 reusable subsystems"],
        ["Error/rescue registry", "12 methods, 4 CRITICAL GAPS"],
        ["TODOS.md updates", "8 items proposed, all accepted"],
        ["Delight opportunities", "5 identified, all accepted"],
        ["Diagrams produced", "7 (arch, event bus, coupling, data flow, observatory, deploy, trajectory)"],
        ["Unresolved decisions", "0"],
    ],
)

doc.add_paragraph()
p = doc.add_paragraph(
    "Generated by Claude Opus 4.6 (1M context) -- Mega Plan Review Skill -- 2026-03-22"
)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].italic = True

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), "ClassBridge_CEO_Platform_Review.docx")
doc.save(output_path)
print(f"Word document saved to: {output_path}")
