"""Generate School Report Card System — Complete Analysis & Design Document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from datetime import datetime


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


def add_table_row(table, cells, bold=False, header=False):
    """Add a row to a table with cell values."""
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        if bold or header:
            run.bold = True
        if header:
            set_cell_shading(cell, '4F46E5')
            run.font.color.rgb = RGBColor(255, 255, 255)
    return row


def create_document():
    doc = Document()

    # ── Styles ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Calibri'
        h.font.color.rgb = RGBColor(79, 70, 229)  # Indigo

    # ── Title Page ──
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('School Report Card\nUpload & AI Analysis')
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(79, 70, 229)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Complete System Analysis, Design & Requirements Document')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f'ClassBridge (EMAI) Platform\nVersion 1.0 | {datetime.now().strftime("%B %d, %Y")}\nFeature Reference: #2286 (Section 6.121)')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_page_break()

    # ── Table of Contents ──
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Requirements & User Stories',
        '3. System Architecture',
        '4. Database Design',
        '5. API Specification',
        '6. AI Analysis Engine',
        '7. UI/UX Design',
        '8. Implementation Status',
        '9. Implementation Plan',
        '10. Test Plan',
        '11. Future Enhancements',
        '12. Appendix: File Reference',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('1. Executive Summary', level=1)

    doc.add_paragraph(
        'The School Report Card Upload & AI Analysis system enables parents to upload their '
        'children\'s Ontario school-issued report cards (PDF or image) and receive AI-powered '
        'analysis. The system extracts text via OCR, auto-detects metadata (grade level, school, '
        'term, date), and generates structured analysis including per-subject feedback, learning '
        'skills assessment, improvement areas, and actionable parent tips.'
    )
    doc.add_paragraph(
        'Additionally, parents can request a Career Path Analysis that aggregates data across '
        'multiple report cards to identify academic strengths, grade trends, and suggest 3-5 '
        'career paths with Ontario-specific course recommendations.'
    )

    doc.add_heading('Current State', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    add_table_row(table, ['Component', 'Status', 'Notes'], header=True)
    rows = [
        ['Backend API (6 endpoints)', 'COMPLETE', 'Deployed to Cloud Run'],
        ['Database Models', 'COMPLETE', '2 tables with indexes'],
        ['AI Service Layer', 'COMPLETE', 'Ontario-specific prompts'],
        ['File Processing (PDF/OCR)', 'COMPLETE', 'PyPDF2 + Vision API'],
        ['Frontend Components', 'COMPLETE', '4 components + API client'],
        ['Frontend Routing', 'MISSING', 'Page not accessible (#2352)'],
        ['Sidebar Navigation', 'MISSING', 'No nav link for parents'],
        ['Frontend Tests', 'MISSING', 'Zero coverage'],
        ['Backend Tests', 'PARTIAL', 'No career path or cache tests'],
    ]
    for r in rows:
        row = add_table_row(table, r)
        # Color the status cell
        status_cell = row.cells[1]
        if r[1] == 'COMPLETE':
            set_cell_shading(status_cell, 'D1FAE5')
        elif r[1] == 'MISSING':
            set_cell_shading(status_cell, 'FEE2E2')
        elif r[1] == 'PARTIAL':
            set_cell_shading(status_cell, 'FEF3C7')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 2. REQUIREMENTS & USER STORIES
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('2. Requirements & User Stories', level=1)

    doc.add_heading('2.1 User Stories', level=2)
    stories = [
        ('US-1: Upload Report Cards',
         'As a parent, I want to upload my child\'s school report card (PDF or image) '
         'so that I can get AI-powered analysis of their academic performance.',
         ['Parent must have at least 1 linked child',
          'Supported formats: PDF, JPG, JPEG, PNG',
          'Max file size: 30 MB per file',
          'Max files per upload: 10',
          'System auto-extracts metadata (grade, school, term, date)',
          'Upload succeeds even if text extraction fails']),
        ('US-2: View Report Card List',
         'As a parent, I want to see all uploaded report cards for each of my children '
         'with their analysis status.',
         ['List grouped by child (tab selector)',
          'Each card shows: filename, term, grade, school, date, analysis badge',
          'Sorted by report date (newest first)',
          'Archived cards are hidden']),
        ('US-3: AI Report Card Analysis',
         'As a parent, I want to get a detailed AI analysis of a report card that helps '
         'me understand my child\'s performance.',
         ['Analysis covers: teacher feedback, per-subject grades, learning skills, improvements, parent tips',
          'Minimum 50 characters of extracted text required',
          'Analysis is cached (instant on re-view)',
          'Costs 1 AI credit from wallet',
          'Ontario curriculum-aware (Achievement Levels 1-4, Learning Skills E/G/S/N)']),
        ('US-4: Career Path Analysis',
         'As a parent, I want to see career path suggestions based on all my child\'s '
         'report cards so I can help guide their future.',
         ['Requires at least 1 report card with text',
          'Aggregates data across all uploaded cards',
          'Shows: strengths, grade trends, 3-5 career suggestions with reasoning',
          'Includes Ontario course code recommendations',
          'Cached by combined content hash']),
        ('US-5: Delete Report Card',
         'As a parent, I want to delete a report card I no longer need.',
         ['Soft-delete (archive) preserves data for admin recovery',
          'Confirmation dialog before deletion',
          'Deleted cards excluded from career path analysis']),
    ]
    for title, desc, criteria in stories:
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)
        doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
        for c in criteria:
            doc.add_paragraph(c, style='List Bullet 2')

    doc.add_heading('2.2 Edge Cases & Error Handling', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    add_table_row(table, ['Scenario', 'Expected Response', 'HTTP Status'], header=True)
    errors = [
        ['Invalid file type (.docx, .txt)', '"Invalid file type"', '400'],
        ['File exceeds 30 MB', '"File size exceeds limit"', '400'],
        ['Upload > 10 files', '"Maximum 10 files per upload"', '400'],
        ['Text extraction < 50 chars', 'Upload OK; Analysis blocked', '400 on analyze'],
        ['AI usage limit exhausted', '"AI usage limit reached"', '429'],
        ['Parent accesses other child', '"Not your child"', '403'],
        ['Career path with 0 cards', '"No report cards found"', '400'],
        ['Report card not found', '"Report card not found"', '404'],
        ['Non-parent role accesses', 'Forbidden', '403'],
    ]
    for e in errors:
        add_table_row(table, e)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 3. SYSTEM ARCHITECTURE
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('3. System Architecture', level=1)

    doc.add_heading('3.1 Architecture Overview', level=2)
    doc.add_paragraph(
        'The system follows the standard ClassBridge architecture: FastAPI backend with '
        'SQLAlchemy models, Pydantic schemas, and service layer. The frontend is React 19 '
        'with TypeScript, using Axios for API calls.'
    )

    # Architecture flow
    doc.add_heading('3.2 Data Flow', level=2)
    flow_text = (
        'Upload Flow:\n'
        '  Parent -> React Upload Modal -> POST /api/school-report-cards/upload\n'
        '    -> File validation (type, size) -> save_file() (local/GCS)\n'
        '    -> process_file() (text extraction: PDF/OCR)\n'
        '    -> extract_metadata() (regex: grade, school, term, date)\n'
        '    -> DB insert (SchoolReportCard) -> Response\n\n'
        'Analysis Flow:\n'
        '  Parent -> "Analyze Now" -> POST /api/school-report-cards/{id}/analyze\n'
        '    -> Cache check (content_hash) -> if cached, return instantly\n'
        '    -> check_ai_usage() (wallet/limit check)\n'
        '    -> analyze_report_card() (Claude API, Ontario prompts)\n'
        '    -> DB insert (SchoolReportCardAnalysis) -> increment_ai_usage()\n'
        '    -> Response with structured analysis\n\n'
        'Career Path Flow:\n'
        '  Parent -> "Career Path" -> POST /api/school-report-cards/{student_id}/career-path\n'
        '    -> Load all cards with text -> compute combined_hash\n'
        '    -> Cache check -> if cached, return instantly\n'
        '    -> Load all existing full analyses\n'
        '    -> generate_career_path() (Claude API, career guidance prompts)\n'
        '    -> DB insert (analysis_type="career_path") -> Response'
    )
    p = doc.add_paragraph()
    run = p.add_run(flow_text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'

    doc.add_heading('3.3 Technology Stack', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    add_table_row(table, ['Component', 'Technology'], header=True)
    tech = [
        ['Backend Framework', 'FastAPI (Python 3.13+)'],
        ['ORM', 'SQLAlchemy 2.0'],
        ['Validation', 'Pydantic v2'],
        ['AI Provider', 'Anthropic Claude (claude-sonnet-4-6)'],
        ['Text Extraction', 'PyPDF2 + Anthropic Vision API (OCR)'],
        ['File Storage', 'Local (dev) / Google Cloud Storage (prod)'],
        ['Database', 'SQLite (dev) / PostgreSQL (prod)'],
        ['Frontend', 'React 19 + TypeScript + Vite'],
        ['HTTP Client', 'Axios with JWT Bearer interceptor'],
        ['Rate Limiting', 'SlowAPI (per-user/IP)'],
        ['Auth', 'JWT + RBAC (require_role decorator)'],
    ]
    for t in tech:
        add_table_row(table, t)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 4. DATABASE DESIGN
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('4. Database Design', level=1)

    doc.add_heading('4.1 SchoolReportCard Table', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    add_table_row(table, ['Column', 'Type', 'Constraints', 'Description'], header=True)
    cols = [
        ['id', 'Integer', 'PK, auto-increment', 'Primary key'],
        ['student_id', 'Integer', 'FK -> students.id, CASCADE', 'Owning student'],
        ['uploaded_by_user_id', 'Integer', 'FK -> users.id, SET NULL', 'Parent who uploaded'],
        ['original_filename', 'String(500)', 'NOT NULL', 'Original filename'],
        ['file_path', 'String(500)', 'nullable', 'Local storage path (dev)'],
        ['gcs_path', 'String(500)', 'nullable', 'GCS path (prod)'],
        ['file_size', 'Integer', 'nullable', 'File size in bytes'],
        ['mime_type', 'String(100)', 'nullable', 'MIME type'],
        ['text_content', 'Text', 'nullable', 'Full extracted text (OCR/PDF)'],
        ['term', 'String(100)', 'nullable', 'e.g., "Term 1", "Semester Two Interim"'],
        ['grade_level', 'String(20)', 'nullable', 'e.g., "08", "10", "JK"'],
        ['school_name', 'String(255)', 'nullable', 'School name'],
        ['report_date', 'Date', 'nullable', 'Report issue date'],
        ['school_year', 'String(20)', 'nullable', 'e.g., "2025-2026"'],
        ['created_at', 'DateTime(tz)', 'server_default=now()', 'Upload timestamp'],
        ['archived_at', 'DateTime(tz)', 'nullable', 'Soft-delete marker'],
    ]
    for c in cols:
        add_table_row(table, c)

    doc.add_heading('4.2 SchoolReportCardAnalysis Table', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    add_table_row(table, ['Column', 'Type', 'Constraints', 'Description'], header=True)
    cols2 = [
        ['id', 'Integer', 'PK, auto-increment', 'Primary key'],
        ['report_card_id', 'Integer', 'FK -> school_report_cards.id, CASCADE, nullable', 'NULL for career_path type'],
        ['student_id', 'Integer', 'FK -> students.id, CASCADE', 'Student being analyzed'],
        ['analysis_type', 'String(20)', 'NOT NULL', '"full" or "career_path"'],
        ['content', 'Text', 'NOT NULL', 'JSON string with structured analysis'],
        ['content_hash', 'String(64)', 'nullable', 'SHA-256 for cache deduplication'],
        ['ai_model', 'String(50)', 'nullable', 'e.g., "claude-sonnet-4-6"'],
        ['prompt_tokens', 'Integer', 'nullable', 'Input tokens consumed'],
        ['completion_tokens', 'Integer', 'nullable', 'Output tokens consumed'],
        ['estimated_cost_usd', 'Float', 'nullable', 'Estimated API cost'],
        ['created_at', 'DateTime(tz)', 'server_default=now()', 'Analysis timestamp'],
    ]
    for c in cols2:
        add_table_row(table, c)

    doc.add_heading('4.3 Indexes', level=2)
    doc.add_paragraph('ix_school_report_cards_student (student_id)', style='List Bullet')
    doc.add_paragraph('ix_school_report_cards_uploaded_by (uploaded_by_user_id)', style='List Bullet')
    doc.add_paragraph('ix_src_analyses_report_card (report_card_id)', style='List Bullet')
    doc.add_paragraph('ix_src_analyses_student_type (student_id, analysis_type)', style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 5. API SPECIFICATION
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('5. API Specification', level=1)

    endpoints = [
        {
            'method': 'POST',
            'path': '/api/school-report-cards/upload',
            'auth': 'PARENT',
            'rate': '10/minute',
            'desc': 'Upload 1-10 report card files (PDF/JPG/PNG) for a student.',
            'params': 'FormData: files[] (multipart), student_id (int), school_name (string, optional)',
            'response': 'UploadReportCardResponse { uploaded[], failures[], total_uploaded }',
        },
        {
            'method': 'GET',
            'path': '/api/school-report-cards/{student_id}',
            'auth': 'PARENT / ADMIN',
            'rate': '60/minute',
            'desc': 'List all non-archived report cards for a student with analysis status.',
            'params': 'Path: student_id (int)',
            'response': 'list[ReportCardItem] with has_text, has_analysis flags',
        },
        {
            'method': 'GET',
            'path': '/api/school-report-cards/{report_card_id}/analysis',
            'auth': 'PARENT / ADMIN',
            'rate': '60/minute',
            'desc': 'Retrieve cached analysis for a report card. Returns null if not yet analyzed.',
            'params': 'Path: report_card_id (int)',
            'response': 'FullAnalysisResponse or { "analysis": null }',
        },
        {
            'method': 'POST',
            'path': '/api/school-report-cards/{report_card_id}/analyze',
            'auth': 'PARENT',
            'rate': '5/minute',
            'desc': 'Trigger AI analysis. Returns cached result if available, otherwise generates new.',
            'params': 'Path: report_card_id (int)',
            'response': 'FullAnalysisResponse { report_card_id, content (dict), created_at }',
        },
        {
            'method': 'POST',
            'path': '/api/school-report-cards/{student_id}/career-path',
            'auth': 'PARENT',
            'rate': '5/minute',
            'desc': 'Generate career path suggestions from all report cards for a student.',
            'params': 'Path: student_id (int)',
            'response': 'CareerPathResponse { student_id, content (dict), report_cards_used, created_at }',
        },
        {
            'method': 'DELETE',
            'path': '/api/school-report-cards/{report_card_id}',
            'auth': 'PARENT',
            'rate': '30/minute',
            'desc': 'Soft-delete (archive) a report card. Sets archived_at timestamp.',
            'params': 'Path: report_card_id (int)',
            'response': '{ "status": "deleted" }',
        },
    ]

    for ep in endpoints:
        doc.add_heading(f'{ep["method"]} {ep["path"]}', level=3)
        doc.add_paragraph(ep['desc'])
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        for i, (k, v) in enumerate([
            ('Auth', ep['auth']),
            ('Rate Limit', ep['rate']),
            ('Parameters', ep['params']),
            ('Response', ep['response']),
            ('Method', ep['method']),
        ]):
            table.rows[i].cells[0].text = k
            table.rows[i].cells[1].text = v
            for cell in table.rows[i].cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 6. AI ANALYSIS ENGINE
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('6. AI Analysis Engine', level=1)

    doc.add_heading('6.1 Report Card Analysis', level=2)
    doc.add_paragraph(
        'Uses Anthropic Claude with an Ontario education specialist system prompt. '
        'Temperature: 0.3 (deterministic). Max tokens: 3000.'
    )
    doc.add_paragraph('Output structure:')
    analysis_fields = [
        ('teacher_feedback_summary', 'Consolidated 2-3 paragraph narrative from all teacher comments'),
        ('grade_analysis[]', 'Per-subject: subject, grade, median, level (1-4), teacher_comment, AI feedback'),
        ('learning_skills', 'Ratings (E/G/S/N) for 6 Ontario skills + pattern summary'),
        ('improvement_areas[]', 'Prioritized (high/medium/low) with specific guidance'),
        ('parent_tips[]', 'Actionable home-based activities per subject'),
        ('overall_summary', '2-3 sentence holistic assessment'),
    ]
    for field, desc in analysis_fields:
        p = doc.add_paragraph()
        run = p.add_run(f'{field}: ')
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(desc)
        run2.font.size = Pt(10)

    doc.add_heading('6.2 Career Path Analysis', level=2)
    doc.add_paragraph(
        'Uses Anthropic Claude with an Ontario career guidance specialist system prompt. '
        'Temperature: 0.5 (balanced creativity). Max tokens: 2000.'
    )
    doc.add_paragraph('Output structure:')
    career_fields = [
        ('strengths[]', 'Subjects consistently strong across multiple reports'),
        ('grade_trends[]', 'Per-subject trajectory (improving/stable/declining) with data points'),
        ('career_suggestions[]', '3-5 careers with reasoning, related subjects, Ontario course recommendations'),
        ('overall_assessment', 'Holistic academic profile and trajectory summary'),
    ]
    for field, desc in career_fields:
        p = doc.add_paragraph()
        run = p.add_run(f'{field}: ')
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(desc)
        run2.font.size = Pt(10)

    doc.add_heading('6.3 Metadata Extraction (Regex, No AI)', level=2)
    doc.add_paragraph(
        'Automatic extraction from report card text using regex patterns. No AI cost. '
        'Handles YRDSB elementary and secondary formats.'
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    add_table_row(table, ['Field', 'Pattern Examples', 'Output'], header=True)
    meta = [
        ['grade_level', '"Grade: 08", "Grade 10"', '"8", "10"'],
        ['school_name', '"School: Franklin Street Public School"', 'Full school name'],
        ['report_date', '"Date: 02/19/2026", "March 9, 2026"', 'Parsed date object'],
        ['term', '"Term 1", "Semester Two Interim"', 'Term string'],
        ['school_year', '"2025-2026" or inferred from date', '"2025-2026"'],
        ['board_name', '"York Region District School Board"', 'Board name'],
    ]
    for m in meta:
        add_table_row(table, m)

    doc.add_heading('6.4 Caching Strategy', level=2)
    doc.add_paragraph(
        'SHA-256 content hashing ensures identical content is never re-analyzed. '
        'Per-card analysis: keyed by report_card_id + analysis_type. '
        'Career path: keyed by combined hash of all card texts (sorted). '
        'New upload invalidates career path cache automatically.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 7. UI/UX DESIGN
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('7. UI/UX Design', level=1)

    doc.add_heading('7.1 Screen 1: Report Cards Landing', level=2)
    doc.add_paragraph(
        'Parent navigates via sidebar "Report Cards" link. The page displays:'
    )
    doc.add_paragraph('Header: "School Report Cards" with descriptive subtitle', style='List Bullet')
    doc.add_paragraph('Child Selector Tabs: horizontal tabs for each linked child', style='List Bullet')
    doc.add_paragraph('Action Bar: "Upload Report Card" (primary) + "Career Path Analysis" (secondary)', style='List Bullet')
    doc.add_paragraph('Report Card List: expandable cards with metadata and analysis badge', style='List Bullet')

    wireframe1 = (
        '+------------------------------------------------------------------+\n'
        '| [Sidebar]  |  School Report Cards                               |\n'
        '|            |  Upload, analyze, and track school report cards     |\n'
        '|  Home      |                                                    |\n'
        '|  My Kids   |  [ Child 1 ]  [ Child 2 ]  [ Child 3 ]            |\n'
        '| >Rep Cards |                                                    |\n'
        '|  AI Tools  |  [+ Upload Report Card]    [Career Path Analysis]  |\n'
        '|  Tasks     |                                                    |\n'
        '|  Messages  |  +----------------------------------------------+  |\n'
        '|            |  | report_card_term1.pdf          [Analyzed]    |  |\n'
        '|            |  | Term 1 | Grade 8 | Franklin PS | Feb 2026   |  |\n'
        '|            |  | [View Analysis]                              |  |\n'
        '|            |  +----------------------------------------------+  |\n'
        '|            |  | report_card_term2.pdf       [Not Analyzed]   |  |\n'
        '|            |  | Term 2 | Grade 8 | Franklin PS | Jun 2026   |  |\n'
        '|            |  | [Analyze Now]              [Delete]          |  |\n'
        '|            |  +----------------------------------------------+  |\n'
        '+------------------------------------------------------------------+'
    )
    p = doc.add_paragraph()
    run = p.add_run(wireframe1)
    run.font.size = Pt(8)
    run.font.name = 'Consolas'

    doc.add_heading('7.2 Screen 2: Upload Modal', level=2)
    doc.add_paragraph(
        'Drag-and-drop upload zone with file list, school name field, and progress states.'
    )
    wireframe2 = (
        '+------------------------------------------+\n'
        '|  Upload Report Cards              [X]    |\n'
        '|                                          |\n'
        '|  +------------------------------------+  |\n'
        '|  |                                    |  |\n'
        '|  |   Drag report cards here           |  |\n'
        '|  |   or click to browse               |  |\n'
        '|  |                                    |  |\n'
        '|  |   PDF, JPG, PNG (max 30MB each)    |  |\n'
        '|  +------------------------------------+  |\n'
        '|                                          |\n'
        '|  Files:                                  |\n'
        '|  [PDF] report_card_t1.pdf  2.1MB  [x]   |\n'
        '|  [IMG] report_card_t2.jpg  1.4MB  [x]   |\n'
        '|                                          |\n'
        '|  School Name (optional):                 |\n'
        '|  [Franklin Street Public School    ]     |\n'
        '|                                          |\n'
        '|  [Cancel]              [Upload 2 Files]  |\n'
        '+------------------------------------------+'
    )
    p = doc.add_paragraph()
    run = p.add_run(wireframe2)
    run.font.size = Pt(8)
    run.font.name = 'Consolas'

    doc.add_heading('7.3 Screen 3: Analysis View (Expanded Card)', level=2)
    doc.add_paragraph(
        'When a parent clicks "View Analysis" on an analyzed card, it expands to show 6 collapsible sections.'
    )
    wireframe3 = (
        '+------------------------------------------------------+\n'
        '| report_card_term1.pdf                    [Analyzed]  |\n'
        '| Term 1 | Grade 8 | Franklin PS | Feb 2026           |\n'
        '|                                                      |\n'
        '| Overall Summary                                      |\n'
        '| "John shows strong engagement across subjects..."    |\n'
        '|                                                      |\n'
        '| v Teacher Feedback Summary                           |\n'
        '|   "Teachers consistently praise John\'s..."           |\n'
        '|                                                      |\n'
        '| v Grade Analysis                                     |\n'
        '|   +------+-------+--------+-------+---------------+  |\n'
        '|   |Subj  |Grade  |Median  |Level  |Feedback       |  |\n'
        '|   +------+-------+--------+-------+---------------+  |\n'
        '|   |Math  |71%    |84%     |  2    |Below median...|  |\n'
        '|   |Eng   |82%    |78%     |  3    |Above median...|  |\n'
        '|   |Sci   |75%    |80%     |  3    |Near median... |  |\n'
        '|   +------+-------+--------+-------+---------------+  |\n'
        '|                                                      |\n'
        '| v Learning Skills                                    |\n'
        '|   Responsibility [G] Organization [G] Initiative [S] |\n'
        '|   Collaboration [G] Self-Regulation [G] Indep. [S]  |\n'
        '|                                                      |\n'
        '| v Improvement Areas                                  |\n'
        '|   [HIGH] Math - Metric Conversions                   |\n'
        '|   [MED]  Science - Lab Report Writing                |\n'
        '|                                                      |\n'
        '| v Parent Tips                                        |\n'
        '|   * Practice conversions using cooking (Math)        |\n'
        '|   * Read science articles together (Science)         |\n'
        '|                                                      |\n'
        '| [Hide Analysis]                          [Delete]    |\n'
        '+------------------------------------------------------+'
    )
    p = doc.add_paragraph()
    run = p.add_run(wireframe3)
    run.font.size = Pt(8)
    run.font.name = 'Consolas'

    doc.add_heading('7.4 Screen 4: Career Path Analysis', level=2)
    doc.add_paragraph(
        'Shown below the card list when parent clicks "Career Path Analysis".'
    )
    wireframe4 = (
        '+------------------------------------------------------+\n'
        '| Career Path Analysis          Based on 3 report cards|\n'
        '|                                                      |\n'
        '| Academic Strengths:                                  |\n'
        '| [Visual Arts] [English - Creative Writing] [Music]   |\n'
        '|                                                      |\n'
        '| Grade Trends:                                        |\n'
        '| Math        declining  78% -> 71% -> 71%             |\n'
        '| English     improving  72% -> 78% -> 82%             |\n'
        '| Visual Arts stable     85% -> 87% -> 86%             |\n'
        '|                                                      |\n'
        '| Suggested Career Paths:                              |\n'
        '| +------------------------+ +------------------------+|\n'
        '| | UX/UI Design           | | Graphic Design         ||\n'
        '| | Strong visual arts...  | | Creative aptitude...   ||\n'
        '| | Subjects: Art, Tech    | | Subjects: Art, Eng     ||\n'
        '| | Next: TGJ1O in Gr 9   | | Next: AVI2O in Gr 10  ||\n'
        '| +------------------------+ +------------------------+|\n'
        '| +------------------------+ +------------------------+|\n'
        '| | Journalism             | | Marketing              ||\n'
        '| | Strong English and...  | | Communication skills.. ||\n'
        '| | Subjects: Eng, Social  | | Subjects: Eng, Math    ||\n'
        '| | Next: ENG2D in Gr 10  | | Next: BMI3C in Gr 11  ||\n'
        '| +------------------------+ +------------------------+|\n'
        '|                                                      |\n'
        '| Overall Assessment:                                  |\n'
        '| "John shows strong creative aptitude with..."        |\n'
        '+------------------------------------------------------+'
    )
    p = doc.add_paragraph()
    run = p.add_run(wireframe4)
    run.font.size = Pt(8)
    run.font.name = 'Consolas'

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 8. IMPLEMENTATION STATUS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('8. Implementation Status', level=1)

    doc.add_heading('8.1 Backend (100% Complete)', level=2)
    items = [
        'All 6 API endpoints deployed and operational',
        'Parent-child RBAC authorization on all endpoints',
        'File upload with type/size validation',
        'Text extraction: PDF (PyPDF2) + Image OCR (Anthropic Vision API)',
        'Regex metadata extraction (grade, school, term, date, board)',
        'AI analysis with Ontario-specific prompts',
        'Career path analysis with multi-card aggregation',
        'SHA-256 content hashing for cache deduplication',
        'AI usage tracking with wallet integration',
        'Rate limiting on all endpoints',
        'Soft-delete with archiving',
        'GCS storage support (production)',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('8.2 Frontend Components (Complete but Orphaned)', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    add_table_row(table, ['Component', 'Lines', 'Status'], header=True)
    components = [
        ['ReportCardAnalysis.tsx (main page)', '529', 'Complete, NOT routed'],
        ['ReportCardUploadModal.tsx', '327', 'Complete, NOT used by main page'],
        ['ReportCardAnalysisView.tsx', '248', 'Complete, NOT used by main page'],
        ['CareerPathView.tsx', '115', 'Complete, NOT used by main page'],
        ['schoolReportCards.ts (API client)', '116', 'Complete'],
        ['ReportCardAnalysis.css', '524', 'Complete (responsive + dark mode)'],
        ['ReportCardUploadModal.css', '342', 'Complete'],
        ['ReportCardAnalysisView.css', '422', 'Complete'],
        ['CareerPathView.css', '250', 'Complete'],
    ]
    for c in components:
        add_table_row(table, c)

    doc.add_heading('8.3 Missing Pieces', level=2)
    missing = [
        ('Route in App.tsx', 'BLOCKING', 'Page is unreachable by users'),
        ('Sidebar nav link', 'BLOCKING', 'No way to navigate to the feature'),
        ('Component wiring', 'QUALITY', 'Dedicated sub-components exist but main page uses inline rendering'),
        ('Delete confirmation UI', 'UX', 'No confirmation before soft-delete'),
        ('Frontend tests', 'QUALITY', 'Zero test coverage for all report card components'),
        ('Backend career path tests', 'QUALITY', 'No tests for career path, cache, or file validation'),
    ]
    for name, sev, desc in missing:
        p = doc.add_paragraph()
        run = p.add_run(f'[{sev}] {name}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 9. IMPLEMENTATION PLAN
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('9. Implementation Plan', level=1)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    add_table_row(table, ['Phase', 'Description', 'Priority', 'Effort', 'Files'], header=True)
    phases = [
        ['1', 'Wire route + sidebar nav', 'P0 - Blocker', '30 min', 'App.tsx, DashboardLayout.tsx'],
        ['2', 'Component consolidation', 'P1 - Quality', '2-3 hrs', 'ReportCardAnalysis.tsx'],
        ['3', 'Test coverage (FE + BE)', 'P1 - Quality', '2-3 hrs', '4 test files'],
        ['4', 'Integration links', 'P2 - Nice to have', '15 min', 'MyKidsPage.tsx, ParentAITools.tsx'],
        ['5', 'Weekly digest integration', 'P3 - Backlog', '1-2 hrs', 'weekly_digest_service.py'],
    ]
    for ph in phases:
        add_table_row(table, ph)

    for phase_num, phase_title, phase_desc in [
        ('Phase 1', 'Wire Frontend Route + Navigation (P0)',
         'Add lazy import and route for /school-report-cards in App.tsx. '
         'Add "Report Cards" nav item with icon in DashboardLayout.tsx parent section.'),
        ('Phase 2', 'Component Consolidation (P1)',
         'Replace inline rendering in ReportCardAnalysis.tsx with dedicated sub-components: '
         'ReportCardUploadModal (drag-drop, multi-file), ReportCardAnalysisView (grade table, '
         'teacher comments), CareerPathView (sorted trends, career cards). Add delete confirmation.'),
        ('Phase 3', 'Test Coverage (P1)',
         'Create frontend tests for all 4 components (ReportCardAnalysis, UploadModal, AnalysisView, CareerPathView). '
         'Add backend tests for career path endpoint, cache behavior, and file validation.'),
        ('Phase 4', 'Integration Points (P2)',
         'Add "Report Cards" quick action on My Kids page. Add 4th tool card on ParentAITools page '
         'linking to /school-report-cards.'),
    ]:
        doc.add_heading(f'{phase_num}: {phase_title}', level=2)
        doc.add_paragraph(phase_desc)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 10. TEST PLAN
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('10. Test Plan', level=1)

    doc.add_heading('10.1 Existing Backend Tests (12 tests)', level=2)
    existing = [
        'TestMetadataExtraction: elementary, secondary, empty input',
        'TestUploadEndpoint: single PDF, requires parent role, unauthorized parent',
        'TestListEndpoint: list cards, unauthorized access',
        'TestAnalyzeEndpoint: analyze card, no text validation',
        'TestDeleteEndpoint: soft delete, not found',
    ]
    for t in existing:
        doc.add_paragraph(t, style='List Bullet')

    doc.add_heading('10.2 Backend Tests to Add', level=2)
    new_be = [
        'TestCareerPathEndpoint: success, no cards (400), cache hit, unauthorized (403)',
        'TestCacheBehavior: analyze returns cached, career path hash invalidation on new upload',
        'TestFileValidation: invalid extension (400), oversized file (400), >10 files (400)',
        'TestGetAnalysis: returns null when not yet analyzed',
    ]
    for t in new_be:
        doc.add_paragraph(t, style='List Bullet')

    doc.add_heading('10.3 Frontend Tests to Create', level=2)
    new_fe = [
        'ReportCardAnalysis.test.tsx: child selector, card list, analyze, view analysis, upload modal, career path, empty/error states (12 tests)',
        'ReportCardUploadModal.test.tsx: drag-drop, file validation, upload flow, error handling (8 tests)',
        'ReportCardAnalysisView.test.tsx: section rendering, toggles, grade coloring (6 tests)',
        'CareerPathView.test.tsx: strengths, trends, career cards (4 tests)',
    ]
    for t in new_fe:
        doc.add_paragraph(t, style='List Bullet')

    doc.add_heading('10.4 Manual Testing Checklist', level=2)
    manual = [
        'Login as parent with linked child -> navigate to /school-report-cards',
        'Upload a PDF report card -> verify metadata extraction',
        'Click "Analyze Now" -> verify AI analysis renders correctly',
        'Upload 2+ cards -> click "Career Path" -> verify trends and suggestions',
        'Click "Delete" -> confirm -> verify card removed from list',
        'Login as student -> verify /school-report-cards is blocked (403)',
        'Upload invalid file type (.docx) -> verify rejection',
        'Upload file > 30MB -> verify rejection',
    ]
    for m in manual:
        doc.add_paragraph(m, style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 11. FUTURE ENHANCEMENTS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('11. Future Enhancements', level=1)

    enhancements = [
        ('Re-analysis', 'Allow parents to re-analyze if AI result was unsatisfactory. Add "Re-analyze" button that clears cached analysis and regenerates.'),
        ('Term-over-term Comparison', 'Side-by-side comparison of two report cards for the same student, highlighting grade changes per subject.'),
        ('Admin Dashboard', 'Admin panel showing all uploaded report cards across families with aggregate statistics (uploads per week, popular subjects, common improvement areas).'),
        ('Mobile Support', 'React Native screens in ClassBridgeMobile for upload (camera capture) and analysis viewing. Backend API is already mobile-ready.'),
        ('Photo Capture (#1432)', 'Snap a report card photo with device camera and auto-import. Leverage existing Vision OCR for text extraction.'),
        ('Weekly Digest Integration', 'Include report card activity in the weekly family digest email: "2 report cards uploaded this week, 1 analyzed."'),
        ('Multi-language Support', 'Extend AI prompts to handle French immersion and other non-English report card formats common in Ontario.'),
        ('Sharing', 'Allow parents to share analysis results with teachers or tutors via a secure link.'),
    ]
    for title, desc in enhancements:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 12. APPENDIX: FILE REFERENCE
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('12. Appendix: File Reference', level=1)

    doc.add_heading('Backend Files', level=2)
    be_files = [
        ('app/api/routes/school_report_cards.py', 'API endpoints (521 lines)'),
        ('app/models/school_report_card.py', 'Database models (68 lines)'),
        ('app/schemas/school_report_card.py', 'Pydantic schemas (143 lines)'),
        ('app/services/school_report_card_service.py', 'AI service layer (427 lines)'),
        ('app/services/file_processor.py', 'Text extraction (PDF/OCR)'),
        ('app/services/ai_service.py', 'Claude API integration'),
        ('app/services/ai_usage.py', 'AI cost tracking + wallet'),
        ('app/services/storage_service.py', 'Local file storage'),
        ('app/services/gcs_service.py', 'Google Cloud Storage'),
        ('tests/test_school_report_cards.py', 'Backend tests (327 lines)'),
    ]
    for path, desc in be_files:
        p = doc.add_paragraph()
        run = p.add_run(f'{path}')
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
        p.add_run(f' — {desc}').font.size = Pt(9)

    doc.add_heading('Frontend Files', level=2)
    fe_files = [
        ('frontend/src/pages/parent/ReportCardAnalysis.tsx', 'Main page (529 lines)'),
        ('frontend/src/pages/parent/ReportCardAnalysis.css', 'Page styles (524 lines)'),
        ('frontend/src/components/parent/ReportCardUploadModal.tsx', 'Upload modal (327 lines)'),
        ('frontend/src/components/parent/ReportCardUploadModal.css', 'Modal styles (342 lines)'),
        ('frontend/src/components/parent/ReportCardAnalysisView.tsx', 'Analysis view (248 lines)'),
        ('frontend/src/components/parent/ReportCardAnalysisView.css', 'Analysis styles (422 lines)'),
        ('frontend/src/components/parent/CareerPathView.tsx', 'Career path view (115 lines)'),
        ('frontend/src/components/parent/CareerPathView.css', 'Career path styles (250 lines)'),
        ('frontend/src/api/schoolReportCards.ts', 'API client (116 lines)'),
    ]
    for path, desc in fe_files:
        p = doc.add_paragraph()
        run = p.add_run(f'{path}')
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
        p.add_run(f' — {desc}').font.size = Pt(9)

    doc.add_heading('Configuration', level=2)
    config = [
        ('ANTHROPIC_API_KEY', 'Required for AI analysis'),
        ('CLAUDE_MODEL', 'Default: claude-sonnet-4-6'),
        ('MAX_UPLOAD_SIZE_MB', 'Default: 30'),
        ('USE_GCS', 'Default: false (true in production)'),
        ('GCS_BUCKET_NAME', 'GCS bucket for file storage'),
    ]
    for key, desc in config:
        p = doc.add_paragraph()
        run = p.add_run(f'{key}')
        run.bold = True
        run.font.name = 'Consolas'
        p.add_run(f' — {desc}')

    # ── Footer ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Generated {datetime.now().strftime("%Y-%m-%d %H:%M")} | ClassBridge Platform')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(156, 163, 175)

    return doc


if __name__ == '__main__':
    doc = create_document()
    output_path = 'docs/School_Report_Card_System.docx'
    doc.save(output_path)
    print(f'Document saved to {output_path}')
