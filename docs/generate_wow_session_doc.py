"""Generate Word document summarizing the WOW Features Planning Session."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = str(val)
    doc.add_paragraph()
    return table


def main():
    doc = Document()

    # Title
    title = doc.add_heading("ClassBridge WOW Features", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Planning Session Summary")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}\n").font.size = Pt(10)
    meta.add_run("Author: Claude Code + Theepan Gnanasabapathy\n").font.size = Pt(10)
    meta.add_run("Status: Planning Complete — Ready for Development").font.size = Pt(10)

    doc.add_page_break()

    # Table of Contents
    add_heading(doc, "Table of Contents", 1)
    toc_items = [
        "1. Executive Summary",
        "2. The Problem: Why Parents Don't See the WOW Factor",
        "3. Solution: WOW Features Overview",
        "4. Feature 1: Smart Daily Briefing (§6.61)",
        "5. Feature 2: Help My Kid — One-Tap Actions (§6.62)",
        "6. Feature 3: Global Search + Smart Shortcuts (§6.17)",
        "7. Feature 4: Weekly Progress Pulse (§6.63)",
        "8. Feature 5: Parent-Child Study Link (§6.64)",
        "9. Feature 6: Dashboard Redesign (§6.65)",
        "10. Feature 7: Responsible AI Parent Tools (§6.66)",
        "11. Hybrid Search Strategy (AI Cost Analysis)",
        "12. Source Material Linking Design",
        "13. Parallel Development Plan",
        "14. GitHub Issues Created",
        "15. Files Updated",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")

    doc.add_page_break()

    # 1. Executive Summary
    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "Pilot user feedback revealed that parents don't see a clear reason to use ClassBridge. "
        'The feedback: "I don\'t see a WOW factor." This planning session designed 7 feature epics '
        "with 27 GitHub issues to transform ClassBridge from a passive information viewer into an "
        "active parenting tool."
    )
    doc.add_paragraph(
        "Key insight: ClassBridge today mirrors what Google Classroom already shows. "
        "The WOW comes when ClassBridge does things parents can't do anywhere else — "
        "proactively telling them what to do, letting them take action in one tap, "
        "and creating a feedback loop with their children."
    )

    add_heading(doc, "Core Principle: Parents First, Responsible AI", 2)
    doc.add_paragraph(
        "Every feature passes the Responsible AI test:\n"
        "• Does it require the student to DO something? ✓\n"
        "• Does it help the PARENT understand and engage? ✓\n"
        "• Could the student use it to avoid studying? ✗ → Don't build it\n\n"
        "ClassBridge is NOT an AI search engine. AI is used for generation only when "
        "explicitly requested. Search is SQL-based ($0 cost)."
    )

    doc.add_page_break()

    # 2. The Problem
    add_heading(doc, "2. The Problem: Why Parents Don't See the WOW Factor", 1)
    add_table(doc,
        ["What ClassBridge Does Today", "Parent Reaction"],
        [
            ["Shows assignments from Google Classroom", '"I can see that in Google Classroom"'],
            ["AI study guides/quizzes", '"Nice, but I don\'t know when to use it"'],
            ["Calendar view", '"Google Calendar does this"'],
            ["Messaging teachers", '"I can email them"'],
            ["Tasks", '"I have a to-do app"'],
        ]
    )
    doc.add_paragraph(
        "The missing piece: ClassBridge doesn't proactively tell parents what to DO. "
        "It's passive — parents have to go looking. Parents are busy. They want ClassBridge "
        "to come to them."
    )

    add_heading(doc, "2.1 The Two Layers: Infrastructure vs Intelligence", 2)
    doc.add_paragraph(
        "ClassBridge should absolutely keep all existing features — they're the foundation. "
        "The problem isn't that they exist, it's that they're not enough on their own to "
        "justify switching from tools parents already use."
    )
    add_table(doc,
        ["Layer", "Purpose", "Examples"],
        [
            ["Layer 1: Infrastructure (what we have)", "Mirror & organize school data",
             "Google Classroom sync, calendar, messaging, tasks, study guides"],
            ["Layer 2: Intelligence (the WOW)", "Tell parents what to DO with that data",
             "Daily Briefing, Help My Kid, Weak Spot Reports, Readiness Checks"],
        ]
    )
    doc.add_paragraph(
        "Layer 1 is table stakes — it gets parents in the door. But parents compare it "
        "1:1 against Google Classroom and shrug.\n\n"
        "Layer 2 is the moat — no other tool does this. Google Classroom doesn't tell a parent "
        '"Your child failed 3 geometry questions this week — here\'s a practice set." '
        "That's the WOW."
    )

    add_heading(doc, "Recommendation", 3)
    doc.add_paragraph(
        "1. Don't remove or de-emphasize existing features — they feed Layer 2\n"
        "2. Reposition them in the UI as supporting tools, not the headline\n"
        "3. Lead with proactive features on the dashboard — Daily Briefing front and center, "
        "study guides accessible but not the hero\n"
        "4. Marketing shift: Stop saying 'View your child's assignments' → Start saying "
        "'Know exactly how to help your child tonight'"
    )

    add_heading(doc, "The Car Analogy", 3)
    doc.add_paragraph(
        "Layer 1 is the engine. Layer 2 is the steering wheel. "
        "Parents don't buy a car because it has an engine — every car has one. "
        "They buy it because of how it drives."
    )

    doc.add_page_break()

    # 3. Solution Overview
    add_heading(doc, "3. Solution: WOW Features Overview", 1)
    add_table(doc,
        ["Priority", "Feature", "Epic #", "AI Cost", "WOW Impact"],
        [
            ["1", "Smart Daily Briefing", "#1403", "$0", "Highest — answers 'why open ClassBridge'"],
            ["2", "Help My Kid", "#1407", "~$0.02/use", "Highest — instant parent action"],
            ["3", "Global Search + Shortcuts", "#1410", "$0", "Medium — power users love it"],
            ["4", "Weekly Progress Pulse", "#1413", "$0", "High — passive value"],
            ["5", "Parent-Child Study Link", "#1414", "$0", "High — emotional connection"],
            ["6", "Dashboard Redesign", "#1415", "$0", "High — first impression"],
            ["7", "Responsible AI Tools", "#1421", "~$0.02/use", "Highest — 'Parents First' differentiator"],
            ["8", "Smart Data Import", "#1431", "~$0.02/photo", "High — Layer 1→2 accelerator"],
        ]
    )

    doc.add_page_break()

    # 3.1 Smart Data Import
    add_heading(doc, "3.1 Smart Data Import — Parent-Powered School Data (§6.67)", 2)
    doc.add_paragraph("GitHub Epic: #1431 | Sub-issues: #1432, #1433, #1434")
    doc.add_paragraph(
        "The Layer 1 bottleneck: if data doesn't flow in easily, Layer 2 has nothing to work with. "
        "School boards won't grant API access. Solution: bypass the school board entirely — "
        "empower parents to bring their own data in."
    )
    add_table(doc,
        ["Method", "How It Works", "AI Cost", "Friction"],
        [
            ["Photo Capture (#1432)", "Snap assignment/report card → AI extracts data", "~$0.02/photo", "Very low"],
            ["Email Forwarding (#1433)", "Forward school email → auto-parse assignments", "~$0.01/email", "Very low"],
            ["Calendar Import (#1434)", "Paste school ICS URL → auto-sync events", "$0", "Low — one-time"],
        ]
    )
    doc.add_paragraph(
        "This is a Layer 1 → Layer 2 accelerator: a parent who photographs a report card today "
        "gets a 'Help My Kid' button on that data tomorrow."
    )

    doc.add_page_break()

    # 4. Smart Daily Briefing
    add_heading(doc, "4. Smart Daily Briefing (§6.61)", 1)
    doc.add_paragraph("GitHub Epic: #1403 | Sub-issues: #1404, #1405, #1406")
    doc.add_paragraph(
        "A proactive daily summary that tells parents what matters today across all children. "
        "The #1 answer to 'why should I open ClassBridge?'"
    )
    doc.add_paragraph(
        "Example:\n"
        '"Good morning, Sarah. Here\'s what matters today."\n'
        "• Emma has a Math test tomorrow — she hasn't opened the study guide yet\n"
        "• Liam's Science project is due in 3 days — no progress tracked\n"
        "• 2 new announcements from Mrs. Chen"
    )
    doc.add_paragraph("AI Cost: $0.00 — pure SQL queries against existing tables.")
    add_heading(doc, "Components", 2)
    doc.add_paragraph("• Backend: GET /api/briefing/daily — aggregation endpoint")
    doc.add_paragraph("• Frontend: DailyBriefing card replacing Today's Focus header")
    doc.add_paragraph("• Email: Optional 7 AM morning digest via SendGrid")

    doc.add_page_break()

    # 5. Help My Kid
    add_heading(doc, "5. Help My Kid — One-Tap Study Actions (§6.62)", 1)
    doc.add_paragraph("GitHub Epic: #1407 | Sub-issues: #1408, #1409")
    doc.add_paragraph(
        "Parent sees upcoming test → taps 'Help Study' → ClassBridge generates practice material "
        "and sends it to the child's dashboard with a notification. One tap from worried to action."
    )
    add_heading(doc, "Source Material Linking", 2)
    doc.add_paragraph(
        "Generated materials maintain a lineage chain via self-referential FK on study_guides:\n"
        "• source_study_guide_id (FK → study_guides.id) — 'derived from' link\n"
        "• generated_by_user_id (FK → users.id) — who triggered it\n"
        "• generated_for_user_id (FK → users.id) — who it's for\n\n"
        "No new tables. Max 2 levels deep. Existing CRUD/permissions work unchanged."
    )

    doc.add_page_break()

    # 6. Global Search
    add_heading(doc, "6. Global Search + Smart Shortcuts (§6.17)", 1)
    doc.add_paragraph("GitHub Epic: #1410 | Sub-issues: #1411, #1412")
    add_heading(doc, "Hybrid Search Strategy", 2)
    doc.add_paragraph(
        "All search is SQL ILIKE — zero AI cost. AI only invoked when user explicitly "
        "clicks an action button on a result."
    )
    add_table(doc,
        ["User Query", "How It Works", "AI Cost"],
        [
            ["What's due for my kids", "SQL: tasks WHERE due_date >= today", "$0"],
            ["What's due for Emma", "SQL: same + filter by student", "$0"],
            ["Search a study guide", "ILIKE on study_guides.title", "$0"],
            ["Search class materials", "ILIKE on course_contents.title", "$0"],
            ["Search classes", "ILIKE on courses.name", "$0"],
            ["Upload class material", "Intent routing → upload page", "$0"],
            ["Create word problems", "Action button → AI generation", "~$0.02"],
        ]
    )
    add_heading(doc, "Smart Presets", 2)
    add_table(doc,
        ["Keyword", "Behavior"],
        [
            ['"due" / "overdue"', "Shows tasks/assignments due this week"],
            ["Child name", "Shows that child's courses, tasks, materials"],
            ['"upload"', "Shows 'Upload Material' action card"],
            ['"create"', "Shows creation options"],
        ]
    )

    doc.add_page_break()

    # 7. Weekly Progress Pulse
    add_heading(doc, "7. Weekly Progress Pulse (§6.63)", 1)
    doc.add_paragraph("GitHub Epic: #1413")
    doc.add_paragraph(
        "Weekly email digest sent Sunday evening summarizing the past week and previewing next.\n\n"
        "Example:\n"
        '"This week: Emma completed 4/6 assignments. Liam has 2 overdue. '
        "Here's what's coming next week.\"\n\n"
        "AI Cost: $0.00 — pure SQL + SendGrid."
    )

    # 8. Parent-Child Study Link
    add_heading(doc, "8. Parent-Child Study Link (§6.64)", 1)
    doc.add_paragraph("GitHub Epic: #1414")
    doc.add_paragraph(
        "Feedback loop:\n"
        "1. Parent generates quiz → child gets notification\n"
        "2. Child completes quiz → parent gets score + struggle areas\n"
        "3. Parent sees 'Study Help I've Sent' with completion status\n\n"
        "Creates emotional connection through the platform. AI Cost: $0."
    )

    doc.add_page_break()

    # 9. Dashboard Redesign
    add_heading(doc, "9. Dashboard Redesign — Persona-Based (§6.65)", 1)
    doc.add_paragraph("GitHub Epic: #1415 | Sub-issues: #1416-#1419")
    add_heading(doc, "Design Philosophy", 2)
    doc.add_paragraph(
        "• One-screen rule: Everything visible without scrolling (1080p)\n"
        "• 3-section max per dashboard\n"
        "• White space is a feature\n"
        "• Role-specific language\n"
        "• Action-first: Lead with what user can DO"
    )
    add_table(doc,
        ["Dashboard", "3 Sections", "Issue"],
        [
            ["Parent v5", "Daily Briefing + Child Snapshot + Quick Actions", "#1416"],
            ["Student v4", "Coming Up + Recent Study + Quick Actions", "#1417"],
            ["Teacher v2", "Student Alerts + My Classes + Quick Actions", "#1418"],
            ["Admin v2", "Platform Health + Recent Activity + Quick Actions", "#1419"],
        ]
    )

    doc.add_page_break()

    # 10. Responsible AI Tools
    add_heading(doc, "10. Responsible AI Parent Tools (§6.66)", 1)
    doc.add_paragraph("GitHub Epic: #1421 | Sub-issues: #1422-#1428")
    add_heading(doc, "Tool Suite", 2)
    add_table(doc,
        ["#", "Tool", "For Parent", "For Student", "AI Cost", "Responsible?"],
        [
            ["1", "Is My Kid Ready?", "Readiness score + gaps", "Must answer 5 questions", "~$0.02", "Yes — tests"],
            ["2", "Parent Briefing Notes", "Topic summary + home tips", "Never sees it", "~$0.01", "Yes — parent only"],
            ["3", "Practice Problems", "Extra practice assigned", "Must solve problems", "~$0.02", "Yes — work required"],
            ["4", "Weak Spot Report", "Trends over time", "Sees own progress", "$0.00", "Yes — real data"],
            ["5", "Conversation Starters", "Dinner table prompts", "N/A", "~$0.005", "Yes — discussion"],
        ]
    )

    add_heading(doc, "Revised Help Study Menu", 2)
    doc.add_paragraph(
        "Primary (responsible tools — student must do work):\n"
        "  1. Quick Assessment (5 questions)\n"
        "  2. Practice Problems\n"
        "  3. Parent Briefing (parent-only)\n\n"
        "Secondary (More Options):\n"
        "  4. Quiz (multiple choice)\n"
        "  5. Study Guide (summary)\n"
        "  6. Flashcards (memorization)"
    )

    add_heading(doc, "Flashcards Assessment", 2)
    doc.add_paragraph(
        "Flashcards bring value for vocabulary, definitions, and formulas (spaced repetition is proven). "
        "However, they're the most passive tool — a student can flip cards without retaining anything. "
        "Decision: Keep flashcards but de-emphasize in the UI. They're a 'More Options' tool, not a primary action."
    )

    doc.add_page_break()

    # 11. AI Cost Analysis
    add_heading(doc, "11. AI Cost Analysis", 1)
    add_table(doc,
        ["Approach", "Cost/Search", "Monthly (1,000 searches)"],
        [
            ["SQL ILIKE (all search)", "$0.00", "$0.00"],
            ["Keyword routing (intent)", "$0.00", "$0.00"],
            ["AI generation (on-demand)", "~$0.02", "~$20 (only explicit actions)"],
            ["If AI used for ALL searches", "~$0.01-0.05", "$50-500+ (NOT recommended)"],
        ]
    )
    doc.add_paragraph(
        "Bottom line: 6 of 7 original use cases need zero AI. "
        "Search is SQL + keyword routing. AI is only for content generation on explicit user action."
    )

    doc.add_page_break()

    # 12. Source Material Linking
    add_heading(doc, "12. Source Material Linking Design", 1)
    doc.add_paragraph(
        "Generated materials (quizzes, practice problems, etc.) link back to their source via "
        "a self-referential FK on study_guides. No new tables needed."
    )
    doc.add_paragraph(
        "Schema changes (3 nullable columns):\n"
        "  ALTER TABLE study_guides ADD COLUMN source_study_guide_id INTEGER REFERENCES study_guides(id);\n"
        "  ALTER TABLE study_guides ADD COLUMN generated_by_user_id INTEGER REFERENCES users(id);\n"
        "  ALTER TABLE study_guides ADD COLUMN generated_for_user_id INTEGER REFERENCES users(id);"
    )
    doc.add_paragraph(
        "Lineage chain: Original Document → Study Guide → Quiz → Word Problems\n"
        "Max depth: 2 levels. Soft-delete only on source materials."
    )

    doc.add_page_break()

    # 13. Parallel Dev Plan
    add_heading(doc, "13. Parallel Development Plan", 1)
    doc.add_paragraph("GitHub Issue: #1429")
    add_heading(doc, "Timeline (6-7 weeks, 2 developers)", 2)
    add_table(doc,
        ["Week", "Dev 1 (Frontend)", "Dev 2 (Backend)"],
        [
            ["1-2", "Dashboard Redesign (#1415)", "Search API + Briefing API + AI Tools (all 6 backends)"],
            ["3-4", "Briefing card + Help Study menu + Search component", "Integration + Study Link backend"],
            ["5-6", "Study Link UI + Weekly Pulse preferences + Polish", "Weekly Pulse backend + Tests + Help/Chatbot"],
        ]
    )
    add_heading(doc, "Maximum Parallelism", 2)
    doc.add_paragraph(
        "Week 1: 4 parallel streams possible\n"
        "  • Stream A: Parent Dashboard v5 (frontend)\n"
        "  • Stream C-BE: Search API (backend)\n"
        "  • Stream D-BE: All 6 AI tool backends (backend)\n"
        "  • Stream F-BE: Weekly digest service (backend)\n\n"
        "Risk mitigation: Ship Dashboard + Briefing + Search first (3 features = immediate WOW). "
        "Defer Study Link + Weekly Pulse if behind."
    )

    doc.add_page_break()

    # 14. GitHub Issues Created
    add_heading(doc, "14. GitHub Issues Created (This Session)", 1)
    add_table(doc,
        ["#", "Title", "Type"],
        [
            ["#1403", "Epic: Smart Daily Briefing (§6.61)", "Epic"],
            ["#1404", "Backend: daily briefing endpoint", "Sub"],
            ["#1405", "Frontend: briefing card", "Sub"],
            ["#1406", "Email: daily morning digest", "Sub"],
            ["#1407", "Epic: Help My Kid (§6.62)", "Epic"],
            ["#1408", "Backend: one-tap generation + source linking", "Sub"],
            ["#1409", "Frontend: Help Study buttons + modal", "Sub"],
            ["#1410", "Epic: Global Search + Shortcuts (§6.17)", "Epic"],
            ["#1411", "Backend: unified search endpoint", "Sub"],
            ["#1412", "Frontend: GlobalSearch component", "Sub"],
            ["#1413", "Epic: Weekly Progress Pulse (§6.63)", "Epic"],
            ["#1414", "Epic: Parent-Child Study Link (§6.64)", "Epic"],
            ["#1415", "Epic: Dashboard Redesign (§6.65)", "Epic"],
            ["#1416", "Parent Dashboard v5", "Sub"],
            ["#1417", "Student Dashboard v4", "Sub"],
            ["#1418", "Teacher Dashboard v2", "Sub"],
            ["#1419", "Admin Dashboard v2", "Sub"],
            ["#1420", "Help Page + Chatbot knowledge base", "Issue"],
            ["#1421", "Epic: Responsible AI Parent Tools (§6.66)", "Epic"],
            ["#1422", "Is My Kid Ready? assessment", "Sub"],
            ["#1423", "Parent Briefing Notes", "Sub"],
            ["#1424", "Practice Problem Sets", "Sub"],
            ["#1425", "Weak Spot Report", "Sub"],
            ["#1426", "Conversation Starters", "Sub"],
            ["#1427", "Frontend: revised Help Study menu", "Sub"],
            ["#1428", "Tests: all new tools", "Sub"],
            ["#1429", "Dev Plan: parallel development strategy", "Issue"],
            ["#1430", "Product Strategy: Two Layers", "Strategy"],
            ["#1431", "Epic: Smart Data Import (§6.67)", "Epic"],
            ["#1432", "Photo Capture: snap & import", "Sub"],
            ["#1433", "Email Forwarding: parse school emails", "Sub"],
            ["#1434", "Calendar Import: ICS feed sync", "Sub"],
        ]
    )

    doc.add_page_break()

    # 15. Files Updated
    add_heading(doc, "15. Files Updated", 1)
    add_table(doc,
        ["File", "Changes"],
        [
            ["requirements/features-part3.md", "Added §6.60.1 Two Layers strategy, §6.61-6.67 (all WOW features + Smart Data Import)"],
            ["requirements/roadmap.md", "Added Phase 2 WOW Features with priority table, Smart Data Import at priority 2"],
            ["REQUIREMENTS.md", "Updated document structure table to reference §6.61-6.67"],
        ]
    )

    # Save
    output_path = "docs/ClassBridge_WOW_Features_Session.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")


if __name__ == "__main__":
    main()
