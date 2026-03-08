"""Generate Word document from ClassBridge Feature Catalog PRD markdown using python-docx."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

docs_dir = Path(__file__).parent
md_path = docs_dir / "classbridge-feature-catalog-prd.md"
docx_path = docs_dir / "ClassBridge_Feature_Catalog.docx"

md_content = md_path.read_text(encoding="utf-8")

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Style setup
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

# Heading styles - ClassBridge brand colors
brand_colors = {
    1: RGBColor(0x1B, 0x4F, 0x72),  # Dark blue
    2: RGBColor(0x2E, 0x74, 0x91),  # Medium blue
    3: RGBColor(0x49, 0xB8, 0xC0),  # Teal (ClassBridge primary)
    4: RGBColor(0x5A, 0x9E, 0x8F),  # Sage
}

for level in range(1, 5):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Calibri'
    heading_style.font.color.rgb = brand_colors[level]
    if level == 1:
        heading_style.font.size = Pt(20)
    elif level == 2:
        heading_style.font.size = Pt(16)
    elif level == 3:
        heading_style.font.size = Pt(13)
    else:
        heading_style.font.size = Pt(11)


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_from_lines(doc, header_line, rows):
    """Parse markdown table and add to document."""
    headers = [cell.strip() for cell in header_line.strip('|').split('|')]
    headers = [h for h in headers if h]

    if not headers:
        return

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row with brand color background
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header.strip('*').strip()
        set_cell_shading(cell, '1B4F72')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = 'Calibri'

    # Data rows with alternating colors
    for idx, row_text in enumerate(rows):
        cells = [cell.strip() for cell in row_text.strip('|').split('|')]
        cells = [c for c in cells if c or cells.index(c) > 0]
        if not cells:
            continue
        row = table.add_row()
        for i, cell_text in enumerate(cells):
            if i < len(headers):
                cell = row.cells[i]
                clean = cell_text.strip()
                clean = re.sub(r'\*\*(.*?)\*\*', r'\1', clean)
                clean = re.sub(r'\*(.*?)\*', r'\1', clean)
                clean = re.sub(r'`(.*?)`', r'\1', clean)
                clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', clean)
                clean = re.sub(r'~~(.*?)~~', r'\1', clean)
                cell.text = clean
                # Alternating row colors
                if idx % 2 == 1:
                    set_cell_shading(cell, 'F0F7FA')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Calibri'


def add_cover_page(doc):
    """Add a professional cover page."""
    # Spacer
    for _ in range(4):
        doc.add_paragraph('')

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ClassBridge')
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x49, 0xB8, 0xC0)
    run.font.name = 'Calibri'
    run.bold = True

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Complete Feature & Design Catalog')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    run.font.name = 'Calibri'

    doc.add_paragraph('')

    # Divider line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_' * 50)
    run.font.color.rgb = RGBColor(0x49, 0xB8, 0xC0)

    doc.add_paragraph('')

    # Metadata
    metadata = [
        ('Version', '1.0'),
        ('Date', '2026-03-08'),
        ('Author', 'Sarah (Product Owner)'),
        ('Source Data', '1,007 GitHub Issues + 60+ Requirement Sections'),
        ('Quality Score', '95/100'),
    ]

    for label, value in metadata:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}: ')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.name = 'Calibri'
        run.bold = True
        run = p.add_run(value)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run.font.name = 'Calibri'

    doc.add_paragraph('')
    doc.add_paragraph('')

    # Tagline
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('AI-Powered Education Platform')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x49, 0xB8, 0xC0)
    run.font.name = 'Calibri'
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Connecting Parents, Students, Teachers & Administrators')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = 'Calibri'
    run.italic = True

    # Page break
    doc.add_page_break()


def process_markdown(doc, content):
    """Process markdown content and add to Word document."""
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    # Skip the first H1 title line since we have a cover page
    skip_first_h1 = True

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_lines)
                if code_text.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Horizontal rules
        if line.strip() == '---':
            i += 1
            continue

        # Headings
        if line.startswith('#'):
            level = 0
            for ch in line:
                if ch == '#':
                    level += 1
                else:
                    break
            text = line[level:].strip()
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

            if level == 1 and skip_first_h1:
                skip_first_h1 = False
                i += 1
                continue

            if level <= 4:
                doc.add_heading(text, level=min(level, 4))
            else:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(11)
            i += 1
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1

            if len(table_lines) >= 2:
                header = table_lines[0]
                data_rows = [l for l in table_lines[2:] if not re.match(r'^[\s|:-]+$', l)]
                add_table_from_lines(doc, header, data_rows)
                doc.add_paragraph('')
            continue

        # Empty lines
        if line.strip() == '':
            i += 1
            continue

        # Checkbox items
        if re.match(r'^\s*-\s*\[[ x]\]', line.strip()):
            checked = '[x]' in line
            text = re.sub(r'^-\s*\[[ x]\]\s*', '', line.strip())
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            prefix = "[DONE] " if checked else "[ ] "
            doc.add_paragraph(prefix + text, style='List Bullet')
            i += 1
            continue

        # Bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()
            if text.startswith('- '):
                text = text[2:]
            else:
                text = text[2:]
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

            indent_level = (len(line) - len(line.lstrip())) // 2
            if indent_level > 0:
                try:
                    p = doc.add_paragraph(text, style='List Bullet 2')
                except KeyError:
                    p = doc.add_paragraph(text, style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.5)
            else:
                p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Numbered lists
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Blockquotes
        if line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # Regular paragraph
        text = line.strip()
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'`(.*?)`', r'\1', text)
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        if text:
            doc.add_paragraph(text)
        i += 1


# Build document
add_cover_page(doc)
process_markdown(doc, md_content)

# Add footer with page numbers
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run('ClassBridge Feature Catalog  |  ')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = 'Calibri'

    # Page number field
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    run2 = p.add_run()
    run2._r.append(fld_char_begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    run3 = p.add_run()
    run3._r.append(instr)

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run4 = p.add_run()
    run4._r.append(fld_char_end)

# Save
doc.save(str(docx_path))
print(f"Word document generated: {docx_path}")
print(f"File size: {docx_path.stat().st_size / 1024:.1f} KB")
