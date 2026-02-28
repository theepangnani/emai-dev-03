"""
ClassBridge UI/UX Assessment Report Generator
Generates a comprehensive Word document with HCD-driven UI analysis.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1B, 0x1E, 0x2B)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1B, 0x1E, 0x2B)
    hs.font.name = 'Calibri'


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_table_row(table, cells_data, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        if bold:
            run.bold = True
    return row


def styled_table(cols, headers):
    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    return table


# ══════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ClassBridge')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x49, 0xB8, 0xC0)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('UI/UX Assessment Report')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1B, 0x1E, 0x2B)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Human-Centered Design Evaluation\nAcross Parent, Student & Teacher Roles')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x5B, 0x62, 0x74)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Date: {datetime.date.today().strftime("%B %d, %Y")}')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x5B, 0x62, 0x74)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Version 1.0 | Confidential')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x5B, 0x62, 0x74)

doc.add_page_break()


# ══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Assessment Methodology & Scope',
    '3. Design System Audit',
    '   3.1 Color Palette & Theming',
    '   3.2 Typography',
    '   3.3 Spacing, Shadows & Elevation',
    '   3.4 Component Patterns',
    '   3.5 Design System Inconsistencies',
    '4. Cross-Role Consistency Analysis',
    '   4.1 Shared Layout (DashboardLayout)',
    '   4.2 Navigation Structure',
    '   4.3 Common Interaction Patterns',
    '5. Parent Role Assessment',
    '   5.1 Dashboard Experience',
    '   5.2 Child Management (My Kids)',
    '   5.3 Task & Material Workflows',
    '   5.4 Communication Flows',
    '   5.5 Parent Pain Points & Recommendations',
    '6. Student Role Assessment',
    '   6.1 Dashboard Experience',
    '   6.2 Study Tools Workflow',
    '   6.3 Course & Material Navigation',
    '   6.4 Student Pain Points & Recommendations',
    '7. Teacher Role Assessment',
    '   7.1 Dashboard Experience',
    '   7.2 Class Management',
    '   7.3 Communication & Announcements',
    '   7.4 Teacher Pain Points & Recommendations',
    '8. Accessibility Audit (WCAG 2.1)',
    '9. Responsive Design & Mobile Experience',
    '10. User Journey Maps',
    '11. Industry Benchmarks & Inspiration',
    '12. Risk Register',
    '13. Prioritized Recommendations',
    '14. Appendix: File Inventory',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'This report presents a comprehensive Human-Centered Design (HCD) evaluation of the ClassBridge '
    'web application across all three primary user roles: Parent, Student, and Teacher. The assessment '
    'examines UI consistency, interaction patterns, accessibility compliance, responsive behavior, '
    'design system coherence, and end-to-end user journeys.'
)
doc.add_paragraph(
    'ClassBridge is an AI-powered education management platform that connects parents, students, '
    'and teachers through Google Classroom integration, AI study tools (study guides, quizzes, flashcards), '
    'parent-teacher messaging, and teacher email/announcement monitoring. The platform serves a diverse '
    'user base with varying technical literacy levels, making usability and clarity paramount.'
)

doc.add_heading('Key Findings Summary', level=2)

doc.add_heading('Strengths', level=3)
for s in [
    'Well-structured role-based architecture with distinct, purpose-built dashboards for each role',
    'Comprehensive design token system with 3 themes (Light, Dark, Focus) using CSS variables',
    'Strong accessibility foundation: skip-to-content links, ARIA labels, keyboard shortcuts (Ctrl+K, ?)',
    'Effective loading states with skeleton shimmer animations and reduced-motion support',
    'Smart child-selector pattern for parents with session persistence and keyboard navigation',
    'Thoughtful empty states with contextual CTAs that guide users toward next actions',
    'Toast notification system with auto-dismiss and aria-live regions',
    'Onboarding tour system with role-specific step sequences',
]:
    add_bullet(s)

doc.add_heading('Areas of Concern', level=3)
for c in [
    'Design system inconsistencies: 20+ hard-coded color values bypass CSS variables',
    'Typography scale fragmentation: 10+ distinct font sizes without a systematic scale',
    'Modal overflow: 10+ modal types with subtly different padding, sizing, and animation values',
    'Mobile experience gaps: touch targets, horizontal scroll discoverability, modal overflow',
    'Parent dashboard cognitive load: too much information density on first load',
    'Cross-role navigation inconsistency: different sidebar items use different icon/label conventions',
    'Missing features in student/teacher workflows that break expected user journeys',
    'Z-index sprawl: values range from -1 to 10001 without documented layering system',
]:
    add_bullet(c)

doc.add_heading('Overall Assessment Scores', level=2)
t = styled_table(3, ['Dimension', 'Score (1-10)', 'Notes'])
for row in [
    ('Visual Consistency', '6.5/10', 'Strong token system undermined by hard-coded overrides'),
    ('Usability (Parent)', '7.5/10', 'Rich features but high cognitive load; child management well done'),
    ('Usability (Student)', '7/10', 'Clean dashboard; study tool flow is smooth; missing grade tracking'),
    ('Usability (Teacher)', '6.5/10', 'Functional but utilitarian; announcement workflow needs preview'),
    ('Accessibility', '7/10', 'Good ARIA foundation; gaps in focus traps, form labels, color-only cues'),
    ('Responsive/Mobile', '6/10', 'Breakpoints exist but inconsistent; touch targets need work'),
    ('Information Architecture', '7.5/10', 'Clear role separation; navigation mostly intuitive'),
    ('Performance Perception', '8/10', 'Skeleton loading, lazy routes, debounced search all implemented'),
]:
    add_table_row(t, row)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 2. METHODOLOGY
# ══════════════════════════════════════════════════════════
doc.add_heading('2. Assessment Methodology & Scope', level=1)

doc.add_heading('Methodology', level=2)
doc.add_paragraph(
    'This assessment was conducted through a systematic code-level review of the entire frontend '
    'codebase (40+ CSS files, 30+ page components, 25+ reusable components). The evaluation applies '
    'Human-Centered Design (HCD) principles across five dimensions:'
)
for b, d in [
    ('Heuristic Evaluation: ', "Nielsen's 10 usability heuristics applied to each role's primary workflows"),
    ('Design System Audit: ', 'Complete inventory of colors, typography, spacing, shadows, and component patterns'),
    ('Accessibility Review: ', 'WCAG 2.1 AA compliance check across keyboard navigation, screen readers, color contrast'),
    ('Responsive Audit: ', 'Breakpoint analysis at 480px, 640px, 768px, 900px, and 1024px'),
    ('Journey Mapping: ', 'End-to-end user flow analysis for critical tasks per role'),
]:
    add_bullet(d, bold_prefix=b)

doc.add_heading('Scope', level=2)
for item in [
    'All pages accessible to Parent, Student, and Teacher roles',
    'Shared components (DashboardLayout, modals, notifications, search, calendar)',
    'Role-specific dashboard implementations and their sub-components',
    'Authentication flow (login, register, forgot/reset password, email verification)',
    'CSS design system (index.css, Dashboard.css, Auth.css, all page-specific CSS)',
    'Navigation patterns, routing architecture, and state persistence',
]:
    add_bullet(item)

doc.add_heading('Out of Scope', level=2)
for item in [
    'Backend API performance and response times',
    'Mobile native app (ClassBridgeMobile/) - separate codebase',
    'Admin role dashboard (minimal UI, not user-facing)',
]:
    add_bullet(item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 3. DESIGN SYSTEM AUDIT
# ══════════════════════════════════════════════════════════
doc.add_heading('3. Design System Audit', level=1)
doc.add_paragraph(
    'ClassBridge implements a well-structured design token system through CSS custom properties in index.css, '
    'supporting three themes: Light, Dark, and Focus. However, the system suffers from inconsistent adoption '
    'across page-specific CSS files.'
)

doc.add_heading('3.1 Color Palette & Theming', level=2)
doc.add_heading('Light Theme (Default)', level=3)
doc.add_paragraph('The light theme uses a clean, professional palette anchored by teal (#49B8C0) as the primary accent:')
t = styled_table(4, ['Token', 'Value', 'Usage', 'Contrast'])
for row in [
    ('--color-accent', '#49B8C0', 'Primary CTAs, links, active states', 'AA borderline on white'),
    ('--color-accent-strong', '#3A9AA1', 'Hover states, emphasis', 'AA compliant'),
    ('--color-ink', '#1B1E2B', 'Primary text', 'AAA compliant'),
    ('--color-ink-muted', '#5B6274', 'Secondary text, labels', 'AA compliant'),
    ('--color-surface', '#FFFFFF', 'Cards, panels, modals', 'N/A'),
    ('--color-surface-bg', '#EEF1F5', 'Page background', 'N/A'),
    ('--color-danger', '#D64545', 'Delete, error, destructive', 'AA on white'),
    ('--color-warning', '#F3B04C', 'Warnings, pending states', 'Needs dark text'),
    ('--color-success', '#2E7D32', 'Completed, positive states', 'AA on white'),
    ('--priority-high', '#EF5350', 'Overdue/urgent items', 'AA on white'),
    ('--priority-medium', '#FF9800', 'Due today items', 'Needs dark text'),
    ('--priority-low', '#66BB6A', 'Upcoming/low priority', 'Needs dark text'),
]:
    add_table_row(t, row)

doc.add_paragraph()
doc.add_heading('Dark Theme', level=3)
doc.add_paragraph(
    'The dark theme shifts the accent to purple (#A78BFA) and uses dark charcoal surfaces (#1E1E1E). '
    'Color adjustments maintain readability, though some page-specific hard-coded colors do not adapt '
    'to dark mode, creating visual breaks.'
)

doc.add_heading('Focus Theme', level=3)
doc.add_paragraph(
    'A unique differentiator. The Focus theme uses warm, muted tones (beige/cream backgrounds, muted teal accents) '
    'designed for distraction-free studying. This is an excellent HCD feature for students with attention difficulties.'
)

doc.add_heading('Issue: Hard-Coded Color Bypass', level=3)
doc.add_paragraph('Multiple CSS files use hard-coded hex/rgb values instead of CSS variables, causing theme-switching failures:')
for item in [
    'CourseMaterialDetailPage.css: Edit button (#b8c4ce, #f0f4f8, #3a5068) and Archive button (#d4a574, #fdf6ee, #8b6914)',
    'StudentDashboard.css: Resolved banners use hard-coded #e8f5e9, #2e7d32, #c8e6c9',
    'Various files: rgba(73,184,192,*) used directly instead of var(--color-accent) with opacity',
    'Print view: Entirely hard-coded (#fff, #1a1a2e, #666, #eee) with no theme awareness',
]:
    add_bullet(item)

doc.add_heading('3.2 Typography', level=2)
doc.add_paragraph(
    'ClassBridge uses Space Grotesk (display/headings) and Source Sans 3 (body). This is a solid pairing. '
    'However, the type scale is fragmented with 10+ distinct sizes without systematic naming.'
)
t = styled_table(4, ['Size', 'Font', 'Weight', 'Usage'])
for row in [
    ('28px', 'Space Grotesk', '700', 'Auth page titles'),
    ('24px', 'Space Grotesk', '700', 'Page titles'),
    ('18px', 'Space Grotesk', '700', 'Section headers'),
    ('17px', 'Space Grotesk', '700', 'Panel headers (student) -- inconsistent with 18px'),
    ('16px', 'Source Sans 3', '500-600', 'Body large, card titles'),
    ('14px', 'Source Sans 3', '400-500', 'Body regular, form labels'),
    ('13px', 'Source Sans 3', '400-500', 'Body small, metadata'),
    ('12px', 'Source Sans 3', '500-600', 'Labels, captions'),
    ('11px', 'Source Sans 3', '600-700', 'Badges, tags'),
    ('10px', 'Source Sans 3', '600', 'Micro badges'),
]:
    add_table_row(t, row)
doc.add_paragraph()
doc.add_paragraph('Recommendation: Consolidate to 7 sizes: 10, 12, 14, 16, 20, 24, 28px with CSS variable names.')

doc.add_heading('3.3 Spacing, Shadows & Elevation', level=2)
doc.add_paragraph(
    'Spacing uses a loosely consistent 4px-based system (4, 8, 12, 16, 20, 24, 32, 40, 48px). '
    'However, values like 6px, 10px, 14px, 18px, 22px appear in various files, breaking the rhythm.'
)
doc.add_paragraph('Shadow tokens are defined but many components use ad-hoc values:')
t = styled_table(3, ['Token', 'Value', 'Usage'])
for row in [
    ('--shadow-soft', '0 8px 24px rgba(19,27,46,0.08)', 'Cards, panels (resting)'),
    ('--shadow-lift', '0 12px 30px rgba(19,27,46,0.12)', 'Cards, panels (hover)'),
    ('(ad-hoc)', '0 4px 20px rgba(0,0,0,0.12)', 'Toast notifications'),
    ('(ad-hoc)', '0 4px 16px rgba(0,0,0,0.15)', 'Modals'),
    ('(ad-hoc)', '0 8px 30px rgba(0,0,0,0.18)', 'Tour tooltip'),
]:
    add_table_row(t, row)
doc.add_paragraph()
doc.add_paragraph('Recommendation: Expand to --shadow-xs through --shadow-xl and replace all ad-hoc values.')

doc.add_heading('3.4 Component Patterns', level=2)
doc.add_heading('Buttons', level=3)
for b, d in [
    ('Primary (.generate-btn): ', 'Teal background, white text, 8px radius, hover lifts with shadow'),
    ('Secondary (.cancel-btn): ', 'Transparent/light background, border, muted text'),
    ('Danger (.danger-btn): ', 'Red background, white text, for destructive actions'),
    ('Ghost (.sd-text-btn): ', 'No background, accent-colored text, subtle hover'),
]:
    add_bullet(d, bold_prefix=b)

doc.add_heading('Modals', level=3)
doc.add_paragraph(
    'Consistent modal pattern (overlay + modal + form + actions) from Dashboard.css. All 10+ modal types '
    'follow this pattern. However, padding and max-width vary (500/600/800px) without clear sizing rules.'
)

doc.add_heading('Cards', level=3)
doc.add_paragraph(
    'Card components use --radius-lg (18px) with --shadow-soft and 1px borders. Hover states consistently '
    'add translateY(-2px) and --shadow-lift. This is well-implemented and consistent.'
)

doc.add_heading('3.5 Design System Inconsistencies', level=2)
t = styled_table(4, ['#', 'Finding', 'Severity', 'Category'])
for i, row in enumerate([
    ('Hard-coded colors bypass theme tokens in 5+ CSS files', 'High', 'Visual'),
    ('10+ font sizes without systematic scale naming', 'Medium', 'Typography'),
    ('Ad-hoc shadow values alongside defined tokens', 'Medium', 'Elevation'),
    ('Border radius mixing (6/10/20px alongside tokens)', 'Low', 'Geometry'),
    ('Spacing rhythm breaks (6/10/14px in 4px system)', 'Low', 'Layout'),
    ('Z-index spans -1 to 10001 without documented layers', 'Medium', 'Architecture'),
    ('Media query breakpoints inconsistent (480/640/768/900/1024)', 'High', 'Responsive'),
    ('EditMaterialModal.css (22 lines) separate from Dashboard.css', 'Low', 'Organization'),
    ('CHILD_COLORS array duplicated in MyKidsPage.tsx and useParentDashboard.ts', 'Low', 'Code'),
    ('Modal max-width varies (500/600/800px) without rules', 'Medium', 'Layout'),
], 1):
    add_table_row(t, [str(i)] + list(row))

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 4. CROSS-ROLE CONSISTENCY
# ══════════════════════════════════════════════════════════
doc.add_heading('4. Cross-Role Consistency Analysis', level=1)

doc.add_heading('4.1 Shared Layout (DashboardLayout)', level=2)
doc.add_paragraph(
    'All authenticated pages share DashboardLayout.tsx, providing a consistent shell: sticky header '
    'with backdrop blur, persistent sidebar (desktop), hamburger menu (mobile), footer, email '
    'verification banner. This is the strongest consistency anchor in the application.'
)
doc.add_paragraph('Shared header elements:')
for item in [
    'Logo + role-adaptive title ("Parent\'s Dashboard", etc.)',
    'Global search (Ctrl+K / Cmd+K) with grouped results',
    'Theme toggle (Light/Dark/Focus cycle)',
    'Notification bell with unread count badge and dropdown',
    'User chip with name, role badge, and multi-role switcher',
    'Sign out button',
]:
    add_bullet(item)

doc.add_heading('4.2 Navigation Structure', level=2)
t = styled_table(4, ['Sidebar Item', 'Parent', 'Student', 'Teacher'])
for row in [
    ('Home', 'Yes', 'Yes', 'Yes'),
    ('My Kids', 'Yes', '--', '--'),
    ('Classes', '--', 'Yes', 'Yes'),
    ('Materials', '--', 'Yes', 'Yes'),
    ('Quiz History', '--', 'Yes', '--'),
    ('Tasks', 'Yes', 'Yes', 'Yes'),
    ('Messages', 'Yes', 'Yes', 'Yes'),
    ('Teacher Comms', '--', '--', 'Yes'),
    ('Help', 'Yes', 'Yes', 'Yes'),
]:
    add_table_row(t, row)

doc.add_paragraph()
doc.add_heading('Navigation Consistency Issues', level=3)
for issue in [
    'Parents access Classes/Materials through child context (My Kids > child > classes), while Students/Teachers access them directly. This creates different mental models.',
    'Parent sidebar lacks direct "Classes" and "Materials" links, forcing extra navigation through My Kids.',
    'Student "Quiz History" is a standalone nav item that may feel orphaned. Consider integrating into Materials.',
    'Quick action paradigms differ: Parent uses FAB, Student uses 6 cards, Teacher uses 7 cards.',
]:
    add_bullet(issue)

doc.add_heading('4.3 Common Interaction Patterns', level=2)
doc.add_heading('Consistent Patterns (Positive)', level=3)
for item in [
    'Click card/row to navigate to detail page (courses, materials, tasks)',
    'Modal for creation (courses, tasks, study materials, invites)',
    'Breadcrumb navigation on detail pages (PageNav component)',
    'Search with debounce and URL sync on list pages',
    'Collapsible sections with chevron animation and aria-expanded',
    'Confirm dialog for destructive actions (ConfirmModal with promise pattern)',
]:
    add_bullet(item)

doc.add_heading('Inconsistent Patterns', level=3)
for item in [
    'Empty states: Some use emoji + title + CTA; others show text only. Should standardize.',
    'Error display: Mix of FAQErrorHint, generic error text, and toast-only patterns.',
    'Loading: Most pages use skeletons, but some show spinner-only states.',
    'List pagination: Mix of infinite scroll, "show more" buttons, and load-on-scroll.',
]:
    add_bullet(item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 5. PARENT ROLE ASSESSMENT
# ══════════════════════════════════════════════════════════
doc.add_heading('5. Parent Role Assessment', level=1)
doc.add_paragraph(
    'The Parent role is the most feature-rich and complex interface. Parents are typically non-technical '
    'guardians who need quick visibility into their children\'s academic status, upcoming deadlines, '
    'and teacher communications. The design must balance comprehensive information with cognitive simplicity.'
)

doc.add_heading('5.1 Dashboard Experience', level=2)
doc.add_heading("Today's Focus Header", level=3)
doc.add_paragraph(
    'The personalized header is an excellent HCD feature. It adapts based on the parent\'s situation: '
    '"Alex has 3 overdue tasks" (red/urgent) vs "All caught up!" (green/positive). This provides '
    'immediate emotional context without requiring any interaction.'
)
doc.add_paragraph('Strengths:')
for s in [
    'Dynamic, context-aware headlines reduce cognitive load',
    'Color-coded urgency tags (overdue=red, today=orange, upcoming=blue) use established conventions',
    'Per-child overdue breakdown for multi-child parents',
    'Collapsible to thin summary bar for frequent users',
]:
    add_bullet(s)
doc.add_paragraph('Concerns:')
for c in [
    'Information density on first load may overwhelm new parents',
    'Inspirational quotes compete with urgent information for attention',
    'No clear visual hierarchy between overdue and due-today beyond color',
    'Collapsed state may lose important context for infrequent users',
]:
    add_bullet(c)

doc.add_heading('Child Selector Tabs', level=3)
doc.add_paragraph(
    'Horizontal pill-style tabs with color-coded dots, grade badges, overdue badges, and ARIA tab '
    'pattern with arrow key navigation.'
)
doc.add_paragraph('Issues:')
for i in [
    'On mobile, horizontal tab scroll not discoverable (no scroll indicators)',
    'Deselecting a child to reach "All Children" view is not intuitive',
    'No explicit "All" tab as first option',
    'Small "+" add child button may be missed',
]:
    add_bullet(i)

doc.add_heading('5.2 Child Management (My Kids Page)', level=2)
doc.add_paragraph(
    'The most complex single page at 1,145 lines. Handles child profiles, course management, '
    'material assignment, teacher linking, and password resets.'
)
doc.add_paragraph('Concerns:')
for c in [
    'Page tries to do too much (6+ modal types). Consider splitting into sub-pages with tabs.',
    'Google discovery flow requires manual course-to-child assignment, which confuses multi-child parents.',
    '"Unassigned Courses/Materials" concept requires understanding the internal data model.',
    'Child edit form expands to 10+ fields (address, postal, etc.) creating form fatigue.',
    'No drag-and-drop for material reorganization between courses.',
]:
    add_bullet(c)

doc.add_heading('5.3 Task & Material Workflows', level=2)
doc.add_paragraph('Workflow Gaps:')
for g in [
    'No batch task operations (mark all complete, bulk delete)',
    'No recurring task templates (e.g., "read 30 min daily")',
    'Task vs Assignment distinction unclear to users',
    'Study material generation progress invisible (background only)',
    'No way to organize study materials into folders/collections',
    'No sharing of study materials between children',
]:
    add_bullet(g)

doc.add_heading('5.4 Communication Flows', level=2)
doc.add_paragraph('Communication Gaps:')
for g in [
    'No message search functionality',
    'No message attachments (file/image sharing)',
    'No read receipts or delivery confirmation',
    'No message templates for common communications',
    'No group messaging (e.g., message all teachers at once)',
]:
    add_bullet(g)

doc.add_heading('5.5 Parent Pain Points & Recommendations', level=2)
t = styled_table(4, ['#', 'Pain Point', 'Impact', 'Recommendation'])
for row in [
    ('1', 'Dashboard information overload', 'High', 'Progressive disclosure: collapsed sections by default for new users'),
    ('2', 'My Kids page too complex', 'High', 'Split into sub-pages: Profile, Classes, Materials, Tasks via tabs'),
    ('3', '"Unassigned" concept confusing', 'Medium', 'Guided onboarding for course-child assignment; auto-assign when possible'),
    ('4', 'No "All" tab in child selector', 'Medium', 'Add explicit "All Children" tab as first item'),
    ('5', 'Mobile tab scroll not discoverable', 'Medium', 'Add fade/gradient at scroll edges; swipe hint on first visit'),
    ('6', 'Task vs Assignment unclear', 'High', 'Rename to "To-Do" vs "Assignment"; add visual differentiation'),
    ('7', 'No batch task operations', 'Medium', 'Multi-select with bulk actions: complete, delete, change priority'),
    ('8', 'Study gen progress invisible', 'Low', 'Add progress indicator or toast with progress bar'),
]:
    add_table_row(t, row)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 6. STUDENT ROLE ASSESSMENT
# ══════════════════════════════════════════════════════════
doc.add_heading('6. Student Role Assessment', level=1)
doc.add_paragraph(
    'The Student role targets learners who need quick access to study materials, upcoming assignments, '
    'and AI-powered study tools. Students are typically younger (K-12) and expect modern, app-like '
    'interfaces with minimal friction.'
)

doc.add_heading('6.1 Dashboard Experience', level=2)
doc.add_paragraph(
    'Time-aware greeting with urgency pills and stat chips. The study streak counter is a positive '
    'gamification element. Six quick action cards provide rapid access to primary tasks.'
)
doc.add_paragraph('Concerns:')
for c in [
    '"New Course" and "Invite Teacher" may not be frequently used; consider moving to settings',
    'Six action cards cluttered on mobile (stacking to single column)',
    'No personalization based on activity patterns',
    '"Sync Classes" and "Google Connection" are related but presented as separate cards',
]:
    add_bullet(c)

doc.add_heading('6.2 Study Tools Workflow', level=2)
doc.add_paragraph(
    'The core value proposition. Text paste or file upload generates AI study guides, quizzes, '
    'and flashcards. Study guide viewer renders markdown with syntax highlighting. Quiz provides '
    'instant feedback and history. Flashcards have flip animation and shuffle mode.'
)
doc.add_paragraph('Gaps:')
for g in [
    'No spaced repetition algorithm for flashcards (proven to improve retention)',
    'No study session timer or Pomodoro integration',
    'No collaborative study groups or shared materials',
    'No integration with grades to suggest what to study',
    'No "smart review" focusing on missed quiz questions',
    'No audio/text-to-speech for study guides',
    'No mobile-optimized study mode (larger text, simplified nav)',
]:
    add_bullet(g)

doc.add_heading('6.3 Course & Material Navigation', level=2)
doc.add_paragraph('Gaps:')
for g in [
    'No grade tracking or report card view',
    'No assignment submission upload (only viewing)',
    'No "recently accessed" or "favorites" for quick access',
    'No course syllabus or schedule display',
]:
    add_bullet(g)

doc.add_heading('6.4 Student Pain Points & Recommendations', level=2)
t = styled_table(4, ['#', 'Pain Point', 'Impact', 'Recommendation'])
for row in [
    ('1', 'No grade visibility', 'High', 'Add grade summary card; per-course breakdown page'),
    ('2', 'Quick actions not personalized', 'Medium', 'Context-aware actions based on deadlines and study patterns'),
    ('3', 'No spaced repetition for flashcards', 'Medium', 'Implement SM-2 algorithm; track card confidence'),
    ('4', 'No recently accessed materials', 'Medium', 'Add "Continue Studying" section with last 3 accessed items'),
    ('5', 'Study streak fragile', 'Low', 'Streak freeze option (1 per week); celebrate milestones'),
    ('6', 'Missing assignment submission', 'High', 'File upload for submissions with teacher notification'),
    ('7', 'No favorites/bookmarks', 'Medium', 'Star/bookmark on materials; show on dashboard'),
]:
    add_table_row(t, row)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 7. TEACHER ROLE ASSESSMENT
# ══════════════════════════════════════════════════════════
doc.add_heading('7. Teacher Role Assessment', level=1)
doc.add_paragraph(
    'The Teacher role focuses on class management, parent communication, and content distribution. '
    'The current implementation is functional but feels more utilitarian than the Parent/Student experiences.'
)

doc.add_heading('7.1 Dashboard Experience', level=2)
doc.add_paragraph(
    'Custom header shows unread messages, upcoming deadlines, and student count. Seven dashboard cards '
    'in a 2x4 grid with emoji icons. Two-column layout below shows recent messages and upcoming deadlines.'
)
doc.add_paragraph('Concerns:')
for c in [
    'Emoji icons lack the polish of SVG icons used in student dashboard',
    'Grid cards show static text ("View", "Send") without counts or urgency',
    'No quick stats on cards (e.g., "3 unread messages")',
    '2x4 grid stacks to single column on mobile, creating excessive scroll',
]:
    add_bullet(c)

doc.add_heading('7.2 Class Management', level=2)
doc.add_paragraph('Gaps:')
for g in [
    'No bulk student import (CSV upload)',
    'No class roster management from dashboard',
    'No attendance tracking',
    'No grade book or assessment management',
    'No assignment creation from dashboard',
    'No class archive/end-of-term workflow',
]:
    add_bullet(g)

doc.add_heading('7.3 Communication & Announcements', level=2)
doc.add_paragraph('Announcement workflow gaps:')
for g in [
    'No preview before sending',
    'No rich text formatting (bold, lists, links)',
    'No attachment support',
    'No scheduling for future send',
    'No draft saving capability',
    'No template library for common announcements',
]:
    add_bullet(g)
doc.add_paragraph(
    'Positive note: The invite resend cooldown (1 hour) with "Wait Xm" timer is a thoughtful UX detail '
    'that prevents spam while allowing retries.'
)

doc.add_heading('7.4 Teacher Pain Points & Recommendations', level=2)
t = styled_table(4, ['#', 'Pain Point', 'Impact', 'Recommendation'])
for row in [
    ('1', 'Dashboard cards lack dynamic data', 'High', 'Show counts: "3 unread", "5 due"; add mini-badges'),
    ('2', 'Emoji icons feel unprofessional', 'Medium', 'Replace with consistent SVG icon set'),
    ('3', 'Announcement has no preview', 'High', 'Add preview step; show formatted message + recipient count'),
    ('4', 'No assignment creation from dashboard', 'Medium', 'Add "Create Assignment" to quick actions'),
    ('5', 'No grade book', 'High', 'Basic grading: submissions, rubric, grade recording'),
    ('6', 'No draft announcements', 'Medium', 'Auto-save drafts; show drafts list'),
    ('7', 'Communication channels fragmented', 'Medium', 'Unify Messages + Announcements + Comms into single Inbox'),
]:
    add_table_row(t, row)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 8. ACCESSIBILITY AUDIT
# ══════════════════════════════════════════════════════════
doc.add_heading('8. Accessibility Audit (WCAG 2.1 AA)', level=1)

doc.add_heading('Implemented (Strengths)', level=2)
for b, d in [
    ('Skip-to-content link: ', 'Hidden link visible on focus, jumps to main content'),
    ('Keyboard shortcuts: ', 'Ctrl+K for search, ? for legend, arrow keys for tabs'),
    ('ARIA roles: ', 'tablist/tab, alertdialog, modal on overlays'),
    ('aria-live regions: ', 'Toasts use aria-live="polite"'),
    ('Focus-visible: ', 'Global 2px solid accent outline'),
    ('Reduced motion: ', '@media (prefers-reduced-motion: reduce) disables animations'),
    ('Skeleton aria-busy: ', 'Loading states use aria-busy="true"'),
]:
    add_bullet(d, bold_prefix=b)

doc.add_heading('Gaps Requiring Attention', level=2)
t = styled_table(4, ['Issue', 'WCAG Criterion', 'Severity', 'Fix'])
for row in [
    ('Emoji icons lack role="img" + aria-label', '1.1.1 Non-text Content', 'Medium', 'Add role="img" aria-label on emoji'),
    ('Color-only priority indicators', '1.4.1 Use of Color', 'High', 'Add text labels alongside color dots'),
    ('Focus trap missing on some modals', '2.4.3 Focus Order', 'High', 'Implement focus trap on all modals'),
    ('Some inputs use placeholder-only', '1.3.1 Info and Relationships', 'Medium', 'Add visible labels above all inputs'),
    ('Mobile touch targets below 44px', '2.5.5 Target Size', 'Medium', 'Increase tap targets to 44x44px minimum'),
    ('Low contrast on --color-accent text', '1.4.3 Contrast', 'Medium', 'Use --color-accent-strong for text'),
]:
    add_table_row(t, row)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 9. RESPONSIVE & MOBILE
# ══════════════════════════════════════════════════════════
doc.add_heading('9. Responsive Design & Mobile Experience', level=1)

doc.add_heading('Breakpoint System', level=2)
t = styled_table(3, ['Breakpoint', 'Target', 'Usage'])
for row in [
    ('480px', 'Small phone', 'Auth pages, some mobile adjustments'),
    ('640px', 'Phone/small tablet', 'Student dashboard'),
    ('768px', 'Tablet', 'Primary breakpoint for sidebar toggle'),
    ('900px', 'Medium desktop', 'Some layout adjustments'),
    ('1024px', 'Desktop', 'Full-width layouts'),
]:
    add_table_row(t, row)

doc.add_paragraph()
doc.add_heading('Mobile Issues', level=2)
for b, d in [
    ('Hamburger menu: ', 'May not be familiar to less tech-savvy parents. Add "Menu" label.'),
    ('Modal overflow: ', 'Modals with many fields extend beyond viewport when keyboard opens.'),
    ('Tab scrolling: ', 'Child selector tabs scroll without visual cues (no gradient fade).'),
    ('Calendar cells: ', 'Day cells may be too small for comfortable tapping.'),
    ('Quick actions: ', '6-7 cards stack to single column, creating excessive scroll depth.'),
]:
    add_bullet(d, bold_prefix=b)

doc.add_heading('Recommendations', level=2)
for r in [
    'Standardize to 3 breakpoints: 480px (mobile), 768px (tablet), 1024px (desktop)',
    'Add scroll indicators on horizontally scrollable containers',
    'Use bottom sheet pattern for modals on mobile instead of centered overlay',
    'Use swipeable cards for quick actions instead of vertical stack',
    'Test all flows with virtual keyboard visible',
]:
    add_bullet(r)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 10. USER JOURNEY MAPS
# ══════════════════════════════════════════════════════════
doc.add_heading('10. User Journey Maps', level=1)
doc.add_paragraph('Journeys identified from the UI implementation, categorized by role, frequency, and completeness.')

doc.add_heading('Parent Journeys', level=2)
t = styled_table(5, ['#', 'Journey', 'Steps', 'Frequency', 'Completeness'])
for row in [
    ('P1', 'First-time Setup', 'Register > Onboarding > Add Child > Connect Google > Sync > Assign', 'Once', 'Complete (6 steps, reduce to 4)'),
    ('P2', 'Daily Check-in', 'Login > Dashboard > Today\'s Focus > Coming Up > Toggle Tasks', 'Daily', 'Complete'),
    ('P3', 'Generate Study Material', 'Dashboard > FAB > Upload > Paste/Upload > Type > Generate > View', 'Weekly', 'Complete (7 steps)'),
    ('P4', 'Message a Teacher', 'Dashboard > Messages > New > Select Teacher > Write > Send', 'Weekly', 'Complete'),
    ('P5', 'Add a New Child', 'My Kids > + > Method > Fill Details > Submit > Share Link', 'Rare', 'Complete'),
    ('P6', 'Review Child Progress', 'My Kids > Select Child > Completion % > Classes > Tasks', 'Weekly', 'Partial (no grades)'),
    ('P7', 'Reassign Material', 'My Kids > Child > Materials > Folder Icon > Search/Create > Confirm', 'Occasional', 'Complete'),
    ('P8', 'View Calendar', 'Tasks > Calendar > Browse Week > Click Day > View/Create Tasks', 'Weekly', 'Complete'),
]:
    add_table_row(t, row)

doc.add_paragraph()
doc.add_heading('Student Journeys', level=2)
t = styled_table(5, ['#', 'Journey', 'Steps', 'Frequency', 'Completeness'])
for row in [
    ('S1', 'First-time Setup', 'Register/Accept Invite > Onboarding > Connect Google > Sync', 'Once', 'Complete'),
    ('S2', 'Daily Study Check', 'Login > Dashboard > Coming Up > Check Streak > Select Material > Study', 'Daily', 'Complete'),
    ('S3', 'Create Study Guide', 'Dashboard > Upload > Paste Text > Select Type > Generate > Study', 'Weekly', 'Complete'),
    ('S4', 'Take Practice Quiz', 'Materials > Select Quiz > Take > Review Answers > Score', 'Weekly', 'Complete'),
    ('S5', 'Review Flashcards', 'Materials > Select Set > Flip Cards > Navigate > Shuffle', 'Daily', 'Partial (no spaced rep)'),
    ('S6', 'Check Assignments', 'Classes > Course > Assignments > Due Dates', 'Daily', 'Partial (no submission)'),
    ('S7', 'View Quiz History', 'Sidebar > Quiz History > Browse > Review', 'Weekly', 'Complete'),
    ('S8', 'Connect Google', 'Dashboard > Connect > OAuth > Auto-Sync > View Classes', 'Once', 'Complete'),
]:
    add_table_row(t, row)

doc.add_paragraph()
doc.add_heading('Teacher Journeys', level=2)
t = styled_table(5, ['#', 'Journey', 'Steps', 'Frequency', 'Completeness'])
for row in [
    ('T1', 'First-time Setup', 'Register > Onboarding > Teacher Type > Connect Google > Sync', 'Once', 'Complete'),
    ('T2', 'Daily Message Check', 'Login > Dashboard > Unread Count > Messages > Read/Reply', 'Daily', 'Complete'),
    ('T3', 'Send Announcement', 'Dashboard > Announcement > Select Class > Subject/Message > Send', 'Weekly', 'Partial (no preview)'),
    ('T4', 'Upload Material', 'Dashboard > Upload > Course > Type > File > Confirm', 'Weekly', 'Complete'),
    ('T5', 'Create New Class', 'Dashboard > + Create > Name/Subject > Submit', 'Rare', 'Complete'),
    ('T6', 'Invite a Parent', 'Dashboard > Invite > Enter Email > Send', 'Occasional', 'Complete'),
    ('T7', 'Sync Google', 'Dashboard > Google Card > Sync > View Updated List', 'Weekly', 'Complete'),
    ('T8', 'Monitor Email Comms', 'Sidebar > Teacher Comms > Filter > Read > Reply', 'Daily', 'Complete'),
]:
    add_table_row(t, row)

doc.add_paragraph()
doc.add_heading('Broken/Incomplete Journeys', level=2)
t = styled_table(4, ['Journey', 'Role', 'Issue', 'Impact'])
for row in [
    ('View child grades', 'Parent', 'No grade data displayed anywhere', 'High'),
    ('Submit homework', 'Student', 'Can view but cannot submit work', 'High'),
    ('Grade assignments', 'Teacher', 'No grading interface', 'High'),
    ('Track attendance', 'Teacher', 'No attendance feature', 'Medium'),
    ('Search messages', 'All', 'No search within conversations', 'Medium'),
    ('Share materials between children', 'Parent', 'Materials are per-child only', 'Low'),
    ('View class timetable', 'Student/Teacher', 'No schedule display', 'Medium'),
    ('Export data/reports', 'All', 'No export capability', 'Low'),
]:
    add_table_row(t, row)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 11. INDUSTRY BENCHMARKS
# ══════════════════════════════════════════════════════════
doc.add_heading('11. Industry Benchmarks & Inspiration', level=1)

for platform, desc, patterns in [
    ('Google Classroom', 'Simplicity is its strength.', [
        'Stream-based chronological activity feed',
        'Material attachment preview cards with type indicators',
        'Three-tab structure (Stream, Classwork, People)',
        'Class banner image personalization',
    ]),
    ('Notion for Education', 'Flexible workspace model.', [
        'Clean, minimal UI with generous whitespace',
        'Slash-command for rapid content creation',
        'Collapsible sidebar with favorites',
        'Template gallery for common document types',
    ]),
    ('Duolingo', 'Student engagement benchmark.', [
        'Study streak with freeze and milestone celebrations',
        'XP points and leaderboards',
        'Bright, playful illustrations',
        'Optimized push notification timing',
    ]),
    ('Canvas LMS', 'Most complete teacher experience.', [
        'SpeedGrader for rapid assignment grading',
        'Calendar with cross-course assignment integration',
        'Announcement scheduling and rich text editor',
        'Analytics dashboard with engagement metrics',
    ]),
    ('ClassDojo', 'Parent communication benchmark.', [
        'Story/timeline feed for class updates',
        'Simple messaging with read receipts and translation',
        'Student portfolio for showcasing work',
        'Behavior tracking with positive reinforcement points',
    ]),
]:
    doc.add_heading(platform, level=2)
    doc.add_paragraph(desc)
    for p in patterns:
        add_bullet(p)

doc.add_heading('Key Takeaways for ClassBridge', level=2)
for r in [
    'Adopt Notion\'s whitespace and progressive disclosure to reduce information overload',
    'Implement Duolingo-style streak celebrations and milestone badges for student engagement',
    'Add Canvas-style calendar integration showing all assignments/tasks in unified view',
    'Borrow ClassDojo\'s parent communication patterns: read receipts, translation, story feed',
    'Add Google Classroom-style class banner images for visual personalization',
]:
    add_bullet(r)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 12. RISK REGISTER
# ══════════════════════════════════════════════════════════
doc.add_heading('12. Risk Register', level=1)
t = styled_table(6, ['#', 'Risk', 'Likelihood (1-5)', 'Impact (1-5)', 'Score', 'Mitigation'])
for row in [
    ('R1', 'Hard-coded colors break theme adoption', '5', '3', '15', 'Audit and replace with CSS vars'),
    ('R2', 'Parent dashboard overwhelms new users', '4', '4', '16', 'Progressive disclosure; simplified toggle'),
    ('R3', 'Missing grades erodes parent trust', '4', '5', '20', 'Prioritize grade integration'),
    ('R4', 'A11y gaps expose legal risk (ADA/AODA)', '3', '5', '15', 'Complete WCAG 2.1 AA audit'),
    ('R5', 'Mobile usability reduces engagement', '4', '3', '12', 'Mobile-first testing pass'),
    ('R6', 'Teacher dashboard feels utilitarian', '3', '4', '12', 'Redesign with data-rich cards and SVG icons'),
    ('R7', 'No message search frustrates users', '3', '3', '9', 'Add full-text search'),
    ('R8', 'Z-index conflicts cause overlay bugs', '3', '3', '9', 'Document layered z-index system'),
    ('R9', 'Inconsistent breakpoints cause tablet bugs', '3', '3', '9', 'Standardize to 3 breakpoints'),
    ('R10', 'Study tools lack differentiation', '3', '4', '12', 'Add spaced repetition, timer, smart review'),
]:
    add_table_row(t, row)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 13. PRIORITIZED RECOMMENDATIONS
# ══════════════════════════════════════════════════════════
doc.add_heading('13. Prioritized Recommendations', level=1)

doc.add_heading('Tier 1: Critical (Next 2 sprints)', level=2)
for num, title, desc in [
    ('1.1', 'Unify Design Tokens', 'Audit all CSS and replace hard-coded values with CSS variables.'),
    ('1.2', 'Fix Accessibility Gaps', 'Focus traps on all modals, text labels alongside color indicators, 44px touch targets.'),
    ('1.3', 'Add "All Children" Tab', 'Replace deselect-to-show-all with explicit first tab. Add scroll indicators.'),
    ('1.4', 'Standardize Breakpoints', 'Consolidate to 480/768/1024. Fix mobile modal overflow.'),
    ('1.5', 'Simplify Parent Dashboard', 'Progressive disclosure: collapsed sections by default for new users.'),
]:
    doc.add_heading(f'{num} {title}', level=3)
    doc.add_paragraph(desc)

doc.add_heading('Tier 2: Important (Next 4 sprints)', level=2)
for num, title, desc in [
    ('2.1', 'Teacher Dashboard Enhancement', 'SVG icons, dynamic card counts, announcement preview, draft save.'),
    ('2.2', 'Student Engagement Features', 'Streak celebrations, spaced repetition, "Continue Studying" section.'),
    ('2.3', 'Consistent Empty States', 'Shared EmptyState component: icon + title + description + CTA.'),
    ('2.4', 'Message Search & Improvements', 'Full-text search, read receipts, attachment support.'),
    ('2.5', 'Standardize Type Scale', 'Reduce to 7 sizes with CSS variables.'),
    ('2.6', 'Unify Quick Action Paradigm', 'Consistent action pattern across all roles.'),
]:
    doc.add_heading(f'{num} {title}', level=3)
    doc.add_paragraph(desc)

doc.add_heading('Tier 3: Nice-to-Have (Backlog)', level=2)
for num, title, desc in [
    ('3.1', 'Grade Integration', 'Display grade data from Google Classroom for parents and students.'),
    ('3.2', 'Assignment Submission', 'Allow students to submit work through the platform.'),
    ('3.3', 'Data Export', 'PDF/CSV export for reports, grades, conversations.'),
    ('3.4', 'Co-Parent Support', 'Multiple parent accounts per child with shared visibility.'),
    ('3.5', 'Batch Task Operations', 'Multi-select with bulk actions.'),
    ('3.6', 'Calendar Unification', 'Merge assignments, tasks, events into single cross-course calendar.'),
    ('3.7', 'Study Timer & Focus Mode', 'Pomodoro-style timer with break reminders.'),
    ('3.8', 'Template Library', 'Announcement templates for teachers; task templates for parents.'),
]:
    doc.add_heading(f'{num} {title}', level=3)
    doc.add_paragraph(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 14. APPENDIX
# ══════════════════════════════════════════════════════════
doc.add_heading('14. Appendix: File Inventory', level=1)

doc.add_heading('Page Components (22+)', level=2)
t = styled_table(4, ['Page', 'File', 'Roles', 'Lines'])
for row in [
    ('Parent Dashboard', 'ParentDashboard.tsx', 'Parent', '500+'),
    ('Student Dashboard', 'StudentDashboard.tsx', 'Student', '1,155'),
    ('Teacher Dashboard', 'TeacherDashboard.tsx', 'Teacher', '1,009'),
    ('My Kids', 'MyKidsPage.tsx', 'Parent', '1,145'),
    ('Courses', 'CoursesPage.tsx', 'All', '400+'),
    ('Course Detail', 'CourseDetailPage.tsx', 'All', '400+'),
    ('Course Materials', 'CourseMaterialDetailPage.tsx', 'All', '400+'),
    ('Study Guides', 'StudyGuidesPage.tsx', 'All', '300+'),
    ('Tasks', 'TasksPage.tsx', 'Parent, Student', '400+'),
    ('Messages', 'MessagesPage.tsx', 'All', '400+'),
    ('Quiz', 'QuizPage.tsx', 'Student', '300+'),
    ('Flashcards', 'FlashcardsPage.tsx', 'Student', '200+'),
    ('Quiz History', 'QuizHistoryPage.tsx', 'Student', '200+'),
    ('Teacher Comms', 'TeacherCommsPage.tsx', 'Teacher', '300+'),
    ('Login', 'LoginPage.tsx', 'Public', '200+'),
    ('Register', 'RegisterPage.tsx', 'Public', '300+'),
    ('Onboarding', 'OnboardingPage.tsx', 'All', '200+'),
    ('Landing', 'LandingPage.tsx', 'Public', '300+'),
    ('Help', 'HelpPage.tsx', 'All', '200+'),
    ('FAQ', 'FAQPage.tsx', 'Public', '100+'),
    ('Notifications', 'NotificationsPage.tsx', 'All', '200+'),
    ('Analytics', 'AnalyticsPage.tsx', 'Parent, Student, Admin', '200+'),
]:
    add_table_row(t, row)

doc.add_paragraph()
doc.add_heading('Key CSS Files', level=2)
t = styled_table(3, ['File', 'Lines', 'Scope'])
for row in [
    ('index.css', '337', 'Global design tokens, themes, base styles'),
    ('Dashboard.css', '2,032', 'DashboardLayout, modals, shared components'),
    ('ParentDashboard.css', '2,212', 'Parent dashboard, child selector, timeline'),
    ('StudentDashboard.css', '800+', 'Student dashboard, hero, quick actions'),
    ('TeacherDashboard.css', '600+', 'Teacher dashboard, grid cards'),
    ('Auth.css', '336', 'Login, register, password reset'),
    ('Calendar.css', '500+', 'Calendar view, day cells, grids'),
    ('CourseMaterialDetailPage.css', '400+', 'Material detail, tabs, print view'),
]:
    add_table_row(t, row)

# ── End ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('--- End of Report ---')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x5B, 0x62, 0x74)
run.italic = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'This report was generated through systematic code-level analysis of the ClassBridge frontend codebase. '
    'Findings are based on source code review and do not include live user testing data.'
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

output_path = r'C:\dev\emai\stream-f\docs\ClassBridge_UI_UX_Assessment_Report.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
