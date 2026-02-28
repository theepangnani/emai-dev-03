"""Generate Word document from ClassBridge requirements markdown using python-docx."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

docs_dir = Path(__file__).parent
md_path = docs_dir / "ClassBridge_Complete_Requirements_Document.md"
docx_path = docs_dir / "ClassBridge_Complete_Requirements_Document.docx"

md_content = md_path.read_text(encoding="utf-8")

doc = Document()

# Style setup
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title styles
for level in range(1, 5):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Calibri'
    heading_style.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

def add_table_from_lines(doc, header_line, rows):
    """Parse markdown table and add to document."""
    headers = [cell.strip() for cell in header_line.strip('|').split('|')]
    headers = [h for h in headers if h]

    if not headers:
        return

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header.strip('*').strip()
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for row_text in rows:
        cells = [cell.strip() for cell in row_text.strip('|').split('|')]
        cells = [c for c in cells if c or cells.index(c) > 0]
        if not cells:
            continue
        row = table.add_row()
        for i, cell_text in enumerate(cells):
            if i < len(headers):
                cell = row.cells[i]
                # Clean markdown formatting
                clean = cell_text.strip()
                clean = re.sub(r'\*\*(.*?)\*\*', r'\1', clean)
                clean = re.sub(r'\*(.*?)\*', r'\1', clean)
                clean = re.sub(r'`(.*?)`', r'\1', clean)
                clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', clean)
                cell.text = clean
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)


def process_markdown(doc, content):
    """Process markdown content and add to Word document."""
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Horizontal rules
        if line.strip() == '---':
            doc.add_paragraph('_' * 60)
            i += 1
            continue

        # Headings
        if line.startswith('#'):
            level = 0
            for ch in line:
                if ch == '#':
                    level += 1
                else:
                    break
            text = line[level:].strip()
            # Clean markdown links
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            if level <= 4:
                doc.add_heading(text, level=min(level, 4))
            else:
                p = doc.add_paragraph(text)
                p.runs[0].bold = True
            i += 1
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            # Collect table
            table_lines = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1

            if len(table_lines) >= 2:
                header = table_lines[0]
                # Skip separator line
                data_rows = [l for l in table_lines[2:] if not re.match(r'^[\s|:-]+$', l)]
                add_table_from_lines(doc, header, data_rows)
                doc.add_paragraph('')  # spacing
            continue

        # Empty lines
        if line.strip() == '':
            i += 1
            continue

        # Bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Clean markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

            indent_level = (len(line) - len(line.lstrip())) // 2
            if indent_level > 0:
                p = doc.add_paragraph(text, style='List Bullet 2')
            else:
                p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Checkbox items
        if line.strip().startswith('- ['):
            checked = 'x' in line.strip()[3:4]
            text = line.strip()[5:].strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            prefix = "[x] " if checked else "[ ] "
            doc.add_paragraph(prefix + text, style='List Bullet')
            i += 1
            continue

        # Numbered lists
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Blockquotes
        if line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # Regular paragraph
        text = line.strip()
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'`(.*?)`', r'\1', text)
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        if text:
            doc.add_paragraph(text)
        i += 1


# Process the markdown
process_markdown(doc, md_content)

# Save
doc.save(str(docx_path))
print(f"Word document generated: {docx_path}")
print(f"File size: {docx_path.stat().st_size / 1024:.1f} KB")
