"""Generate Word document for ClassBridge Email Setup Guide using python-docx."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

docs_dir = Path(__file__).parent
docx_path = docs_dir / "ClassBridge_Email_Setup_Guide.docx"

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Style setup
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

# Heading styles - ClassBridge brand colors
brand_colors = {
    1: RGBColor(0x1B, 0x4F, 0x72),  # Dark blue
    2: RGBColor(0x2E, 0x74, 0x91),  # Medium blue
    3: RGBColor(0x49, 0xB8, 0xC0),  # Teal (ClassBridge primary)
    4: RGBColor(0x5A, 0x9E, 0x8F),  # Sage
}

for level in range(1, 5):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Calibri'
    heading_style.font.color.rgb = brand_colors[level]
    if level == 1:
        heading_style.font.size = Pt(20)
    elif level == 2:
        heading_style.font.size = Pt(16)
    elif level == 3:
        heading_style.font.size = Pt(13)
    else:
        heading_style.font.size = Pt(11)


def add_cover_page(doc):
    """Add a professional cover page."""
    for _ in range(4):
        doc.add_paragraph('')

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ClassBridge')
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x49, 0xB8, 0xC0)
    run.font.name = 'Calibri'
    run.bold = True

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Email Setup Guide')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    run.font.name = 'Calibri'

    doc.add_paragraph('')

    # Divider line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_' * 50)
    run.font.color.rgb = RGBColor(0x49, 0xB8, 0xC0)

    doc.add_paragraph('')

    # Metadata
    metadata = [
        ('Version', '1.0'),
        ('Date', '2026-03-22'),
        ('Providers', 'SendGrid (primary) + Gmail SMTP (fallback)'),
        ('Platform', 'GCP Cloud Run'),
        ('Domain', 'classbridge.ca'),
    ]

    for label, value in metadata:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}: ')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.name = 'Calibri'
        run.bold = True
        run = p.add_run(value)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run.font.name = 'Calibri'

    doc.add_paragraph('')
    doc.add_paragraph('')

    # Tagline
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('AI-Powered Education Platform')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x49, 0xB8, 0xC0)
    run.font.name = 'Calibri'
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Connecting Parents, Students, Teachers & Administrators')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = 'Calibri'
    run.italic = True

    doc.add_page_break()


def add_bullet(doc, text, indent=0):
    """Add a bullet point with optional indentation."""
    if indent > 0:
        try:
            p = doc.add_paragraph(text, style='List Bullet 2')
        except KeyError:
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
    else:
        p = doc.add_paragraph(text, style='List Bullet')
    return p


def add_code_block(doc, code_text):
    """Add a formatted code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_inline_code_paragraph(doc, parts):
    """Add a paragraph with mixed normal and code-formatted text.
    parts is a list of tuples: (text, is_code)
    """
    p = doc.add_paragraph()
    for text, is_code in parts:
        run = p.add_run(text)
        if is_code:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        else:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
    return p


def add_bullet_with_code(doc, normal_text, code_text=None, indent=0):
    """Add a bullet point that may contain inline code."""
    if indent > 0:
        try:
            p = doc.add_paragraph(style='List Bullet 2')
        except KeyError:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
    else:
        p = doc.add_paragraph(style='List Bullet')

    run = p.add_run(normal_text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    if code_text:
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_important(doc, text):
    """Add an IMPORTANT note with bold prefix."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run('IMPORTANT: ')
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p


# ============================================================
# Build the document
# ============================================================

add_cover_page(doc)

# --- Section 1: Overview ---
doc.add_heading('Overview', level=1)
add_bullet(doc, 'ClassBridge uses two email providers: SendGrid (primary) and Gmail SMTP (fallback)')
add_bullet_with_code(doc, 'The email service in ', 'app/services/email_service.py')
p = doc.paragraphs[-1]
run = p.add_run(' tries SendGrid first, then falls back to SMTP')
run.font.name = 'Calibri'
run.font.size = Pt(11)

add_bullet(doc, 'Emails are sent for: waitlist confirmations, admin notifications, welcome emails, '
           'password resets, daily/weekly digests, teacher communication notifications')

# --- Section 2: SendGrid Setup ---
doc.add_heading('SendGrid Setup', level=1)

# 2.1
doc.add_heading('Create a SendGrid Account', level=2)
add_bullet(doc, 'Go to sendgrid.com and sign up (free tier: 100 emails/day)')
add_bullet(doc, 'Select the 60-day trial plan ($0/month) \u2014 converts to free tier automatically after trial')
add_bullet(doc, 'For Marketing Campaigns plan, also select 60-day trial (not used by ClassBridge)')

# 2.2
doc.add_heading('Domain Authentication (DNS Records)', level=2)
add_bullet(doc, 'In SendGrid onboarding, enter domain: classbridge.ca')
add_bullet(doc, 'Select "Yes" for link branding')
add_bullet(doc, 'Leave Advanced Settings defaults (automated security checked, everything else unchecked)')
add_bullet(doc, 'SendGrid will generate 6 DNS records to add to your domain registrar (GoDaddy):')
add_bullet(doc, '5 CNAME records (link tracking, mail, DKIM keys)', indent=1)
add_bullet(doc, '1 TXT record (DMARC)', indent=1)
add_bullet(doc, 'In GoDaddy (dcc.godaddy.com), go to Domain Portfolio > classbridge.ca > DNS > Add New Record')
add_bullet(doc, 'Add each record one by one with default TTL')
add_bullet(doc, 'Go back to SendGrid and verify the DNS records (may take 5\u201310 minutes to propagate)')

# 2.3
doc.add_heading('Create Sender Identity', level=2)
add_bullet(doc, 'Go to Settings > Sender Authentication > Single Sender Verification')
add_bullet(doc, 'From Name: ClassBridge')
add_bullet(doc, 'From Email: clazzbridge@gmail.com (or noreply@classbridge.ca if mailbox exists)')
add_bullet(doc, 'Reply To: clazzbridge@gmail.com')
add_bullet(doc, 'Fill in company address details')
add_bullet(doc, 'Save and verify')

# 2.4
doc.add_heading('Create API Key', level=2)
add_bullet(doc, 'Go to Settings > API Keys > Create API Key')
add_bullet(doc, 'Name: "ClassBridge Production"')
add_bullet(doc, 'Permissions: Full Access')
add_bullet(doc, 'Click Create & View')
add_bullet(doc, 'Copy the key (starts with SG.) \u2014 only shown once')
add_important(doc, 'When storing the key, ensure no trailing newline/whitespace is added')

# 2.5
doc.add_heading('Store API Key in GCP Secret Manager', level=2)
add_bullet(doc, 'Command:')
add_code_block(doc, "printf '%s' 'SG.your-key-here' | gcloud secrets versions add SENDGRID_API_KEY "
               "--project=emai-dev-01 --data-file=-")
add_bullet_with_code(doc, 'Use ', "printf '%s'")
p = doc.paragraphs[-1]
run = p.add_run(' (not ')
run.font.name = 'Calibri'
run.font.size = Pt(11)
run = p.add_run('echo')
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
run = p.add_run(') to avoid trailing newline characters that will cause authentication failures')
run.font.name = 'Calibri'
run.font.size = Pt(11)

add_bullet(doc, 'Redeploy Cloud Run:')
add_code_block(doc, "gcloud run services update classbridge --project=emai-dev-01 "
               "--region=us-central1 --update-secrets=SENDGRID_API_KEY=SENDGRID_API_KEY:latest")

# --- Section 3: Gmail SMTP Setup ---
doc.add_heading('Gmail SMTP Setup (Fallback)', level=1)

# 3.1
doc.add_heading('Prerequisites', level=2)
add_bullet(doc, 'Gmail account: clazzbridge@gmail.com')
add_bullet(doc, '2-Step Verification must be enabled')

# 3.2
doc.add_heading('Generate App Password', level=2)
add_bullet(doc, 'Go to myaccount.google.com/apppasswords')
add_bullet(doc, 'Type app name: "ClassBridge"')
add_bullet(doc, 'Click Create')
add_bullet(doc, 'Copy the 16-character password')
add_important(doc, 'Changing the Gmail account password invalidates ALL existing App Passwords \u2014 '
              'you must generate a new one')

# 3.3
doc.add_heading('Store Credentials in GCP Secret Manager', level=2)
add_bullet(doc, 'SMTP_USER:')
add_code_block(doc, "echo 'clazzbridge@gmail.com' | gcloud secrets versions add SMTP_USER "
               "--project=emai-dev-01 --data-file=-")
add_bullet(doc, 'SMTP_PASSWORD:')
add_code_block(doc, "printf '%s' 'your-app-password-no-spaces' | gcloud secrets versions add SMTP_PASSWORD "
               "--project=emai-dev-01 --data-file=-")
add_bullet(doc, 'Redeploy Cloud Run:')
add_code_block(doc, "gcloud run services update classbridge --project=emai-dev-01 "
               "--region=us-central1 --update-secrets=SMTP_PASSWORD=SMTP_PASSWORD:latest")

# --- Section 4: Cloud Run Configuration ---
doc.add_heading('Cloud Run Configuration', level=1)

# 4.1
doc.add_heading('Environment Variables', level=2)
add_bullet_with_code(doc, 'FROM_EMAIL=noreply@classbridge.ca \u2014 set in both deploy configs:')
add_bullet_with_code(doc, '', '.github/workflows/deploy.yml', indent=1)
p = doc.paragraphs[-1]
run = p.add_run(' (CI/CD)')
run.font.name = 'Calibri'
run.font.size = Pt(11)

add_bullet_with_code(doc, '', 'scripts/deploy.sh', indent=1)
p = doc.paragraphs[-1]
run = p.add_run(' (manual deploy)')
run.font.name = 'Calibri'
run.font.size = Pt(11)

add_bullet_with_code(doc, 'The deploy workflow uses ', '--set-env-vars')
p = doc.paragraphs[-1]
run = p.add_run(' which replaces ALL env vars on each deploy, so any new env var must be added to these files')
run.font.name = 'Calibri'
run.font.size = Pt(11)

# 4.2
doc.add_heading('Secrets (via GCP Secret Manager)', level=2)
add_bullet_with_code(doc, 'SENDGRID_API_KEY \u2014 SendGrid API key (starts with SG.)')
add_bullet_with_code(doc, 'SMTP_USER \u2014 Gmail address (clazzbridge@gmail.com)')
add_bullet_with_code(doc, 'SMTP_PASSWORD \u2014 Gmail App Password')
add_bullet(doc, 'Referenced in deploy as:')
add_code_block(doc, '--set-secrets "...SENDGRID_API_KEY=SENDGRID_API_KEY:latest,'
               'SMTP_USER=SMTP_USER:latest,SMTP_PASSWORD=SMTP_PASSWORD:latest"')

# --- Section 5: Troubleshooting ---
doc.add_heading('Troubleshooting', level=1)

# 5.1
doc.add_heading('No Emails Being Sent', level=2)
add_bullet(doc, 'Check Cloud Run logs:')
add_code_block(doc, 'gcloud logging read \'resource.type="cloud_run_revision" AND '
               'resource.labels.service_name="classbridge" AND textPayload:"email"\' '
               '--project=emai-dev-01 --limit=20 --format=json --freshness=1d')
add_bullet(doc, 'Common errors:')

add_bullet(doc, '"Username and Password not accepted"', indent=1)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.8)
run = p.add_run('Fix: ')
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(11)
run = p.add_run('Gmail App Password expired/invalidated. Generate new one at myaccount.google.com/apppasswords')
run.font.name = 'Calibri'
run.font.size = Pt(11)

add_bullet(doc, '"Invalid header value" with \\r\\n', indent=1)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.8)
run = p.add_run('Fix: ')
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(11)
run = p.add_run("Trailing newline in API key. Re-store using printf '%s' instead of echo")
run.font.name = 'Calibri'
run.font.size = Pt(11)

add_bullet(doc, '"No email provider configured"', indent=1)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.8)
run = p.add_run('Fix: ')
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(11)
run = p.add_run('Both SENDGRID_API_KEY and SMTP_USER/SMTP_PASSWORD are empty or not set')
run.font.name = 'Calibri'
run.font.size = Pt(11)

# 5.2
doc.add_heading('SendGrid Failing, Falling Back to SMTP', level=2)
add_bullet(doc, 'Verify SENDGRID_API_KEY starts with "SG." \u2014 check with:')
add_code_block(doc, 'gcloud secrets versions access latest --secret=SENDGRID_API_KEY --project=emai-dev-01')
add_bullet(doc, 'If key is placeholder "your-sendgrid-key", replace with real key from SendGrid dashboard')

# 5.3
doc.add_heading('Emails Going to Spam', level=2)
add_bullet(doc, 'Ensure domain authentication is verified in SendGrid (Settings > Sender Authentication)')
add_bullet_with_code(doc, 'Ensure DMARC TXT record exists: ', '_dmarc.classbridge.ca')
p = doc.paragraphs[-1]
run = p.add_run(' \u2192 ')
run.font.name = 'Calibri'
run.font.size = Pt(11)
run = p.add_run('v=DMARC1; p=none;')
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

add_bullet(doc, 'Use noreply@classbridge.ca as from-address (not gmail.com)')

# --- Section 6: Architecture Reference ---
doc.add_heading('Architecture Reference', level=1)

add_bullet_with_code(doc, 'Email service: ', 'app/services/email_service.py')
add_bullet_with_code(doc, 'Config: ', 'app/core/config.py')
p = doc.paragraphs[-1]
run = p.add_run(' (settings: sendgrid_api_key, from_email, smtp_host, smtp_port, smtp_user, smtp_password)')
run.font.name = 'Calibri'
run.font.size = Pt(11)

add_bullet_with_code(doc, 'Flow: ', 'send_email()')
p = doc.paragraphs[-1]
run = p.add_run(' \u2192 tries SendGrid (')
run.font.name = 'Calibri'
run.font.size = Pt(11)
run = p.add_run('_send_via_sendgrid')
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
run = p.add_run(') \u2192 falls back to SMTP (')
run.font.name = 'Calibri'
run.font.size = Pt(11)
run = p.add_run('_send_via_smtp')
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
run = p.add_run(')')
run.font.name = 'Calibri'
run.font.size = Pt(11)

add_bullet_with_code(doc, 'Batch sending: ', 'send_emails_batch()')
p = doc.paragraphs[-1]
run = p.add_run(' for digest emails')
run.font.name = 'Calibri'
run.font.size = Pt(11)

add_bullet_with_code(doc, 'Email template wrapper: ', 'wrap_branded_email()')
p = doc.paragraphs[-1]
run = p.add_run(' for consistent branding')
run.font.name = 'Calibri'
run.font.size = Pt(11)

# --- Section 7: Current Configuration ---
doc.add_heading('Current Configuration (as of March 22, 2026)', level=1)

config_items = [
    'Primary provider: SendGrid (domain-authenticated for classbridge.ca)',
    'Fallback provider: Gmail SMTP (clazzbridge@gmail.com)',
    'From address: noreply@classbridge.ca',
    'SendGrid plan: Free trial (100 emails/day, converts to free tier after 60 days)',
    'DNS records: 5 CNAME + 1 TXT on GoDaddy for classbridge.ca',
    'GCP project: emai-dev-01',
    'Cloud Run service: classbridge (us-central1)',
]

for item in config_items:
    add_bullet(doc, item)

# Add footer with page numbers
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run('ClassBridge Email Setup Guide  |  ')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = 'Calibri'

    # Page number field
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    run2 = p.add_run()
    run2._r.append(fld_char_begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    run3 = p.add_run()
    run3._r.append(instr)

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run4 = p.add_run()
    run4._r.append(fld_char_end)

# Save
doc.save(str(docx_path))
print(f"Word document generated: {docx_path}")
print(f"File size: {docx_path.stat().st_size / 1024:.1f} KB")
