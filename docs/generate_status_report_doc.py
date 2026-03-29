"""Generate Word document from ClassBridge Project Status Report markdown using python-docx."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

docs_dir = Path(__file__).parent
md_path = docs_dir / "ClassBridge_Project_Status_Report.md"
docx_path = docs_dir / "ClassBridge_Project_Status_Report.docx"

md_content = md_path.read_text(encoding="utf-8")

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Style setup
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

# Heading styles
brand_colors = {
    1: RGBColor(0x1B, 0x4F, 0x72),
    2: RGBColor(0x2E, 0x74, 0x91),
    3: RGBColor(0x49, 0xB8, 0xC0),
    4: RGBColor(0x5A, 0x9E, 0x8F),
}

for level in range(1, 5):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Calibri'
    heading_style.font.color.rgb = brand_colors[level]
    if level == 1:
        heading_style.font.size = Pt(20)
    elif level == 2:
        heading_style.font.size = Pt(16)
    elif level == 3:
        heading_style.font.size = Pt(13)
    else:
        heading_style.font.size = Pt(11)


def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


# RAG status colors
RAG_COLORS = {
    'green': '27AE60',
    'amber': 'F39C12',
    'red': 'E74C3C',
}


def add_table_from_lines(doc, header_line, rows):
    headers = [cell.strip() for cell in header_line.strip('|').split('|')]
    headers = [h for h in headers if h]
    if not headers:
        return

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header.strip('*').strip()
        set_cell_shading(cell, '49B8C0')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = 'Calibri'

    for idx, row_text in enumerate(rows):
        cells = [cell.strip() for cell in row_text.strip('|').split('|')]
        cells = [c for c in cells if c or cells.index(c) > 0]
        if not cells:
            continue
        row = table.add_row()
        for i, cell_text in enumerate(cells):
            if i < len(headers):
                cell = row.cells[i]
                clean = cell_text.strip()
                clean = re.sub(r'\*\*(.*?)\*\*', r'\1', clean)
                clean = re.sub(r'\*(.*?)\*', r'\1', clean)
                clean = re.sub(r'`(.*?)`', r'\1', clean)
                clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', clean)
                cell.text = clean
                if idx % 2 == 1:
                    set_cell_shading(cell, 'F0F7FA')
                # Color-code RAG status cells
                clean_lower = clean.lower()
                if clean_lower in RAG_COLORS:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            r_val = int(RAG_COLORS[clean_lower][:2], 16)
                            g_val = int(RAG_COLORS[clean_lower][2:4], 16)
                            b_val = int(RAG_COLORS[clean_lower][4:6], 16)
                            run.font.color.rgb = RGBColor(r_val, g_val, b_val)
                            run.bold = True
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Calibri'


def add_cover_page(doc):
    for _ in range(4):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ClassBridge')
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x49, 0xB8, 0xC0)
    run.font.name = 'Calibri'
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Project Status Report')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    run.font.name = 'Calibri'

    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_' * 50)
    run.font.color.rgb = RGBColor(0x49, 0xB8, 0xC0)

    doc.add_paragraph('')

    metadata = [
        ('Version', '1.0'),
        ('Date', '2026-03-27'),
        ('Author', 'Theepan Gnanasabapathy'),
        ('Status', 'Active Development'),
    ]

    for label, value in metadata:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}: ')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.name = 'Calibri'
        run.bold = True
        run = p.add_run(value)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run.font.name = 'Calibri'

    doc.add_paragraph('')
    doc.add_paragraph('')

    # Color swatches
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    colors = [
        ('#49B8C0', 'Teal'),
        ('#1B4F72', 'Blue'),
        ('#27AE60', 'Green'),
        ('#F39C12', 'Amber'),
        ('#E74C3C', 'Red'),
    ]
    for hex_color, name in colors:
        r_val = int(hex_color[1:3], 16)
        g_val = int(hex_color[3:5], 16)
        b_val = int(hex_color[5:7], 16)
        run = p.add_run(f'  {name}  ')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(r_val, g_val, b_val)
        run.bold = True
        run.font.name = 'Calibri'

    doc.add_page_break()


def process_markdown(doc, content):
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    skip_first_h1 = True

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_lines)
                if code_text.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if line.strip() == '---':
            i += 1
            continue

        if line.startswith('#'):
            level = 0
            for ch in line:
                if ch == '#':
                    level += 1
                else:
                    break
            text = line[level:].strip()
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

            if level == 1 and skip_first_h1:
                skip_first_h1 = False
                i += 1
                continue

            if level <= 4:
                doc.add_heading(text, level=min(level, 4))
            else:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(11)
            i += 1
            continue

        if '|' in line and line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                header = table_lines[0]
                data_rows = [l for l in table_lines[2:] if not re.match(r'^[\s|:-]+$', l)]
                add_table_from_lines(doc, header, data_rows)
                doc.add_paragraph('')
            continue

        if line.strip() == '':
            i += 1
            continue

        if re.match(r'^\s*-\s*\[[ x]\]', line.strip()):
            checked = '[x]' in line
            text = re.sub(r'^-\s*\[[ x]\]\s*', '', line.strip())
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            prefix = "[DONE] " if checked else "[ ] "
            doc.add_paragraph(prefix + text, style='List Bullet')
            i += 1
            continue

        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()
            text = text[2:]
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            indent_level = (len(line) - len(line.lstrip())) // 2
            if indent_level > 0:
                try:
                    p = doc.add_paragraph(text, style='List Bullet 2')
                except KeyError:
                    p = doc.add_paragraph(text, style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.5)
            else:
                p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        if line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        text = line.strip()
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'`(.*?)`', r'\1', text)
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        if text:
            doc.add_paragraph(text)
        i += 1


# Build document
add_cover_page(doc)
process_markdown(doc, md_content)

# Footer
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ClassBridge Project Status Report  |  ')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = 'Calibri'

    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    run2 = p.add_run()
    run2._r.append(fld_char_begin)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    run3 = p.add_run()
    run3._r.append(instr)
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run4 = p.add_run()
    run4._r.append(fld_char_end)

doc.save(str(docx_path))
print(f"Word document generated: {docx_path}")
print(f"File size: {docx_path.stat().st_size / 1024:.1f} KB")
