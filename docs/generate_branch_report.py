"""Generate a Word document report of all 340 unmerged remote branches with action recommendations."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from collections import defaultdict, Counter
import os
from datetime import datetime


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_row_font(row_cells, size=8):
    for cell in row_cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(size)


def add_styled_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row_data in rows:
        row = tbl.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
        set_row_font(row, 8)
    return tbl


def parse_analysis(filepath):
    entries = []
    with open(filepath, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            parts = line.split("|", 9)
            if len(parts) >= 10:
                entries.append({
                    "action": parts[0],
                    "category": parts[1],
                    "date": parts[2],
                    "branch": parts[3],
                    "issue": parts[4] if parts[4] else "—",
                    "issue_state": parts[5],
                    "files": int(parts[6]),
                    "lines": int(parts[7]),
                    "reason": parts[8],
                    "message": parts[9][:120],
                })
    return entries


def make_report(entries, output_path):
    doc = Document()

    # ── Title ──
    title = doc.add_heading("ClassBridge — Remote Branch Audit Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n"
        f"Repository: theepangnani/emai-dev-03\n"
        f"Total unmerged remote branches audited: {len(entries)}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # ── Executive Summary ──
    doc.add_heading("1. Executive Summary", level=1)

    actions = Counter(e["action"] for e in entries)
    total_delete_lines = sum(e["lines"] for e in entries if e["action"] == "DELETE")
    total_keep_lines = sum(e["lines"] for e in entries if e["action"] == "KEEP")
    total_review_lines = sum(e["lines"] for e in entries if e["action"] == "REVIEW")

    p = doc.add_paragraph()
    p.add_run("Recommendation: ").bold = True
    p.add_run(
        f"Delete {actions['DELETE']} branches ({actions['DELETE']/len(entries)*100:.0f}%), "
        f"keep {actions['KEEP']} branches, and review {actions['REVIEW']} branches manually."
    )

    add_styled_table(doc,
        ["Action", "Count", "% of Total", "Total Lines of Code", "Rationale"],
        [
            ("DELETE", actions["DELETE"], f"{actions['DELETE']/len(entries)*100:.0f}%",
             f"{total_delete_lines:,}", "Linked GitHub issue is CLOSED — work already shipped or abandoned"),
            ("KEEP", actions["KEEP"], f"{actions['KEEP']/len(entries)*100:.0f}%",
             f"{total_keep_lines:,}", "Linked GitHub issue is still OPEN — work may be needed"),
            ("REVIEW", actions["REVIEW"], f"{actions['REVIEW']/len(entries)*100:.0f}%",
             f"{total_review_lines:,}", "No GitHub issue linked — manual review needed"),
        ]
    )

    doc.add_paragraph("")

    # ── Category breakdown ──
    doc.add_heading("2. Breakdown by Category", level=1)
    cat_counts = defaultdict(lambda: {"DELETE": 0, "KEEP": 0, "REVIEW": 0, "total": 0})
    for e in entries:
        cat_counts[e["category"]][e["action"]] += 1
        cat_counts[e["category"]]["total"] += 1

    add_styled_table(doc,
        ["Category", "Total", "Delete", "Keep", "Review"],
        [(cat, d["total"], d["DELETE"], d["KEEP"], d["REVIEW"])
         for cat, d in sorted(cat_counts.items(), key=lambda x: -x[1]["total"])]
    )

    doc.add_paragraph("")

    # ── Age breakdown ──
    doc.add_heading("3. Breakdown by Age", level=1)

    def age_bucket(date_str):
        if date_str <= "2025-12-31": return "2025 or older"
        elif date_str <= "2026-01-31": return "January 2026"
        elif date_str <= "2026-02-28": return "February 2026"
        elif date_str <= "2026-03-15": return "March 1-15"
        elif date_str <= "2026-03-22": return "March 16-22"
        else: return "March 23-25"

    age_data = defaultdict(lambda: {"DELETE": 0, "KEEP": 0, "REVIEW": 0, "total": 0})
    for e in entries:
        bucket = age_bucket(e["date"])
        age_data[bucket][e["action"]] += 1
        age_data[bucket]["total"] += 1

    age_order = ["2025 or older", "January 2026", "February 2026",
                 "March 1-15", "March 16-22", "March 23-25"]
    add_styled_table(doc,
        ["Period", "Total", "Delete", "Keep", "Review"],
        [(period, age_data[period]["total"], age_data[period]["DELETE"],
          age_data[period]["KEEP"], age_data[period]["REVIEW"])
         for period in age_order if age_data[period]["total"] > 0]
    )

    doc.add_paragraph("")

    # ── Section: DELETE branches ──
    doc.add_heading("4. Branches to DELETE (260)", level=1)
    doc.add_paragraph(
        "These branches reference GitHub issues that are CLOSED. The work has either been "
        "merged to master via integration branches, superseded by other work, or abandoned. "
        "Deleting these branches removes clutter without losing any work — the commits remain "
        "accessible via git reflog and GitHub's branch recovery for 90 days."
    )

    delete_entries = sorted(
        [e for e in entries if e["action"] == "DELETE"],
        key=lambda x: x["date"], reverse=True
    )

    # Group by category for DELETE
    del_by_cat = defaultdict(list)
    for e in delete_entries:
        del_by_cat[e["category"]].append(e)

    for cat in ["Bug Fix", "Feature", "Integration", "Design/UI", "Documentation", "Worktree", "Other"]:
        items = del_by_cat.get(cat, [])
        if not items:
            continue
        doc.add_heading(f"{cat} ({len(items)})", level=3)
        add_styled_table(doc,
            ["Date", "Branch", "Issue", "Files", "Lines", "Commit Message"],
            [(e["date"], e["branch"], f"#{e['issue']}" if e["issue"] != "—" else "—",
              e["files"], e["lines"], e["message"])
             for e in items]
        )
        doc.add_paragraph("")

    # ── Section: KEEP branches ──
    doc.add_heading("5. Branches to KEEP (11)", level=1)
    doc.add_paragraph(
        "These branches reference GitHub issues that are still OPEN. They contain work "
        "that may be needed for upcoming features or the April pilot launch."
    )

    keep_entries = sorted(
        [e for e in entries if e["action"] == "KEEP"],
        key=lambda x: -x["lines"]
    )
    add_styled_table(doc,
        ["Date", "Branch", "Issue", "State", "Files", "Lines", "Commit Message"],
        [(e["date"], e["branch"], f"#{e['issue']}", e["issue_state"],
          e["files"], e["lines"], e["message"])
         for e in keep_entries]
    )

    doc.add_paragraph("")

    # ── Section: REVIEW branches ──
    doc.add_heading("6. Branches to REVIEW (69)", level=1)
    doc.add_paragraph(
        "These branches have no linked GitHub issue or the issue couldn't be resolved. "
        "Manual review is recommended. Branches with fewer than 20 lines of changes are "
        "likely trivial fixes that were superseded."
    )

    review_entries = sorted(
        [e for e in entries if e["action"] == "REVIEW"],
        key=lambda x: -x["lines"]
    )

    # Split into likely-delete and needs-review
    trivial = [e for e in review_entries if e["lines"] <= 50]
    substantial = [e for e in review_entries if e["lines"] > 50]

    doc.add_heading(f"Substantial (>50 lines) — {len(substantial)} branches", level=3)
    add_styled_table(doc,
        ["Date", "Branch", "Files", "Lines", "Commit Message"],
        [(e["date"], e["branch"], e["files"], e["lines"], e["message"])
         for e in substantial]
    )

    doc.add_paragraph("")

    doc.add_heading(f"Trivial (<=50 lines) — {len(trivial)} branches, likely safe to delete", level=3)
    add_styled_table(doc,
        ["Date", "Branch", "Files", "Lines", "Commit Message"],
        [(e["date"], e["branch"], e["files"], e["lines"], e["message"])
         for e in trivial]
    )

    doc.add_paragraph("")

    # ── Issue Cross-Reference ──
    doc.add_heading("7. GitHub Issue Cross-Reference", level=1)

    issues_map = defaultdict(lambda: {"branches": [], "state": "UNKNOWN"})
    for e in entries:
        if e["issue"] != "—":
            issues_map[e["issue"]]["branches"].append(e["branch"])
            issues_map[e["issue"]]["state"] = e["issue_state"]

    add_styled_table(doc,
        ["Issue #", "State", "# Branches", "Branch Names"],
        [(f"#{iss}", data["state"], len(data["branches"]),
          ", ".join(data["branches"])[:150])
         for iss, data in sorted(issues_map.items(),
                                  key=lambda x: int(x[0]) if x[0].isdigit() else 0)]
    )

    # ── Save ──
    doc.save(output_path)
    print(f"Report saved to {output_path}")
    print(f"  DELETE: {actions['DELETE']} branches")
    print(f"  KEEP: {actions['KEEP']} branches")
    print(f"  REVIEW: {actions['REVIEW']} branches")


if __name__ == "__main__":
    tmp_path = os.path.join(os.environ.get("TEMP", "/tmp"), "branch_full_analysis.txt")
    entries = parse_analysis(tmp_path)
    make_report(entries, "docs/ClassBridge_Unmerged_Branch_Report.docx")
