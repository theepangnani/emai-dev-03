"""
Generate DTAP Compliance Research Report for ClassBridge (EMAI)
Word document (.docx) format
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import datetime


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._element.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')


def add_table_row(table, cells, bold=False, header=False):
    """Add a row to a table with formatting."""
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        if bold or header:
            run.bold = True
        if header:
            set_cell_shading(cell, "1F4E79")
            run.font.color.rgb = RGBColor(255, 255, 255)
    return row


def create_report():
    doc = Document()

    # -- Page setup --
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # -- Styles --
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Calibri'
        h_style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        if level == 1:
            h_style.font.size = Pt(22)
            h_style.paragraph_format.space_before = Pt(24)
            h_style.paragraph_format.space_after = Pt(12)
        elif level == 2:
            h_style.font.size = Pt(16)
            h_style.paragraph_format.space_before = Pt(18)
            h_style.paragraph_format.space_after = Pt(8)
        else:
            h_style.font.size = Pt(13)
            h_style.paragraph_format.space_before = Pt(12)
            h_style.paragraph_format.space_after = Pt(6)

    # ================================================================
    # COVER PAGE
    # ================================================================
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Digital Technology Approval Process\n(DTAP) — Ontario School Boards")
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("DTAP & ECNO VASP Compliance Research & Gap Analysis\nfor ClassBridge (EMAI)")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x4A, 0x86, 0xC8)

    doc.add_paragraph("")
    doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Prepared: {datetime.date.today().strftime('%B %d, %Y')}\nVersion: 2.0\nClassification: Internal / Confidential")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ================================================================
    # TABLE OF CONTENTS (manual)
    # ================================================================
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Understanding DTAP and ECNO VASP",
        "3. ECNO — The Central Organization",
        "4. ECNO VASP — The Provincial Vetting Process",
        "5. ECNO Student Digital Privacy Standard",
        "6. Legal & Regulatory Framework",
        "7. Technical Requirements for EdTech Vendors",
        "8. OSAPAC — Provincial Software Licensing",
        "9. Board-Level vs Provincial-Level Approval",
        "10. ClassBridge Current Architecture Assessment",
        "11. Compliance Gap Analysis",
        "12. Required Architecture & Design Changes",
        "13. Likelihood of ClassBridge Passing DTAP",
        "14. Recommended Roadmap & GitHub Issues",
        "15. Cost & Timeline Estimates",
        "16. Sources & References",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ================================================================
    # 1. EXECUTIVE SUMMARY
    # ================================================================
    doc.add_heading("1. Executive Summary", level=1)

    doc.add_paragraph(
        "This report provides comprehensive research into the technology approval processes used by Ontario's "
        "72 publicly funded school boards, with a specific focus on how ClassBridge (EMAI) — an AI-powered "
        "education management platform — can achieve approval for deployment in Ontario schools."
    )

    doc.add_heading("Key Findings", level=3)
    findings = [
        "DTAP (Digital Technology Approval Process) is the generic term for each school board's internal process to authorize new technology by assessing security, privacy, and pedagogical/business alignment. Each of Ontario's 72 boards runs its own DTAP.",
        "ECNO VASP (Vetting of Application Security & Privacy) is the provincial-level service that feeds into board-level DTAP processes. VASP produces a risk assessment report that boards use as a key input to their DTAP decision.",
        "Vendors cannot apply directly to VASP. A school board partner must submit the request on your behalf.",
        "ClassBridge has a solid foundation but requires significant work in privacy compliance, accessibility, Canadian data residency, and security certifications before it would likely pass a board's DTAP or receive a favourable VASP risk score.",
        "New legislation (Bill 194 / EDSTA) effective 2025 significantly raises the bar for privacy and cybersecurity, making early compliance a competitive advantage.",
        "The estimated timeline to reach DTAP/VASP-readiness is 6-12 months with focused effort.",
    ]
    for f in findings:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading("ClassBridge Readiness Score", level=3)
    doc.add_paragraph(
        "Based on our gap analysis, ClassBridge currently meets approximately 35-40% of the requirements "
        "for passing a board's DTAP and receiving a favourable VASP risk score. The major gaps are in: "
        "Canadian data residency (currently US-hosted), MFIPPA/PIPEDA-compliant privacy documentation, "
        "WCAG 2.1 AA accessibility, MFA/SSO support, and third-party security certifications (SOC 2 Type II)."
    )

    doc.add_page_break()

    # ================================================================
    # 2. UNDERSTANDING DTAP AND ECNO VASP
    # ================================================================
    doc.add_heading("2. Understanding DTAP and ECNO VASP", level=1)

    doc.add_heading("What is DTAP?", level=2)
    doc.add_paragraph(
        "A Digital Technology Approval Process (DTAP) is the procedure by which an organization — typically a school board "
        "— authorizes new technology (software, apps, platforms) for use within its environment. A DTAP assesses "
        "security, privacy, and pedagogical/business alignment before granting approval."
    )
    doc.add_paragraph(
        "In Ontario, each of the 72 publicly funded school boards operates its own DTAP. While the specifics vary "
        "from board to board (different forms, timelines, review committees), they all evaluate the same core concerns: "
        "Does this tool protect student privacy? Is it secure? Does it serve a legitimate educational purpose? "
        "Does it comply with applicable legislation (MFIPPA, Education Act, AODA)?"
    )

    doc.add_heading("How DTAP and ECNO VASP Work Together", level=2)
    doc.add_paragraph(
        "Ontario has a multi-layered system where board-level DTAP processes are supported by provincial-level services. "
        "The most important provincial service is ECNO VASP (Vetting of Application Security & Privacy), which provides "
        "centralized security and privacy risk assessments that feed into each board's DTAP decision."
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, ["Process", "Scope", "Administered By", "Output"], header=True)
    table.rows[0]._element.getparent().remove(table.rows[0]._element)  # remove empty first row
    add_table_row(table, ["Board-level DTAP", "Individual board", "Each of 72 boards", "Approval/rejection for that board"])
    add_table_row(table, ["ECNO VASP", "Provincial (71+ boards)", "ECNO", "Risk score + confidential reports (feeds into DTAP)"])
    add_table_row(table, ["OSAPAC", "Provincial (all boards)", "Ministry advisory committee", "Province-wide license"])
    add_table_row(table, ["Privacy Impact Assessment", "Per-board (mandatory under EDSTA)", "Each board's privacy office", "PIA report (feeds into DTAP)"])

    doc.add_paragraph("")
    doc.add_paragraph(
        "A typical DTAP workflow: (1) A teacher or staff member requests a new tool. (2) The board's IT/curriculum team "
        "does an initial pedagogical and infrastructure review. (3) If the tool handles personal data, the board requests "
        "an ECNO VASP assessment. (4) The VASP risk score, along with the board's own PIA and internal review, informs "
        "the board's DTAP decision. (5) The tool is approved, approved with conditions, or rejected."
    )
    doc.add_paragraph(
        "For ClassBridge, the goal is to satisfy both: (a) the provincial-level ECNO VASP assessment with a favourable "
        "risk score, and (b) individual board DTAP requirements. A strong VASP score makes board-level DTAP approval "
        "significantly easier, as all 71 participating boards can reference the same VASP report."
    )

    doc.add_page_break()

    # ================================================================
    # 3. ECNO
    # ================================================================
    doc.add_heading("3. ECNO — The Central Organization", level=1)

    doc.add_heading("What is ECNO?", level=2)
    ecno_details = [
        ("Full Name", "Educational Computing Network of Ontario"),
        ("Website", "https://www.ecno.org"),
        ("Contact", "office@ecno.org | (519) 568-7899"),
        ("Structure", "Virtual organization (no physical office) providing technology leadership for Ontario K-12"),
        ("Members", "Ontario's 72 publicly funded school boards plus Indigenous education authorities"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    add_table_row(table, ["Field", "Details"], header=True)
    table.rows[0]._element.getparent().remove(table.rows[0]._element)
    for label, val in ecno_details:
        add_table_row(table, [label, val])

    doc.add_paragraph("")
    doc.add_paragraph(
        "ECNO is not a government body — it is a collaborative network owned and directed by its member school boards. "
        "It delivers shared IT services, negotiates licensing deals (e.g., their Value-Added Reseller relationship with "
        "Microsoft has saved boards over $6.5 million), and operates security and privacy services including VASP."
    )

    doc.add_heading("ECNO Services Relevant to EdTech Vendors", level=2)
    services = [
        "VASP (Vetting of Application Security & Privacy) — security/privacy risk assessments",
        "Student Digital Privacy Standard — the evaluation framework for EdTech tools",
        "Third-Party Risk Management (TPRM) — guidance for boards on vendor assessment",
        "K-12 Cyber Awareness — security education and resources",
        "Collaborative licensing negotiations (Microsoft, etc.)",
    ]
    for s in services:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_page_break()

    # ================================================================
    # 4. ECNO VASP
    # ================================================================
    doc.add_heading("4. ECNO VASP — The Provincial Vetting Process", level=1)

    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "VASP stands for Vetting of Application Security & Privacy. It is the closest thing Ontario has to a "
        "centralized, province-wide technology approval process for EdTech tools."
    )

    vasp_facts = [
        ("Launched", "Fall 2021"),
        ("Created by", "ECNO, in collaboration with OASBO Joint ICT-PIM-SCM committee"),
        ("Purpose", "Centrally provide professional risk assessments of security and privacy concerns for educational digital tools"),
        ("Participating boards", "71 of Ontario's 72 school boards"),
        ("Scale (as of 2024)", "Over 800 digital tools reviewed, plus 45 mobile applications"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    add_table_row(table, ["Attribute", "Detail"], header=True)
    table.rows[0]._element.getparent().remove(table.rows[0]._element)
    for label, val in vasp_facts:
        add_table_row(table, [label, val])

    doc.add_paragraph("")
    p = doc.add_paragraph("")
    run = p.add_run("Important: ")
    run.bold = True
    p.add_run(
        "VASP is NOT a certification, endorsement, or marketing tool. It is explicitly not an approval or "
        "disapproval of any given tool. It provides a risk assessment that boards use to make their own informed decisions."
    )

    doc.add_heading("The VASP Review Process", level=2)

    steps = [
        ("Step 1: Board-Level Request",
         "A teacher or staff member identifies a digital tool. They submit a request through their school board's "
         "internal Digital Tool Approval Process. The board conducts an internal assessment from a pedagogical lens."),
        ("Step 2: Board Submits to ECNO",
         "If the board determines the tool warrants further vetting, the board (not the vendor) submits a VASP request. "
         "Only ECNO member school boards can submit requests — vendors cannot apply directly."),
        ("Step 3: Privacy Analysis",
         "ECNO's Privacy Analyst reviews the vendor's privacy policies, terms of use, data processing agreements, "
         "and alignment with ECNO's Student Digital Privacy Standard."),
        ("Step 4: Security Analysis (Fuzzing)",
         "ECNO's Security Analyst performs automated and manual security testing (fuzzing) of the application to "
         "identify vulnerabilities."),
        ("Step 5: Vendor Engagement",
         "The vetting process generates questions and recommendations for the vendor. The VASP team connects with "
         "the provider to address questions and resolve issues."),
        ("Step 6: Report Generation",
         "Analysts produce a School Board Report (detailed risk assessment) and an Educator Report (practical guidance). "
         "Reports include a risk score and are confidential — for Ontario school board staff only."),
    ]
    for title, desc in steps:
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)

    doc.add_heading("Timeline", level=3)
    doc.add_paragraph(
        "ECNO does not publish expected turnaround times. Based on the nature of the process (privacy document review, "
        "security testing, vendor back-and-forth), it likely takes several weeks to months depending on complexity "
        "and vendor responsiveness."
    )

    doc.add_heading("VASP Team Structure", level=3)
    roles = [
        "Director of Security Services (program oversight)",
        "VASP Project Coordinator",
        "Privacy Analysts",
        "Security Analysts",
        "VASP Project Analysts",
    ]
    for r in roles:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_page_break()

    # ================================================================
    # 5. STUDENT DIGITAL PRIVACY STANDARD
    # ================================================================
    doc.add_heading("5. ECNO Student Digital Privacy Standard", level=1)

    doc.add_paragraph(
        "VASP assessments are based on ECNO's Student Digital Privacy Standard (published February 11, 2021). "
        "This standard aligns with data protections for children endorsed by regulators and experts across "
        "North America and the EU."
    )

    principles = [
        ("1. Specifying Purposes", "Providers must state all data elements collected and provide reasons for each."),
        ("2. Consent", "Schools must ensure verifiable parental consent for collection, use, and disclosure of personal information of children under 18, where no other legal basis exists."),
        ("3. Data Minimization", "Providers must collect only the personal information required to operate the service. No accessing browser history, contact lists, search terms, preferences, device identification, location, etc., unless directly related to the service."),
        ("4. Limiting Use, Disclosure & Retention", "Providers must not repurpose student data or use it for research without express consent. Must securely destroy or anonymize data no longer required. Must explicitly identify retention timelines."),
        ("5. Security Safeguards", "Provider must have a comprehensive security program. Must use administrative, technological, and physical safeguards. Must ensure all sub-processors implement the same safeguards."),
        ("6. Openness & Transparency", "Privacy notices, terms of use, and contracts must be in clear, specific, and unambiguous language."),
        ("7. Individual Access & Control", "Must provide mechanisms to access, correct, erase, and download content in a usable format. Right to erasure includes metadata, inferences, assessments, and profiles. No fee may be charged."),
        ("8. Accountability & Contact", "Must provide name and contact information of an operator who responds to privacy inquiries."),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    add_table_row(table, ["Principle", "Requirement"], header=True)
    table.rows[0]._element.getparent().remove(table.rows[0]._element)
    for title, desc in principles:
        add_table_row(table, [title, desc])

    doc.add_paragraph("")
    doc.add_heading("Additional Requirements", level=2)
    additional = [
        "Breach notification protocols must be in place",
        "No use of student data for targeted advertising",
        "Successor entities must implement the same safeguards for previously collected data",
        "Must directly inform users before privacy policy / ToS changes",
        "Must disclose third-party cookies and provide management options",
    ]
    for a in additional:
        doc.add_paragraph(a, style='List Bullet')

    doc.add_page_break()

    # ================================================================
    # 6. LEGAL & REGULATORY FRAMEWORK
    # ================================================================
    doc.add_heading("6. Legal & Regulatory Framework", level=1)

    doc.add_heading("6.1 MFIPPA (Municipal Freedom of Information and Protection of Privacy Act)", level=2)
    doc.add_paragraph(
        "MFIPPA is the primary privacy law governing Ontario school boards. It sets rules for collection, use, "
        "retention, and disclosure of personal information. School boards remain accountable for personal information "
        "even when handled by third parties."
    )

    doc.add_heading("Key MFIPPA Requirements", level=3)
    mfippa = [
        "Collection only if expressly authorized by law or necessary for a lawfully authorized activity (s.28(2))",
        "Notice at time of collection: legal authority, purpose, contact person (s.29(2))",
        "Data minimization: collect only what is necessary (s.30)",
        "Purpose limitation: use only for the purpose collected (s.31)",
        "Disclosure restrictions: defined circumstances only — consent, law enforcement, health/safety (s.32)",
        "Individual right to access their personal information (Part I)",
        "Right to request correction of inaccurate information (s.36)",
    ]
    for m in mfippa:
        doc.add_paragraph(m, style='List Bullet')

    doc.add_heading("6.2 Consent Requirements for Minors", level=2)
    consent_table = [
        ("Under 16", "Parent/guardian provides consent on child's behalf"),
        ("Age 16-17", "Both the student AND parent/guardian consent"),
        ("18+", "Student alone provides consent"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    add_table_row(table, ["Age Group", "Consent Requirement"], header=True)
    table.rows[0]._element.getparent().remove(table.rows[0]._element)
    for age, req in consent_table:
        add_table_row(table, [age, req])

    doc.add_paragraph("")
    doc.add_paragraph("Consent must be written, with the specific information identified.")

    doc.add_heading("6.3 Cross-Border Data Transfer", level=2)
    doc.add_paragraph(
        "MFIPPA does not contain an explicit prohibition on cross-border transfer, unlike B.C. FIPPA and N.S. PIIDPA. "
        "However, Ontario guidance strongly states that cloud data should preferably be stored in a Canadian multi-zone "
        "region under the control of Canadian nationals. In practice, boards will almost universally require Canadian "
        "data residency as a contractual term."
    )

    doc.add_heading("6.4 Bill 194 / EDSTA (Effective 2025)", level=2)
    doc.add_paragraph(
        "The Enhancing Digital Security and Trust Act (EDSTA), passed as Bill 194, received Royal Assent on "
        "November 25, 2024 and significantly raises the bar for privacy and cybersecurity."
    )

    edsta_items = [
        "Mandatory Privacy Impact Assessments (PIAs) before collecting personal information",
        "Mandatory breach notification to IPC and affected individuals (real risk of significant harm standard)",
        "Enhanced cybersecurity measures: encryption, identity/access controls, patch management",
        "AI governance: regulation of AI systems used by public institutions (specifics pending)",
        "Stronger children's data protections for school boards",
        "Third-party accountability: extends institutional responsibility to data handled by vendors",
        "Annual breach reporting to IPC starting 2026",
    ]
    for item in edsta_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("6.5 Ontario IPC Digital Privacy Charter (October 2024)", level=2)
    doc.add_paragraph(
        "The Information and Privacy Commissioner of Ontario launched the Digital Privacy Charter for Ontario "
        "Schools in October 2024 with 12 commitments for schools, including explaining types of personal "
        "information collected, informing students/parents of significant breaches in a timely way, educating "
        "students about online risks, and empowering students to exercise their privacy rights. Schools are being "
        "urged to sign the pledge before the 2025-26 school year."
    )

    doc.add_heading("6.6 Joint Resolution of Canada's Privacy Commissioners (October 2025)", level=2)
    doc.add_paragraph(
        "A unanimous joint resolution by all of Canada's information and privacy regulators was adopted in "
        "Banff, Alberta on October 8, 2025, calling for:"
    )
    joint_items = [
        "Embedding privacy by design into products and services",
        "Avoiding design practices that manipulate or coerce users",
        "Building in appropriate access controls and encryption",
        "Establishing privacy settings to most protective level by default",
        "Prioritizing privacy protection when selecting educational technologies",
        "Funding digital education and privacy training",
    ]
    for item in joint_items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph(
        "This joint resolution signals increased regulatory scrutiny of EdTech and reinforces the need for "
        "privacy-by-design architecture in platforms like ClassBridge."
    )

    doc.add_heading("6.7 Privacy Impact Assessment (PIA) Requirements", level=2)
    doc.add_paragraph("Under Bill 194, PIAs must cover the following mandatory elements:")
    pia_items = [
        "Purpose and necessity of the personal information collection",
        "Legal authority for the collection",
        "Types and sources of personal information being collected",
        "Intended uses and disclosures of the information",
        "Access roles — who will have access",
        "Limits and restrictions on access permissions",
        "Retention period for the information",
        "Safeguards and risks, including summary of risks if breach occurs",
        "Specific steps to prevent and mitigate those risks",
        "Other prescribed information (as defined by future regulations)",
    ]
    for i, item in enumerate(pia_items, 1):
        doc.add_paragraph(f"{i}. {item}")

    doc.add_page_break()

    # ================================================================
    # 7. TECHNICAL REQUIREMENTS
    # ================================================================
    doc.add_heading("7. Technical Requirements for EdTech Vendors", level=1)

    doc.add_heading("7.1 Encryption", level=2)
    enc_table = [
        ("In Transit", "TLS 1.2 minimum, TLS 1.3 preferred", "Required"),
        ("At Rest", "AES-256 encryption for stored personal information", "Required"),
        ("In Processing", "Encryption during cloud processing", "Increasingly expected"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    add_table_row(table, ["Type", "Standard", "Status"], header=True)
    table.rows[0]._element.getparent().remove(table.rows[0]._element)
    for row in enc_table:
        add_table_row(table, row)

    doc.add_heading("7.2 Authentication & Access Control", level=2)
    auth_items = [
        "Multi-Factor Authentication (MFA): Not explicitly mandated but expected under Bill 194's 'identity and access controls' requirement. Standard expectation in vendor questionnaires.",
        "SSO Integration: SAML 2.0, OAuth 2.0, or OpenID Connect strongly preferred. Common platforms: Google, Microsoft, Clever, ClassLink.",
        "Role-Based Access Control (RBAC): PIA requirement explicitly asks for access roles and restrictions.",
        "Account Lockout: Expected to prevent brute-force attacks.",
    ]
    for item in auth_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("7.3 Security Certifications", level=2)
    cert_table = [
        ("SOC 2 Type II", "De facto industry standard for EdTech. Not legally mandated but increasingly demanded by boards.", "High priority"),
        ("ISO 27001", "Not required. ~80% overlap with SOC 2. Consider after SOC 2.", "Lower priority"),
        ("Penetration Testing", "Annual testing expected. VASP performs its own security fuzzing.", "Required"),
        ("Vulnerability Scanning", "Automated scanning in CI/CD pipeline. Regular third-party scans.", "Required"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    add_table_row(table, ["Certification", "Details", "Priority"], header=True)
    table.rows[0]._element.getparent().remove(table.rows[0]._element)
    for row in cert_table:
        add_table_row(table, row)

    doc.add_heading("7.4 Data Residency", level=2)
    doc.add_paragraph(
        "While MFIPPA does not explicitly mandate Canadian data residency, boards will almost universally require it. "
        "GCP's northamerica-northeast2 (Toronto) region is ideal for Ontario-focused applications. "
        "All backup/DR must also use Canadian regions. Sub-processors (e.g., OpenAI) must also be assessed for "
        "data location."
    )

    doc.add_heading("7.5 Accessibility (AODA / WCAG)", level=2)
    doc.add_paragraph(
        "The Accessibility for Ontarians with Disabilities Act (AODA) requires WCAG 2.0 Level AA compliance. "
        "Best practice is to target WCAG 2.1 AA. Key requirements include:"
    )
    accessibility_items = [
        "Keyboard navigation for all interactive elements, no keyboard traps",
        "Screen reader compatibility with proper semantic HTML and ARIA labels",
        "Text alternatives (alt text) for all non-text content",
        "Color contrast: minimum 4.5:1 for normal text, 3:1 for large text",
        "Captions for multimedia, transcripts for audio",
        "Content usable at 200% zoom without horizontal scrolling",
        "Visible focus indicators on all interactive elements",
        "Form labels and clear error identification",
        "No timing dependencies (users can extend or disable time limits)",
    ]
    for item in accessibility_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("7.6 Audit Logging", level=2)
    doc.add_paragraph(
        "SOC 2 and PIAs require comprehensive audit logging. Boards expect audit trails of data access, "
        "modifications, and exports. Logs should include: who accessed what, when, from where (IP), "
        "and what action was taken."
    )

    doc.add_heading("7.7 Backup & Disaster Recovery", level=2)
    doc.add_paragraph(
        "Documented backup procedures and disaster recovery plans with defined RPO/RTO are expected. "
        "Boards will ask about backup frequency, retention, encryption, and recovery testing."
    )

    doc.add_heading("7.8 Insurance Requirements", level=2)
    ins_table = [
        ("Commercial General Liability", "$2M - $5M"),
        ("Cyber Liability / Privacy & Network Security", "$5M - $10M"),
        ("Technology Errors & Omissions", "$2M - $5M"),
        ("Professional Liability", "$1M - $2M"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    add_table_row(table, ["Coverage Type", "Typical Minimum"], header=True)
    table.rows[0]._element.getparent().remove(table.rows[0]._element)
    for row in ins_table:
        add_table_row(table, row)

    doc.add_paragraph("")
    doc.add_paragraph(
        "MEARIE (Ontario's public sector reciprocal insurer) recommends a minimum $10M limit for Privacy, "
        "Liability, and Network Security Liability. Vendor's policy should name the board as additional insured."
    )

    doc.add_page_break()

    # ================================================================
    # 8. OSAPAC
    # ================================================================
    doc.add_heading("8. OSAPAC — Provincial Software Licensing", level=1)
    doc.add_paragraph(
        "OSAPAC (Ontario Software Acquisition Program Advisory Committee) advises the Ministry of Education "
        "on acquiring province-wide software licenses. It is separate from VASP — OSAPAC handles licensing and "
        "procurement, while VASP handles security/privacy vetting."
    )
    doc.add_heading("OSAPAC Process", level=2)
    osapac_steps = [
        "Province-wide survey to determine software license priorities",
        "Priorities confirmed with Ministry's Curriculum and Assessment Policy Branch",
        "Requests for Qualification (RFQ) posted on Ontario electronic bidding system",
        "Qualified respondents submit pricing through Request for Tender (RFT)",
        "Highest-scoring qualified product wins (if within budget)",
        "Ministry funds the provincial license; downloads via ECNO portal",
    ]
    for i, step in enumerate(osapac_steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_paragraph(
        "OSAPAC is a longer-term consideration for ClassBridge. Focus on VASP approval and individual board "
        "adoption first."
    )

    doc.add_page_break()

    # ================================================================
    # 9. BOARD vs PROVINCIAL
    # ================================================================
    doc.add_heading("9. Board-Level vs Provincial-Level Approval", level=1)
    doc.add_paragraph(
        "Ontario uses a multi-layered system, not a single centralized approval. VASP reports are produced centrally "
        "by ECNO and available to all participating member boards. However, each board still makes its own independent "
        "decision about whether to approve a tool based on the VASP report and their own internal criteria."
    )
    doc.add_paragraph(
        "A VASP risk score does not automatically approve a tool at any board — it informs the board's decision. "
        "Each board has its own internal digital tool approval process with varying timelines and requirements."
    )

    doc.add_heading("Strategy for ClassBridge", level=2)
    strategy = [
        "Identify 1-2 pilot school boards willing to champion ClassBridge",
        "Work with the pilot board to submit a VASP request to ECNO",
        "Prepare all documentation to VASP standards before submission",
        "A positive VASP assessment will facilitate adoption by additional boards",
    ]
    for s in strategy:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_page_break()

    # ================================================================
    # 10. CLASSBRIDGE CURRENT ASSESSMENT
    # ================================================================
    doc.add_heading("10. ClassBridge Current Architecture Assessment", level=1)

    doc.add_heading("10.1 What ClassBridge Has (Strengths)", level=2)
    strengths = [
        "JWT-based authentication with access/refresh tokens, bcrypt password hashing, strong password requirements",
        "RBAC with four roles (PARENT, STUDENT, TEACHER, ADMIN) and role-based route protection",
        "Login rate limiting (5 attempts/minute) with failed login audit logging",
        "Comprehensive security headers (HSTS, CSP, X-Frame-Options, X-Content-Type-Options)",
        "CORS restricted to configured origins (not wildcard in production)",
        "Pydantic input validation with typed schemas",
        "File upload validation (size limits, extension whitelisting)",
        "Audit logging (AuditLog model with user, action, IP, user agent)",
        "Privacy Policy and Terms of Service pages with contact information",
        "Cloud SQL with automatic daily backups and point-in-time recovery",
        "CI/CD with test gates before deployment",
        "Documented disaster recovery plan (RPO <24h, RTO <1h)",
        "Account deletion with 30-day timeline",
        "Google OAuth integration with scope tracking",
        "Email verification flow",
    ]
    for s in strengths:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("10.2 Critical Gaps", level=2)

    # Gap table
    gap_categories = [
        ("Data Residency", [
            "Currently hosted in us-central1 (Iowa, USA) — must migrate to Canadian region",
            "OpenAI API calls send student data to US servers — needs Canadian alternative or documentation",
        ]),
        ("Privacy & Legal", [
            "No MFIPPA/PIPEDA-compliant privacy documentation",
            "No Data Processing Agreement (DPA) template for school boards",
            "No Privacy Impact Assessment (PIA) document",
            "No age-based consent mechanism (MFIPPA requires different consent for <16, 16-17, 18+)",
            "Privacy policy references COPPA but not Canadian legislation",
            "No cookie/consent management",
            "No data export in usable format (right to data portability)",
        ]),
        ("Authentication & Security", [
            "No MFA/2FA support",
            "No SSO/SAML support for district-level integration",
            "JWT stored in localStorage (XSS vulnerability) — should use httpOnly cookies",
            "No account lockout after failed attempts",
            "No SOC 2 Type II certification",
            "No penetration testing program",
            "No dependency vulnerability scanning in CI/CD",
            "No SAST/DAST in pipeline",
        ]),
        ("Accessibility", [
            "No WCAG 2.1 AA compliance audit",
            "Minimal ARIA implementation",
            "No skip-to-content links",
            "No documented accessibility testing",
            "No color contrast verification",
        ]),
        ("Infrastructure", [
            "No high availability (single instance)",
            "No database replication/failover",
            "No WAF or DDoS protection",
            "No infrastructure-as-code",
        ]),
    ]

    for category, gaps in gap_categories:
        doc.add_heading(category, level=3)
        for gap in gaps:
            doc.add_paragraph(gap, style='List Bullet')

    doc.add_page_break()

    # ================================================================
    # 11. COMPLIANCE GAP ANALYSIS
    # ================================================================
    doc.add_heading("11. Compliance Gap Analysis", level=1)
    doc.add_paragraph(
        "The following table maps each ECNO Student Digital Privacy Standard principle and technical requirement "
        "to ClassBridge's current state and required actions."
    )

    gap_analysis = [
        ("Specifying Purposes", "Partial", "Privacy policy exists but needs MFIPPA-specific language. Must itemize all data elements collected."),
        ("Consent (age-based)", "Not Met", "No age verification or tiered consent mechanism. Must implement MFIPPA-compliant consent for <16, 16-17, 18+ age groups."),
        ("Data Minimization", "Partial", "Collects appropriate data but no formal data inventory. Need to document and justify each data element."),
        ("Use/Disclosure/Retention Limits", "Partial", "30-day deletion exists. Need formal retention schedule, data destruction procedures, and sub-processor disclosure."),
        ("Security Safeguards", "Partial", "Good baseline (encryption, rate limiting, headers). Need MFA, pen testing, SOC 2, vulnerability scanning."),
        ("Openness & Transparency", "Partial", "Privacy policy and ToS exist. Need plain-language updates aligned with MFIPPA and Student Digital Privacy Standard."),
        ("Individual Access & Control", "Not Met", "No data export, no self-service data erasure, no profile download mechanism."),
        ("Accountability & Contact", "Met", "Privacy officer contact (privacy@classbridge.ca) documented in privacy policy."),
        ("Canadian Data Residency", "Not Met", "Currently US-hosted. Must migrate to GCP northamerica-northeast2 (Toronto)."),
        ("MFA / SSO", "Not Met", "No MFA, no SAML SSO. Must implement for admin/teacher accounts at minimum."),
        ("WCAG 2.1 AA", "Not Met", "Minimal ARIA. Full accessibility audit and remediation required."),
        ("Breach Notification", "Partial", "Incident response doc exists. Need formal MFIPPA-compliant notification procedures."),
        ("SOC 2 Type II", "Not Started", "6-12 month process. Begin with readiness assessment."),
        ("Penetration Testing", "Not Started", "Annual pen testing required. Schedule first engagement."),
        ("PIA Document", "Not Started", "Must prepare using IPC template aligned with Bill 194 requirements."),
        ("DPA Template", "Not Started", "Must create school-board-ready Data Processing Agreement."),
        ("Cyber Insurance", "Not Started", "Obtain $2-5M minimum coverage."),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    add_table_row(table, ["Requirement", "Status", "Action Required"], header=True)
    table.rows[0]._element.getparent().remove(table.rows[0]._element)
    for row in gap_analysis:
        r = add_table_row(table, row)
        # Color code status
        status_cell = r.cells[1]
        status = row[1]
        if status == "Met":
            set_cell_shading(status_cell, "C6EFCE")
        elif status == "Partial":
            set_cell_shading(status_cell, "FFEB9C")
        elif status in ("Not Met", "Not Started"):
            set_cell_shading(status_cell, "FFC7CE")

    doc.add_page_break()

    # ================================================================
    # 12. REQUIRED ARCHITECTURE & DESIGN CHANGES
    # ================================================================
    doc.add_heading("12. Required Architecture & Design Changes", level=1)

    doc.add_paragraph(
        "This section details the specific architectural and design changes ClassBridge must implement "
        "to achieve DTAP/VASP compliance. These changes affect both the backend infrastructure and "
        "the frontend application."
    )

    doc.add_heading("12.1 Infrastructure Migration (Canadian Data Sovereignty)", level=2)
    infra_changes = [
        ("Cloud Run / Railway", "Migrate compute from us-central1 to Canadian region (GCP northamerica-northeast2 Toronto or Railway Montreal). All application instances must run in Canada."),
        ("Cloud SQL / PostgreSQL", "Migrate database to Canadian region. Enable Customer-Managed Encryption Keys (CMEK) for AES-256 at-rest encryption. Ensure PITR and backups also reside in Canadian region."),
        ("GCS / Object Storage", "Move all file storage buckets (course materials, exports, backups) to Canadian region."),
        ("CI/CD Pipeline", "Update deployment configs (.github/workflows/deploy.yml) to target Canadian region. Update environment variables and secrets."),
        ("DNS & CDN", "Ensure CDN edge caching does not store PII outside Canada. Consider Cloudflare Canada-only configuration."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    add_table_row(table, ["Component", "Change Required"], header=True)
    table.rows[0]._element.getparent().remove(table.rows[0]._element)
    for comp, change in infra_changes:
        add_table_row(table, [comp, change])

    doc.add_heading("12.2 Authentication Architecture Overhaul", level=2)
    auth_changes = [
        ("JWT Storage", "Move from localStorage to httpOnly, Secure, SameSite=Strict cookies. Update Axios interceptor to use cookie-based auth instead of Authorization header. Add CSRF protection token."),
        ("MFA/2FA Layer", "Add TOTP-based MFA (RFC 6238). New database tables: mfa_settings (user_id, secret, enabled, backup_codes). New API routes: /api/auth/mfa/setup, /api/auth/mfa/verify, /api/auth/mfa/backup-codes. Update JWT claims to include mfa_verified boolean."),
        ("SSO/SAML Integration", "Add SAML 2.0 Service Provider (SP) capability using python-saml3. Support Google Workspace, Microsoft 365, Clever, ClassLink IdPs. New routes: /api/auth/saml/metadata, /api/auth/saml/login, /api/auth/saml/acs. Add tenant/organization model for multi-board deployments."),
        ("Account Lockout", "Add failed_login_attempts and locked_until columns to User model. Implement progressive lockout: 5 failures → 15min, 10 → 1hr, 15 → 24hr. Add admin unlock capability."),
        ("Session Management", "Add Session model tracking device, IP, last_activity. Implement concurrent session limits. Add /api/auth/sessions endpoint for users to view/revoke active sessions."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    add_table_row(table, ["Component", "Change Required"], header=True)
    table.rows[0]._element.getparent().remove(table.rows[0]._element)
    for comp, change in auth_changes:
        add_table_row(table, [comp, change])

    doc.add_heading("12.3 Privacy & Consent Architecture", level=2)
    privacy_changes = [
        ("Age-Based Consent", "Add date_of_birth field to Student model. Implement consent workflow engine: under-16 requires parent consent only, 16-17 requires dual consent (parent + student), 18+ student-only. New Consent model tracking: user_id, consent_type, granted_at, granted_by, version."),
        ("Data Export API", "New endpoint: GET /api/users/me/export returning ZIP containing JSON files for all user data (profile, messages, study guides, grades, audit logs). Must include metadata, inferences, and AI-generated content."),
        ("Right to Erasure", "New endpoint: DELETE /api/users/me with 30-day grace period. Cascading deletion of all related data: messages, study guides, quiz results, flashcards, notifications, audit logs (after retention period). Anonymize rather than delete where referential integrity requires."),
        ("Cookie Consent", "Add cookie consent banner component. Implement consent categories: essential, analytics, functional. Store consent preferences. Block non-essential cookies until consent granted."),
        ("Breach Notification", "Add BreachIncident model. Implement notification workflow: detect → assess severity → notify IPC within 72 hours → notify affected users. Integrate with email service for mass notification."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    add_table_row(table, ["Component", "Change Required"], header=True)
    table.rows[0]._element.getparent().remove(table.rows[0]._element)
    for comp, change in privacy_changes:
        add_table_row(table, [comp, change])

    doc.add_heading("12.4 AI Service Architecture (OpenAI Data Residency)", level=2)
    ai_changes = [
        ("Option A: Azure OpenAI (Recommended)", "Migrate from OpenAI API to Azure OpenAI Service deployed in Canada East (Montreal) region. This keeps data within Canadian borders. Requires Azure subscription and model deployment. API is compatible — minimal code changes in app/services/study_service.py."),
        ("Option B: Data Anonymization Layer", "Add a pre-processing layer that strips PII before sending to OpenAI. Replace student names, school names, teacher names with anonymized tokens. Reconstruct after response. More complex but allows continued use of OpenAI API."),
        ("Option C: Document & Opt-Out", "Document the US data transfer in the PIA with contractual safeguards (OpenAI DPA). Offer boards an opt-out for AI features. Least disruptive but may face pushback from privacy-conscious boards."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    add_table_row(table, ["Option", "Description"], header=True)
    table.rows[0]._element.getparent().remove(table.rows[0]._element)
    for opt, desc in ai_changes:
        add_table_row(table, [opt, desc])

    doc.add_heading("12.5 Frontend Accessibility Overhaul (WCAG 2.1 AA)", level=2)
    a11y_changes = [
        "Add skip-to-content link on every page",
        "Ensure all interactive elements are keyboard-accessible with visible focus indicators",
        "Add ARIA landmarks (banner, navigation, main, contentinfo) to layout components",
        "Add aria-label, aria-describedby, and aria-live attributes to dynamic content",
        "Ensure all form inputs have associated <label> elements",
        "Verify color contrast ratios: 4.5:1 for normal text, 3:1 for large text",
        "Add alt text to all images; decorative images get empty alt=''",
        "Ensure modals trap focus and return focus on close",
        "Add error announcements to screen readers on form validation",
        "Support 200% text zoom without horizontal scrolling",
        "Add prefers-reduced-motion media query support",
    ]
    for item in a11y_changes:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("12.6 Security Pipeline & Monitoring", level=2)
    sec_changes = [
        ("WAF/DDoS Protection", "Deploy Cloud Armor or equivalent WAF in front of application. Configure OWASP Core Rule Set for SQLi, XSS, path traversal blocking. Add geographic IP filtering capability (Canada-only option for board deployments)."),
        ("SAST in CI/CD", "Add CodeQL (Python + TypeScript) and Bandit (Python security linter) to GitHub Actions. Block PRs on high/critical findings. Add GitLeaks for secret scanning."),
        ("DAST in CI/CD", "Add OWASP ZAP automated scans against staging after each deploy. Archive reports for SOC 2 evidence."),
        ("Container Scanning", "Add Trivy to scan Docker images in CI. Block deploys on critical CVEs."),
        ("Dependency Scanning", "Enable GitHub Dependabot for Python (pip) and JavaScript (npm). Add pip-audit and npm audit to CI pipeline."),
        ("Comprehensive Rate Limiting", "Extend slowapi rate limiting to all API endpoints. Tier by endpoint type: auth (5-10/min), AI generation (10-20/min), data reads (60-120/min), data writes (30-60/min). Rate limit by user ID for authenticated requests, by IP for unauthenticated."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    add_table_row(table, ["Component", "Change Required"], header=True)
    table.rows[0]._element.getparent().remove(table.rows[0]._element)
    for comp, change in sec_changes:
        add_table_row(table, [comp, change])

    doc.add_heading("12.7 Enhanced Audit Logging", level=2)
    audit_changes = [
        "Log all data READ actions (who viewed which student's data, when, from where)",
        "Log data EXPORT actions with export scope and format",
        "Log all ADMIN actions (role changes, user management, config changes)",
        "Add structured JSON schema for audit details field",
        "Implement audit log export for SOC 2 evidence collection",
        "Add real-time alerting on suspicious patterns (e.g., bulk data access, privilege escalation)",
        "Ensure audit logs are tamper-resistant (append-only, signed)",
    ]
    for item in audit_changes:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ================================================================
    # 13. LIKELIHOOD OF APPROVAL
    # ================================================================
    doc.add_heading("13. Likelihood of ClassBridge Passing DTAP / Achieving Favourable VASP Score", level=1)

    doc.add_heading("Current Likelihood: Low-Medium (35-40% ready)", level=2)
    doc.add_paragraph(
        "ClassBridge has a solid technical foundation but is not yet ready to pass a board's DTAP or receive "
        "a favourable VASP risk score. The critical blockers are:"
    )
    blockers = [
        "US data residency (most significant blocker — boards will not approve US-hosted student data)",
        "No MFIPPA-compliant privacy documentation or DPA",
        "No MFA/SSO support",
        "No accessibility compliance (AODA is legally binding)",
        "OpenAI data flows to US servers (AI feature is a core differentiator but also a risk)",
    ]
    for b in blockers:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading("Likelihood Progression by Phase", level=2)
    likelihood_data = [
        ("Today (as-is)", "~0%", "US data residency is an automatic disqualifier for Ontario boards"),
        ("After Phase 1 (3 months)", "~40-50%", "Data in Canada + privacy docs, but still missing security certs and MFA"),
        ("After Phase 2 (6 months)", "~60-70%", "MFA/SSO + security scanning adds credibility and meets Bill 194 requirements"),
        ("After Phase 3 (9 months)", "~80-85%", "WCAG compliance + full documentation suite meets AODA and ECNO standards"),
        ("After Phase 4 (12 months)", "~90%+", "SOC 2 + pen test + board partner = strong VASP submission"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    add_table_row(table, ["Timeframe", "Likelihood", "Condition"], header=True)
    table.rows[0]._element.getparent().remove(table.rows[0]._element)
    for row in likelihood_data:
        r = add_table_row(table, row)
        # Color code likelihood
        pct = row[1]
        cell = r.cells[1]
        if "90" in pct:
            set_cell_shading(cell, "C6EFCE")
        elif "80" in pct:
            set_cell_shading(cell, "D4EDBC")
        elif "60" in pct or "70" in pct:
            set_cell_shading(cell, "FFEB9C")
        elif "40" in pct or "50" in pct:
            set_cell_shading(cell, "FFD699")
        else:
            set_cell_shading(cell, "FFC7CE")

    doc.add_paragraph("")

    doc.add_heading("After Remediation: High (80-85% ready)", level=2)
    doc.add_paragraph(
        "With the recommended roadmap completed (6-12 months), ClassBridge would have a strong likelihood of "
        "passing board-level DTAP reviews and receiving a favourable VASP risk score. The platform's core "
        "architecture is sound, the tech stack is modern, and the existing security measures provide a good "
        "foundation to build upon."
    )

    doc.add_heading("Competitive Advantages", level=2)
    advantages = [
        "AI-powered study tools are a strong differentiator in the Ontario market",
        "Google Classroom integration aligns with boards still using Google Workspace",
        "Parent engagement focus addresses a real gap in the market",
        "Modern tech stack (FastAPI, React 19) enables rapid compliance improvements",
        "Existing RBAC and audit logging provide a head start on access control requirements",
    ]
    for a in advantages:
        doc.add_paragraph(a, style='List Bullet')

    doc.add_heading("Key Risk: OpenAI Data Flows", level=2)
    doc.add_paragraph(
        "ClassBridge's AI features (study guides, quizzes, flashcards) send course content to OpenAI's API, "
        "which may process data on US servers. This will be flagged in any VASP assessment. Options include:"
    )
    options = [
        "Use Azure OpenAI Service with a Canadian region deployment",
        "Implement data anonymization before sending to OpenAI (strip PII, use anonymized identifiers)",
        "Document the data flow explicitly in the PIA with safeguards (contractual, technical)",
        "Offer an opt-out for AI features in board deployments",
    ]
    for i, opt in enumerate(options, 1):
        doc.add_paragraph(f"{i}. {opt}")

    doc.add_page_break()

    # ================================================================
    # 14. RECOMMENDED ROADMAP & GITHUB ISSUES
    # ================================================================
    doc.add_heading("14. Recommended Roadmap & GitHub Issues", level=1)

    doc.add_paragraph(
        "All work items are tracked as GitHub issues under the DTAP Compliance epic (#803). "
        "Issues are labeled 'compliance' and organized into four phases."
    )

    doc.add_heading("Phase 1: Critical Foundation (Months 1-3)", level=2)
    phase1 = [
        "#779 — Migrate infrastructure to GCP Canada region (northamerica-northeast2)",
        "#780 — Address OpenAI API data residency (US data transfer risk)",
        "#781 — Update Privacy Policy & ToS for MFIPPA/PIPEDA compliance",
        "#782 — Create Privacy Impact Assessment (PIA) document",
        "#783 — Implement age-based consent mechanism (MFIPPA)",
        "#788 — Move JWT to httpOnly cookies (XSS mitigation)",
    ]
    for item in phase1:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Phase 2: Security Hardening (Months 3-6)", level=2)
    phase2 = [
        "#784 — Implement MFA/2FA support",
        "#785 — Implement SSO/SAML support for school board integration",
        "#789 — Add dependency vulnerability scanning to CI/CD",
        "#790 — Establish annual penetration testing program",
        "#796 — Implement account lockout and enhanced brute-force protection",
        "#800 — Enhance audit logging for SOC 2 and VASP requirements",
        "#804 — Add WAF/DDoS protection (Cloud Armor)",
        "#805 — Add SAST/DAST security scanning to CI/CD",
        "#807 — Implement comprehensive API rate limiting",
    ]
    for item in phase2:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Phase 3: Accessibility & Compliance (Months 6-9)", level=2)
    phase3 = [
        "#786 — WCAG 2.1 AA accessibility audit and remediation",
        "#787 — Implement data export and right to erasure",
        "#793 — Create Data Processing Agreement (DPA) template for school boards",
        "#794 — Implement formal breach notification procedures",
        "#795 — Complete K-12CVAT vendor questionnaire",
        "#797 — Implement consent management and cookie disclosure",
        "#798 — Implement formal data retention and automated purging",
        "#799 — Designate privacy officer and establish privacy program",
        "#806 — Create formal data classification and inventory",
    ]
    for item in phase3:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Phase 4: Certification & Engagement (Months 9-12)", level=2)
    phase4 = [
        "#791 — Begin SOC 2 Type II readiness assessment",
        "#792 — Obtain cyber liability insurance",
        "#801 — Evaluate D2L Brightspace integration (Ontario VLE)",
        "#802 — Identify and engage pilot school board partner",
    ]
    for item in phase4:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ================================================================
    # 15. COST & TIMELINE
    # ================================================================
    doc.add_heading("15. Cost & Timeline Estimates", level=1)

    doc.add_heading("Direct Costs", level=2)
    cost_items = [
        ("GCP region migration (infra changes)", "$500 - $2,000", "Month 1-2"),
        ("Penetration testing (annual)", "$5,000 - $20,000", "Month 4-5"),
        ("SOC 2 Type II audit", "$20,000 - $80,000", "Month 6-12"),
        ("Accessibility audit (third-party)", "$5,000 - $15,000", "Month 6-7"),
        ("Cyber liability insurance (annual)", "$3,000 - $10,000", "Month 3"),
        ("Legal review (privacy docs, DPA)", "$5,000 - $15,000", "Month 1-3"),
        ("Azure OpenAI Service (if migrating)", "$200 - $1,000/month", "Month 2+"),
        ("Security scanning tools (Snyk/similar)", "$0 - $500/month", "Month 3+"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    add_table_row(table, ["Item", "Estimated Cost", "Timeline"], header=True)
    table.rows[0]._element.getparent().remove(table.rows[0]._element)
    for row in cost_items:
        add_table_row(table, row)

    doc.add_paragraph("")
    p = doc.add_paragraph("")
    run = p.add_run("Total estimated first-year cost: $40,000 - $145,000")
    run.bold = True
    doc.add_paragraph(
        "Note: Engineering effort for implementation is additional. "
        "The VASP assessment itself has no direct fee — it is a service provided to ECNO member boards."
    )

    doc.add_heading("No Direct Application Fee", level=2)
    doc.add_paragraph(
        "There is no fee to apply for VASP assessment. The vendor does not apply or pay — a school board submits "
        "the request. VASP is included as part of ECNO membership benefits for school boards."
    )

    doc.add_page_break()

    # ================================================================
    # 16. SOURCES
    # ================================================================
    doc.add_heading("16. Sources & References", level=1)

    sources = [
        ("ECNO — About", "https://ecno.org/about/about-ecno/"),
        ("ECNO VASP Service", "https://www.ecno.org/services-programs/vasp/"),
        ("ECNO Student Digital Privacy Standard", "https://www.ecno.org/student-digital-privacy-standard/"),
        ("ECNO VASP Overview PDF", "https://www.ecno.org/wp-content/uploads/2024/10/ECNO-Vetting-of-Application-Security-Privacy-VASP.pdf"),
        ("ECNO Third-Party Risk Management", "https://www.k-12cyberawareon.ca/navigating-the-digital-frontier"),
        ("Ontario IPC — Privacy Legislation", "https://www.ipc.on.ca/en/education/ontarios-access-and-privacy-legislation"),
        ("Ontario IPC — Consent Guidance", "https://www.ipc.on.ca/en/education/consent-to-collect-use-and-disclose-personal-information"),
        ("Ontario IPC — PIA Guide", "https://www.ipc.on.ca/en/resources/planning-success-privacy-impact-assessment-guide-ontarios-public-institutions"),
        ("Ontario IPC — Digital Privacy Charter", "https://www.ipc.on.ca/en/privacy-organizations/digital-privacy-charter-for-ontario-schools"),
        ("Ontario IPC — Third-Party Contracting", "https://www.ipc.on.ca/en/resources/privacy-and-access-public-sector-contracting-third-party-service-providers"),
        ("Ontario IPC — EdTech in Schools Blog", "https://www.ipc.on.ca/en/media-centre/blog/use-edtech-schools-children-should-not-have-swap-their-privacy-education"),
        ("Bill 194 — Legislative Assembly", "https://www.ola.org/en/legislative-business/bills/parliament-43/session-1/bill-194"),
        ("Bill 194 — IPC Overview", "https://www.ipc.on.ca/en/resources/bill-194-strengthening-cyber-security-and-building-trust-public-sector-act"),
        ("Bill 194 — Fasken Analysis", "https://www.fasken.com/en/knowledge/2024/12/ontarios-public-sector-cyber-security-legislation-receives-royal-assent"),
        ("Bill 194 — Dentons Analysis", "https://www.dentonsdata.com/ontarios-new-public-sector-cybersecurity-and-ai-law-now-in-force-what-public-and-private-sector-organizations-need-to-know/"),
        ("OSAPAC — Mandate", "https://www.osapac.ca/about-us/mandate/"),
        ("OSAPAC — Vendor Info", "https://www.osapac.ca/information-for-vendors/"),
        ("Ontario MFIPPA", "https://www.ontario.ca/laws/statute/90m56"),
        ("Ontario — AODA Website Accessibility", "https://www.ontario.ca/page/how-make-websites-accessible"),
        ("Ontario BPS Cyber Security Strategy", "https://www.ontario.ca/page/ontario-broader-public-sector-cyber-security-strategy-report"),
        ("D2L Brightspace Ontario", "https://www.d2l.com/k-12-ontario/"),
        ("CoSN K-12CVAT", "https://www.cosn.org/tools-and-resources/resource/k-12cvat/"),
        ("EDUCAUSE HECVAT", "https://www.educause.edu/higher-education-community-vendor-assessment-toolkit"),
        ("EdTechPrivacy.ca", "https://www.edtechprivacy.ca/"),
        ("Toronto Data Residency Guidelines", "https://www.toronto.ca/wp-content/uploads/2023/08/966e-Data-Residency-for-Cloud-Technology-Guidelinev1.0.pdf"),
        ("Canadian Data Residency — Pilotcore", "https://pilotcore.io/blog/canadian-data-residency-and-the-public-cloud"),
        ("MEARIE — Cyber Coverage", "https://secure3.mearie.ca/reciprocal-news/how-to-ensure-cyber-coverage-in-your-contracts"),
        ("1EdTech DPSA Template", "https://www.1edtech.org/resource/dpsa"),
        ("IPC Digital Privacy Charter for Ontario Schools", "https://www.ipc.on.ca/en/privacy-organizations/digital-privacy-charter-for-ontario-schools"),
        ("Canada Privacy Commissioners Joint Resolution on EdTech (Oct 2025)", "https://www.priv.gc.ca/en/about-the-opc/what-we-do/provincial-and-territorial-collaboration/joint-resolutions-with-provinces-and-territories/res_20251008_edtech/"),
        ("Bill 194 Analysis — BLG", "https://www.blg.com/en/insights/2024/05/bill-194-the-new-enhancing-digital-security-and-trust-act-2024"),
        ("CoSN Canada State of EdTech 2025", "https://www.cosn.org/tools-and-resources/resource/canada-state-of-edtech-2025/"),
        ("TDSB Digital Learning Tools", "https://www.tdsb.on.ca/Elementary-School/The-Classroom/Technology/Digital-Learning-Tools"),
        ("OECM Ontario Education Collaborative Marketplace", "https://oecm.ca/"),
    ]

    for title, url in sources:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(url)

    # Save
    output_path = "c:/dev/emai/emai-dev-03/docs/DTAP_VASP_Compliance_Report_ClassBridge.docx"
    doc.save(output_path)
    print(f"Report saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_report()
